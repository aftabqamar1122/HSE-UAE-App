import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import re

load_dotenv()

st.set_page_config(page_title="Risk Assessment", page_icon="📋")

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning("Please go to Home and select your Emirate first.")
    st.stop()

color = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(f"📋 Risk Assessment — {color} {emirate}")
st.markdown(
    "**Format: ENGC 14-Column | ADOSH-SF v4.0 | "
    "Hierarchy of Controls**"
)
st.divider()

# ── Helpers ───────────────────────────────────────────────

def set_cell_color(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tcPr.append(borders)


def risk_color(score):
    s = int(score)
    if s >= 16: return "FF0000"
    if s >= 10: return "FFC000"
    if s >= 5:  return "FFFF00"
    return "00AF50"


def risk_label(score):
    s = int(score)
    if s >= 16: return "EXTREME"
    if s >= 10: return "HIGH"
    if s >= 5:  return "MODERATE"
    return "LOW"


def add_cell_text(cell, text, bold=False,
                  size=9, center=False):
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)


def build_docx(meta, rows):
    doc = Document()

    # A4 Landscape
    sec = doc.sections[0]
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21.0)
    sec.left_margin   = Cm(1.27)
    sec.right_margin  = Cm(1.27)
    sec.top_margin    = Cm(1.27)
    sec.bottom_margin = Cm(1.27)

    # ── Company Header ─────────────────────────────────
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run(
        f"{meta['company']} — RISK ASSESSMENT"
    )
    r.bold = True
    r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        f"Project: {meta['project']}   |   "
        f"Activity: {meta['activity']}   |   "
        f"Date: {meta['date']}   |   "
        f"Prepared by: {meta['prepared_by']}   |   "
        f"Location: {meta['location']}   |   "
        f"Emirate: {meta['emirate']}"
    )
    r2.font.size = Pt(9)
    doc.add_paragraph()

    # Column widths
    col_w = [
        Cm(0.93), Cm(1.68), Cm(2.80),
        Cm(2.13), Cm(3.23),
        Cm(0.98), Cm(1.12), Cm(1.12),
        Cm(1.71), Cm(5.81),
        Cm(1.00), Cm(1.07), Cm(1.29), Cm(2.36)
    ]

    tbl = doc.add_table(rows=0, cols=14)
    tbl.style = 'Table Grid'

    # ── SINGLE MERGED HEADER ROW ───────────────────────
    # We use ONE header row with sub-labels built in
    hdr_row = tbl.add_row()
    hdr_cells = hdr_row.cells

    # Set all header cells grey and bold
    col_labels = [
        "S/N",
        "Activity\nElement",
        "Hazards /\nImpact",
        "Who Might Be\nHarmed\nand How?",
        "Risk &\nPotential\nConsequences",
        "Probability\n(P)",
        "Severity\n(S)",
        "Risk Rating\n(R)\nP×S",
        "Initial Risk Level\nLow / Moderate\n/ High / Extreme",
        "Action to be Taken to Reduce the Risk\n"
        "(Hierarchy of Safety Controls)\n"
        "(Include legal & regulatory reference)",
        "Probability\n(P)",
        "Severity\n(S)",
        "Risk Rating\n(R)\nP×S",
        "Residual Risk Level\nLow / Moderate\n/ High / Extreme",
    ]

    # Span label row: cols 5-6-7 under "Risk Classification"
    # and cols 10-11-12 under "Revised Risk Classification"
    # We do this with a 2-row header

    # ── ROW A: Top-level group headers ────────────────
    row_a = tbl.add_row()
    a = row_a.cells

    top_groups = [
        ("S/N", 1),
        ("Activity\nElement", 1),
        ("Hazards /\nImpact", 1),
        ("Who Might Be\nHarmed and How?", 1),
        ("Risk & Potential\nConsequences", 1),
        ("Risk Classification", 3),
        (
            "Initial Risk Level\n"
            "Low / Moderate / High / Extreme", 1
        ),
        (
            "Action to be Taken to Reduce the Risk\n"
            "(Hierarchy of Safety Controls)\n"
            "(Legal & Regulatory Reference)", 1
        ),
        ("Revised Risk\nClassification", 3),
        (
            "Residual Risk Level\n"
            "Low / Moderate / High / Extreme", 1
        ),
    ]

    cur = 0
    for text, span in top_groups:
        cell = a[cur]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(text)
        rn.bold = True
        rn.font.size = Pt(8)
        set_cell_color(cell, "C0C0C0")
        set_borders(cell)
        if span > 1:
            cell.merge(a[cur + span - 1])
        cur += span

    # Remove the blank row we added first
    tbl._tbl.remove(hdr_row._tr)

    # ── ROW B: Sub-headers ─────────────────────────────
    row_b = tbl.add_row()
    b = row_b.cells
    sub_labels = [
        "S/N",
        "Activity\nElement",
        "Hazards /\nImpact",
        "Who Might\nBe Harmed?",
        "Risk &\nPotential\nConsequences",
        "Probability\n(P)",
        "Severity\n(S)",
        "Risk Rating\n(R) P×S",
        "Initial Risk\nLevel",
        "Controls\n(Hierarchy + CoP Ref)",
        "Probability\n(P)",
        "Severity\n(S)",
        "Risk Rating\n(R) P×S",
        "Residual\nRisk Level",
    ]
    for idx, label in enumerate(sub_labels):
        cell = b[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(label)
        rn.bold = True
        rn.font.size = Pt(7)
        set_cell_color(cell, "BFBFBF")
        set_borders(cell)
        cell.width = col_w[idx]

    # ── DATA ROWS ──────────────────────────────────────
    for i, rd in enumerate(rows, start=1):
        ip  = int(rd.get("initial_p", 3))
        is_ = int(rd.get("initial_s", 3))
        ir  = ip * is_
        rp  = int(rd.get("residual_p", 1))
        rs  = int(rd.get("residual_s", 3))
        rr  = rp * rs

        # Build controls lines
        ctrl = rd.get("controls", {})
        lines = []
        order = [
            ("elimination",    "ELIMINATION"),
            ("substitution",   "SUBSTITUTION"),
            ("engineering",    "ENGINEERING CONTROLS"),
            ("administrative", "ADMINISTRATIVE CONTROLS"),
            ("ppe",            "PPE"),
        ]
        if isinstance(ctrl, dict):
            for key, label in order:
                items = ctrl.get(key, [])
                if items:
                    lines.append(f"{label}:")
                    for item in items:
                        lines.append(f"• {item}")
        else:
            lines = [str(ctrl)]

        row = tbl.add_row()
        cells = row.cells
        for idx in range(14):
            cells[idx].width = col_w[idx]

        # Col 0: S/N
        add_cell_text(
            cells[0], str(i),
            bold=True, size=9, center=True
        )
        set_borders(cells[0])

        # Col 1: Activity
        add_cell_text(
            cells[1],
            rd.get("activity", ""),
            bold=True, size=9
        )
        set_borders(cells[1])

        # Col 2: Hazard
        add_cell_text(
            cells[2],
            rd.get("hazard", ""),
            size=9
        )
        set_borders(cells[2])

        # Col 3: Who Harmed
        add_cell_text(
            cells[3],
            rd.get("who_harmed", ""),
            size=9
        )
        set_borders(cells[3])

        # Col 4: Consequences
        add_cell_text(
            cells[4],
            rd.get("consequences", ""),
            size=9
        )
        set_borders(cells[4])

        # Col 5: Initial P
        add_cell_text(
            cells[5], str(ip),
            bold=True, size=10, center=True
        )
        set_borders(cells[5])

        # Col 6: Initial S
        add_cell_text(
            cells[6], str(is_),
            bold=True, size=10, center=True
        )
        set_borders(cells[6])

        # Col 7: Initial R
        add_cell_text(
            cells[7], str(ir),
            bold=True, size=10, center=True
        )
        set_cell_color(cells[7], risk_color(ir))
        set_borders(cells[7])

        # Col 8: Initial Risk Level
        cells[8].text = ""
        p8 = cells[8].paragraphs[0]
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r8 = p8.add_run(risk_label(ir))
        r8.bold = True
        r8.font.size = Pt(9)
        set_cell_color(cells[8], risk_color(ir))
        set_borders(cells[8])

        # Col 9: Controls with hierarchy
        cells[9].text = ""
        first = True
        for line in lines:
            if first:
                para9 = cells[9].paragraphs[0]
                first = False
            else:
                para9 = cells[9].add_paragraph()
            is_hdr = (
                line.endswith(":")
                and line.replace(":", "").strip().isupper()
            )
            rn9 = para9.add_run(line)
            rn9.font.size = Pt(8)
            if is_hdr:
                rn9.bold = True
                rn9.font.color.rgb = RGBColor(192, 0, 0)
        set_borders(cells[9])

        # Col 10: Residual P
        add_cell_text(
            cells[10], str(rp),
            bold=True, size=10, center=True
        )
        set_borders(cells[10])

        # Col 11: Residual S
        add_cell_text(
            cells[11], str(rs),
            bold=True, size=10, center=True
        )
        set_borders(cells[11])

        # Col 12: Residual R
        add_cell_text(
            cells[12], str(rr),
            bold=True, size=10, center=True
        )
        set_cell_color(cells[12], risk_color(rr))
        set_borders(cells[12])

        # Col 13: Residual Risk Level
        cells[13].text = ""
        p13 = cells[13].paragraphs[0]
        p13.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r13 = p13.add_run(risk_label(rr))
        r13.bold = True
        r13.font.size = Pt(9)
        set_cell_color(cells[13], risk_color(rr))
        set_borders(cells[13])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── INPUT FORM ────────────────────────────────────────────
col1, col2 = st.columns(2)
with col1:
    project_name = st.text_input(
        "Project Name",
        placeholder="e.g. Bloom Living Almeria"
    )
    company_name = st.text_input(
        "Company Name", value="ENGC"
    )
    prepared_by = st.text_input(
        "Prepared By",
        placeholder="e.g. Aftab Qamar — Senior HSSE Engineer"
    )
with col2:
    activity = st.text_input(
        "Work Activity",
        placeholder="e.g. Excavation Work"
    )
    ra_date = st.date_input("Date", value=date.today())
    location = st.text_input(
        "Location",
        placeholder="e.g. Plot C13, Zayed City"
    )

st.divider()
work_description = st.text_area(
    "Describe the work activity",
    placeholder="e.g. Excavation for foundation to 2.5m...",
    height=80
)
workers_involved = st.text_input(
    "Who might be harmed?",
    placeholder="e.g. Excavation crew, site workers, public"
)

num_rows = st.slider(
    "Number of activity rows to generate",
    min_value=6,
    max_value=15,
    value=10
)

# ── GENERATE ──────────────────────────────────────────────
if st.button(
    "⚡ Generate Risk Assessment (.docx)",
    type="primary",
    use_container_width=True
):
    if not activity:
        st.error("Please enter a work activity.")
    else:
        with st.spinner(
            f"Generating {num_rows} rows — please wait..."
        ):
            try:
                client = anthropic.Anthropic(
                    api_key=os.getenv("ANTHROPIC_API_KEY")
                )

                if emirate == "Abu Dhabi":
                    reg = (
                        "ADOSH-SF v4.0 — cite specific "
                        "CoP numbers e.g. CoP 29.0, CoP 2.0"
                    )
                else:
                    reg = (
                        "Dubai Municipality Code of "
                        "Construction Safety Practice — "
                        "cite chapter numbers"
                    )

                prompt = f"""You are a senior HSE engineer 
in {emirate}, UAE.

Generate a detailed Risk Assessment for:
Activity: {activity}
Project: {project_name}
Description: {work_description}
Who might be harmed: {workers_involved}
Regulation: {reg}

Return ONLY a JSON array of exactly {num_rows} objects.
No text before or after.
Start immediately with [ and end with ]

Each object must have EXACTLY these keys:
{{
  "activity": "task step name (max 8 words)",
  "hazard": "specific hazard (max 12 words)",
  "consequences": "injury or damage (max 12 words)",
  "who_harmed": "who is at risk (max 8 words)",
  "initial_p": 4,
  "initial_s": 4,
  "controls": {{
    "elimination": ["short measure (max 10 words)"],
    "substitution": ["short measure (max 10 words)"],
    "engineering": [
      "measure with CoP ref (max 12 words)",
      "measure with CoP ref (max 12 words)"
    ],
    "administrative": [
      "measure with CoP ref (max 12 words)",
      "PTW required per CoP 21.0"
    ],
    "ppe": [
      "PPE item per CoP 2.0 (max 10 words)",
      "PPE item per CoP 2.0 (max 10 words)"
    ]
  }},
  "residual_p": 1,
  "residual_s": 3
}}

Cover these different aspects of the activity:
- Site preparation and setup
- Main excavation/construction work
- Equipment operation
- Working at heights (if applicable)
- Electrical hazards
- Manual handling
- Heat stress (UAE climate)
- Traffic and logistics
- Emergency situations
- Environmental hazards
- Chemical/dust exposure
- Lifting operations
- Fire prevention
- Worker welfare
- Final inspection

Keep all string values SHORT to ensure valid JSON.
Return ONLY the JSON array."""

                response = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=8000,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )

                raw = response.content[0].text.strip()

                # Clean markdown
                raw = re.sub(
                    r'```json|```', '', raw
                ).strip()

                # Extract array only
                start = raw.find("[")
                end   = raw.rfind("]") + 1
                if start == -1 or end == 0:
                    raise ValueError(
                        "No JSON array in response"
                    )
                raw = raw[start:end]

                ra_rows = json.loads(raw)

                meta = {
                    "project":     project_name,
                    "company":     company_name,
                    "prepared_by": prepared_by,
                    "date":        ra_date,
                    "activity":    activity,
                    "location":    location,
                    "emirate":     emirate,
                }

                docx_buf = build_docx(meta, ra_rows)

                st.success(
                    f"✅ {len(ra_rows)} rows generated — "
                    "ENGC 14-Column Format!"
                )

                # Preview
                import pandas as pd
                preview = []
                for i, r in enumerate(ra_rows, 1):
                    ip  = int(r.get("initial_p", 3))
                    is_ = int(r.get("initial_s", 3))
                    rp  = int(r.get("residual_p", 1))
                    rs_ = int(r.get("residual_s", 3))
                    preview.append({
                        "#": i,
                        "Activity": r.get("activity",""),
                        "Hazard":   r.get("hazard",""),
                        "P":  ip,  "S":  is_,
                        "R":  ip * is_,
                        "Initial":  risk_label(ip*is_),
                        "P2": rp,  "S2": rs_,
                        "R2": rp * rs_,
                        "Residual": risk_label(rp*rs_),
                    })
                st.dataframe(
                    pd.DataFrame(preview),
                    use_container_width=True
                )

                fname = (
                    f"RA_{activity.replace(' ','_')}"
                    f"_{ra_date}_{emirate}.docx"
                )
                st.download_button(
                    label=(
                        "⬇️ Download Risk Assessment (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument."
                        "wordprocessingml.document"
                    )
                )

            except json.JSONDecodeError as e:
                st.error(
                    f"JSON error: {e}\n"
                    "Please try again — "
                    "AI response was too complex."
                )
            except Exception as e:
                st.error(f"Error: {str(e)}")