import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date, datetime, time
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

load_dotenv()

st.set_page_config(page_title="Incident Report", page_icon="🚨")

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning("Please go to Home and select your Emirate first.")
    st.stop()

color = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(f"🚨 Incident Report — {color} {emirate}")

if emirate == "Abu Dhabi":
    st.markdown("**Reference: ADOSH-SF Version 4.0 | ADPHC**")
    st.info(
        "⚠️ All incidents must be reported to ADOSH "
        "within the required timeframe per ADOSH-SF v4.0"
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality Code — Article 2.4.5**"
    )
    st.error(
        "🚨 DUBAI: Written report to Dubai Municipality "
        "within **72 HOURS**. Records kept **5 YEARS**.\n\n"
        "Police: 999 | Civil Defence: 997 | DM: 800900"
    )

st.divider()

# ═══════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════

def set_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_borders(cell, color="000000", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def cell_para(cell, text, bold=False, size=10,
              color_rgb=None, center=False, italic=False):
    """Clear cell and add formatted paragraph."""
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text) if text else "—")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    return run


def section_title(doc, number, title, fill="1F3864",
                  text_color=(255, 255, 255)):
    """Add a dark blue section header bar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    run = p.add_run(f"  SECTION {number} — {title.upper()}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*text_color)
    return p


def two_col_table(doc, rows_data, widths=(5, 12)):
    """Create a 2-column label/value table."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows_data:
        row = tbl.add_row()
        # Label cell
        lc = row.cells[0]
        lc.width = Cm(widths[0])
        cell_para(lc, label, bold=True, size=9,
                  color_rgb=(0, 0, 0))
        set_shading(lc, "E8F0FE")
        set_borders(lc)
        # Value cell
        vc = row.cells[1]
        vc.width = Cm(widths[1])
        cell_para(vc, value or "—", size=9)
        set_borders(vc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def action_table(doc, actions, title, fill="FFC000"):
    """Create a corrective action table."""
    if not actions:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    r = p.add_run(f"  {title}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)

    hdrs = ["#", "Action Required", "Responsible",
            "Target Date", "Regulation Reference"]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'

    # Header row
    hdr_cells = tbl.rows[0].cells
    widths = [Cm(1), Cm(6.5), Cm(3), Cm(2.5), Cm(4)]
    col_labels = hdrs
    for i, (cell, lbl, w) in enumerate(
        zip(hdr_cells, col_labels, widths)
    ):
        cell.width = w
        cell_para(cell, lbl, bold=True, size=9,
                  color_rgb=(255, 255, 255), center=True)
        set_shading(cell, "1F3864")
        set_borders(cell)

    # Data rows
    for item in actions:
        row = tbl.add_row()
        vals = [
            item.get("no", ""),
            item.get("action", ""),
            item.get("responsible", ""),
            item.get("target", ""),
            item.get("reference", ""),
        ]
        for i, (cell, val, w) in enumerate(
            zip(row.cells, vals, widths)
        ):
            cell.width = w
            cell_para(cell, val, size=9,
                      center=(i == 0))
            set_borders(cell)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def five_why_block(doc, chain_title, steps, root):
    """Format a 5-Why chain as a clean table."""
    p = doc.add_paragraph()
    r = p.add_run(f"  🔴 {chain_title}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(192, 0, 0)
    p.paragraph_format.space_before = Pt(4)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    colors = ["FFF2CC", "FFE0B2", "FFCCCC",
              "FFB3B3", "FF9999"]
    why_labels = ["WHY 1?", "WHY 2?", "WHY 3?",
                  "WHY 4?", "WHY 5?"]
    for i, (step, clr) in enumerate(
        zip(steps, colors)
    ):
        row = tbl.add_row()
        lc = row.cells[0]
        lc.width = Cm(2)
        cell_para(lc, why_labels[i], bold=True, size=9)
        set_shading(lc, clr)
        set_borders(lc)
        vc = row.cells[1]
        vc.width = Cm(14)
        cell_para(vc, step, size=9)
        set_borders(vc)

    # Root cause highlight
    root_row = tbl.add_row()
    rc_lc = root_row.cells[0]
    rc_lc.width = Cm(2)
    cell_para(rc_lc, "✅ ROOT\nCAUSE", bold=True,
              size=9, color_rgb=(255, 255, 255))
    set_shading(rc_lc, "C00000")
    set_borders(rc_lc)
    rc_vc = root_row.cells[1]
    rc_vc.width = Cm(14)
    cell_para(rc_vc, root, bold=True, size=9,
              color_rgb=(192, 0, 0))
    set_shading(rc_vc, "FFE0E0")
    set_borders(rc_vc)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def lessons_table(doc, lessons):
    """Format lessons learned as a table."""
    if not lessons:
        return
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    cell_para(hdr[0], "#", bold=True, size=9,
              color_rgb=(255, 255, 255), center=True)
    set_shading(hdr[0], "1F3864")
    set_borders(hdr[0])
    hdr[0].width = Cm(1)
    cell_para(hdr[1], "Key Lesson", bold=True, size=9,
              color_rgb=(255, 255, 255))
    set_shading(hdr[1], "1F3864")
    set_borders(hdr[1])
    hdr[1].width = Cm(16)

    for i, lesson in enumerate(lessons, 1):
        row = tbl.add_row()
        cell_para(row.cells[0], str(i), bold=True,
                  size=9, center=True)
        set_shading(row.cells[0],
                    "E8F0FE" if i % 2 == 0 else "FFFFFF")
        set_borders(row.cells[0])
        row.cells[0].width = Cm(1)
        cell_para(row.cells[1], lesson, size=9)
        set_shading(row.cells[1],
                    "E8F0FE" if i % 2 == 0 else "FFFFFF")
        set_borders(row.cells[1])
        row.cells[1].width = Cm(16)


# ═══════════════════════════════════════════════════════════
# PARSE AI RESPONSE INTO STRUCTURED SECTIONS
# ═══════════════════════════════════════════════════════════

def parse_ai_sections(text):
    """Extract structured content from AI response."""
    result = {
        "sequence": [],
        "immediate_causes": [],
        "root_chains": [],
        "immediate_actions": [],
        "short_actions": [],
        "long_actions": [],
        "preventive": [],
        "lessons": [],
        "regulatory": "",
    }

    lines = text.split("\n")
    current = None

    for line in lines:
        line = line.strip()
        if not line or line.startswith("```"):
            continue

        low = line.lower()

        # Detect sections
        if "sequence of events" in low:
            current = "sequence"
        elif "immediate cause" in low:
            current = "immediate_causes"
        elif "root cause" in low and "5-why" in low:
            current = "root_chains"
        elif ("immediate action" in low or
              "within 24" in low):
            current = "immediate_actions"
        elif ("short-term" in low or
              "short term" in low or
              "1-2 week" in low):
            current = "short_actions"
        elif ("long-term" in low or
              "long term" in low or
              "1-3 month" in low):
            current = "long_actions"
        elif "preventive" in low:
            current = "preventive"
        elif "lesson" in low:
            current = "lessons"
        elif "regulatory" in low:
            current = "regulatory"
        else:
            # Collect content
            clean = re.sub(
                r'^[#\*\-\|\>]+\s*', '', line
            ).strip()
            if not clean:
                continue

            if current == "sequence" and clean:
                result["sequence"].append(clean)
            elif current == "immediate_causes" and clean:
                result["immediate_causes"].append(clean)
            elif current in [
                "immediate_actions",
                "short_actions",
                "long_actions"
            ] and clean:
                # Try to parse table rows
                if "|" in clean:
                    parts = [
                        p.strip() for p in
                        clean.split("|")
                        if p.strip() and
                        p.strip() not in
                        ["---", "—"]
                    ]
                    if len(parts) >= 4:
                        result[current].append({
                            "no": parts[0],
                            "action": parts[1],
                            "responsible": parts[2]
                                if len(parts) > 2
                                else "",
                            "target": parts[3]
                                if len(parts) > 3
                                else "",
                            "reference": parts[4]
                                if len(parts) > 4
                                else "",
                        })
            elif current == "lessons" and clean:
                result["lessons"].append(clean)
            elif current == "regulatory" and clean:
                result["regulatory"] += clean + "\n"

    return result


# ═══════════════════════════════════════════════════════════
# BUILD DOCX
# ═══════════════════════════════════════════════════════════

def build_incident_docx(data, ai_text, emirate):
    doc = Document()

    # Page: A4 Portrait
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # ── COVER HEADER ──────────────────────────────────────
    # Logo placeholder + Company name
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style = 'Table Grid'
    lc = hdr_tbl.rows[0].cells[0]
    lc.width = Cm(4)
    cell_para(lc, data['company'], bold=True,
              size=18, color_rgb=(31, 56, 100))
    set_shading(lc, "E8F0FE")
    set_borders(lc, "1F3864")

    rc = hdr_tbl.rows[0].cells[1]
    rc.width = Cm(13)
    cell_para(rc, "INCIDENT / ACCIDENT REPORT",
              bold=True, size=16,
              color_rgb=(192, 0, 0), center=True)
    set_shading(rc, "FFF2F2")
    set_borders(rc, "C00000")

    doc.add_paragraph()

    # Incident banner
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = banner._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), "1F3864")
    pPr.append(shd)
    br = banner.add_run(
        f"  Emirate: {emirate}   |   "
        f"Report No: {data['report_no']}   |   "
        f"Date of Incident: {data['incident_date']}   |   "
        f"Classification: {data['incident_type']}  "
    )
    br.bold = True
    br.font.size = Pt(10)
    br.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── SECTION 1: CLASSIFICATION ──────────────────────────
    section_title(doc, "1", "Incident Classification",
                  fill="C00000")
    two_col_table(doc, [
        ("Incident Type",         data['incident_type']),
        ("Injury Classification", data['injury_class']),
        ("Report Number",         data['report_no']),
        ("Date of Incident",      str(data['incident_date'])),
        ("Time of Incident",      str(data['incident_time'])),
        ("Day of Week",           data['day_of_week']),
    ])

    # ── SECTION 2: PROJECT DETAILS ─────────────────────────
    section_title(doc, "2", "Project Details",
                  fill="1F3864")
    two_col_table(doc, [
        ("Project Name",   data['project_name']),
        ("Company",        data['company']),
        ("Location / Site", data['location']),
        ("Reported By",    data['reported_by']),
        ("HSE Officer",    data['hse_officer']),
        ("Report Date",    str(data['report_date'])),
    ])

    # ── SECTION 3: INJURED PERSON ──────────────────────────
    section_title(doc, "3",
                  "Injured / Affected Person Details",
                  fill="1F3864")
    two_col_table(doc, [
        ("Full Name",             data['injured_name']),
        ("Age",                   data['injured_age']),
        ("Nationality",           data['injured_nationality']),
        ("Trade / Designation",   data['injured_trade']),
        ("Employer",              data['injured_employer']),
        ("Experience in Trade",   data['injured_experience']),
        ("Nature of Injury",      data['nature_of_injury']),
        ("Body Part Affected",    data['body_part']),
        ("Medical Treatment",     data['medical_treatment']),
        ("Hospital / Clinic",     data['hospital']),
    ])

    # ── SECTION 4: DESCRIPTION ─────────────────────────────
    section_title(doc, "4", "Incident Description",
                  fill="1F3864")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(
        data['incident_description'] or
        "No description provided."
    )
    run.font.size = Pt(10)
    doc.add_paragraph()

    # ── SECTION 5: IMMEDIATE CAUSES ────────────────────────
    section_title(doc, "5", "Immediate Causes",
                  fill="BF8F00")
    two_col_table(doc, [
        ("Unsafe Act",         data['unsafe_act']),
        ("Unsafe Condition",   data['unsafe_condition']),
        ("Activity at Time",   data['activity_at_time']),
        ("Equipment Involved", data['equipment_involved']),
    ])

    # ── SECTION 6: ROOT CAUSES ─────────────────────────────
    section_title(doc, "6", "Root Causes",
                  fill="BF8F00")
    two_col_table(doc, [
        ("Personal Factors",    data['root_personal']),
        ("Job / Task Factors",  data['root_job']),
        ("Management Factors",  data['root_management']),
    ])

    # ── SECTION 7: AI INVESTIGATION ───────────────────────
    section_title(doc, "7",
                  "AI Investigation Analysis & "
                  "Corrective Actions",
                  fill="1F3864")

    # Parse AI response
    parsed = parse_ai_sections(ai_text)

    # 7.1 Sequence of Events
    if parsed["sequence"]:
        p = doc.add_paragraph()
        r = p.add_run("7.1  Sequence of Events")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(31, 56, 100)

        seq_tbl = doc.add_table(rows=0, cols=3)
        seq_tbl.style = 'Table Grid'
        # Header
        hrow = seq_tbl.add_row()
        for cell, lbl, w in zip(
            hrow.cells,
            ["Step", "Timing", "Description"],
            [Cm(1.5), Cm(3), Cm(12.5)]
        ):
            cell.width = w
            cell_para(cell, lbl, bold=True, size=9,
                      color_rgb=(255, 255, 255),
                      center=True)
            set_shading(cell, "1F3864")
            set_borders(cell)

        for i, seq in enumerate(
            parsed["sequence"][:7], 1
        ):
            drow = seq_tbl.add_row()
            fill = (
                "E8F0FE" if i % 2 == 0 else "FFFFFF"
            )
            for cell, val, w in zip(
                drow.cells,
                [str(i), "—", seq],
                [Cm(1.5), Cm(3), Cm(12.5)]
            ):
                cell.width = w
                cell_para(cell, val, size=9,
                          center=(val == str(i)))
                set_shading(cell, fill)
                set_borders(cell)

        doc.add_paragraph()

    # 7.2 Immediate Causes Summary
    if parsed["immediate_causes"]:
        p = doc.add_paragraph()
        r = p.add_run(
            "7.2  Immediate Causes Analysis"
        )
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(31, 56, 100)

        for item in parsed["immediate_causes"][:8]:
            bp = doc.add_paragraph(style='List Bullet')
            br2 = bp.add_run(item)
            br2.font.size = Pt(9)

        doc.add_paragraph()

    # 7.3 5-Why (simplified — 3 chains max)
    p = doc.add_paragraph()
    r = p.add_run("7.3  Root Cause Analysis (5-Why Method)")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 56, 100)

    # Generic 5-why from root cause data
    chains = [
        {
            "title": "Root Cause Chain — Personal Factor",
            "steps": [
                f"Worker was injured during "
                f"{data.get('activity_at_time','')}.",
                "Worker did not follow safe work procedure.",
                f"Unsafe act identified: "
                f"{data.get('unsafe_act','—')}",
                f"Personal factor: "
                f"{data.get('root_personal','—')}",
                "Insufficient pre-task briefing and "
                "competency verification.",
            ],
            "root": f"ROOT CAUSE: {data.get('root_personal','—')} "
                    f"— Worker competency and awareness gap "
                    f"(Ref: ADOSH CoP 2.0 / ADOSH-SF v4.0)",
        },
        {
            "title": "Root Cause Chain — Management Factor",
            "steps": [
                "Incident occurred despite safety "
                "procedures being in place.",
                f"Unsafe condition: "
                f"{data.get('unsafe_condition','—')}",
                f"Job factor: "
                f"{data.get('root_job','—')}",
                f"Management factor: "
                f"{data.get('root_management','—')}",
                "Management system failed to enforce "
                "controls at task level.",
            ],
            "root": f"ROOT CAUSE: {data.get('root_management','—')} "
                    f"— Management system gap "
                    f"(Ref: ADOSH-SF v4.0 Clause 5.3)",
        },
    ]

    for chain in chains:
        five_why_block(
            doc,
            chain["title"],
            chain["steps"],
            chain["root"]
        )

    # 7.4 Corrective Actions
    p = doc.add_paragraph()
    r = p.add_run("7.4  Corrective & Preventive Actions")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 56, 100)
    doc.add_paragraph()

    # Use parsed actions if available, else defaults
    if parsed["immediate_actions"]:
        ia = parsed["immediate_actions"]
    else:
        ia = [
            {
                "no": "IA-01",
                "action": "Stop work review — "
                          "safety stand-down for all crews",
                "responsible": "HSE Manager",
                "target": "Within 24 hours",
                "reference": "ADOSH-SF v4.0",
            },
            {
                "no": "IA-02",
                "action": "Notify ADOSH of incident "
                          "within required timeframe",
                "responsible": "Project Director",
                "target": "Within 24 hours",
                "reference": "ADOSH-SF v4.0 Incident "
                             "Reporting",
            },
            {
                "no": "IA-03",
                "action": "Preserve incident scene — "
                          "no alteration until investigation",
                "responsible": "HSE Officer",
                "target": "Immediate",
                "reference": "ADOSH-SF v4.0",
            },
        ]

    action_table(
        doc, ia,
        "IMMEDIATE ACTIONS (Within 24–48 Hours)",
        fill="FFC000"
    )

    if parsed["short_actions"]:
        sta = parsed["short_actions"]
    else:
        sta = [
            {
                "no": "ST-01",
                "action": "Review and update risk "
                          "assessment for the activity",
                "responsible": "HSE Manager",
                "target": "7 days",
                "reference": "ADOSH CoP 21.0",
            },
            {
                "no": "ST-02",
                "action": "Conduct toolbox talk — "
                          "share lessons learned with all crews",
                "responsible": "Site HSE Officer",
                "target": "3 days",
                "reference": "ADOSH-SF v4.0",
            },
            {
                "no": "ST-03",
                "action": "Verify PPE compliance "
                          "for all workers on similar tasks",
                "responsible": "HSE Officer",
                "target": "5 days",
                "reference": "ADOSH CoP 2.0",
            },
        ]

    action_table(
        doc, sta,
        "SHORT-TERM ACTIONS (Within 1–2 Weeks)",
        fill="FFE0B2"
    )

    if parsed["long_actions"]:
        lta = parsed["long_actions"]
    else:
        lta = [
            {
                "no": "LT-01",
                "action": "Develop formal supervision "
                          "plan for all high-risk activities",
                "responsible": "Project Manager",
                "target": "30 days",
                "reference": "ADOSH-SF v4.0",
            },
            {
                "no": "LT-02",
                "action": "Enroll workers in formal "
                          "safety training programme",
                "responsible": "Training Manager",
                "target": "45 days",
                "reference": "ADOSH-SF v4.0 Training",
            },
        ]

    action_table(
        doc, lta,
        "LONG-TERM ACTIONS (Within 1–3 Months)",
        fill="C8E6C9"
    )

    # 7.5 Lessons Learned
    p = doc.add_paragraph()
    r = p.add_run("7.5  Key Lessons Learned")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 56, 100)
    doc.add_paragraph()

    if parsed["lessons"]:
        ll = parsed["lessons"][:8]
    else:
        ll = [
            "Always conduct pre-task toolbox talk "
            "before starting high-risk activities.",
            "Permits to Work must be issued and "
            "displayed before work commences.",
            "Supervisors must be present throughout "
            "all high-risk operations.",
            "Workers must immediately report any "
            "unsafe conditions using Stop Work Authority.",
            "PPE is the LAST line of defence — "
            "controls higher in the hierarchy must "
            "be implemented first.",
        ]

    lessons_table(doc, ll)
    doc.add_paragraph()

    # ── SECTION 8: REGULATORY ─────────────────────────────
    section_title(doc, "8",
                  "Regulatory Reporting Requirements",
                  fill="1F3864")

    if emirate == "Abu Dhabi":
        reg_rows = [
            ("Reporting Authority", "ADOSH / ADPHC"),
            ("Framework Reference",
             "ADOSH-SF Version 4.0 (July 2024)"),
            ("Serious Incident Notification",
             "Immediate verbal notification to ADOSH"),
            ("Written Report",
             "Per ADOSH-SF v4.0 requirements"),
            ("Emergency", "Ambulance: 998 | Police: 999"),
        ]
    else:
        reg_rows = [
            ("Reporting Authority",
             "Dubai Municipality — Specialised Dept"),
            ("Legal Reference",
             "DM Code Article 2.4.5"),
            ("Written Report Deadline",
             "Within 72 HOURS of occurrence"),
            ("Police Notification",
             "Call 999 — for serious accidents/LTIs"),
            ("Civil Defence",
             "Call 997 — for fire/explosion/collapse"),
            ("DM Emergency Line",
             "800900"),
            ("Records Retention",
             "5 YEARS from date of occurrence"),
        ]

    two_col_table(doc, reg_rows)

    # ── SECTION 9: WITNESSES ──────────────────────────────
    section_title(doc, "9", "Witnesses", fill="1F3864")
    p = doc.add_paragraph()
    r = p.add_run(
        data.get('witnesses') or "No witnesses recorded."
    )
    r.font.size = Pt(10)
    doc.add_paragraph()

    # ── SECTION 10: SIGNATURES ────────────────────────────
    section_title(doc, "10", "Signatures & Approval",
                  fill="1F3864")
    sig_tbl = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    roles = [
        "Reported By",
        "HSE Officer",
        "Project Manager"
    ]
    names = [
        data.get('reported_by', ''),
        data.get('hse_officer', ''),
        ""
    ]

    # Row 0: Role headers
    for i, role in enumerate(roles):
        cell = sig_tbl.rows[0].cells[i]
        cell.width = Cm(5.67)
        cell_para(cell, role, bold=True, size=10,
                  color_rgb=(255, 255, 255), center=True)
        set_shading(cell, "1F3864")
        set_borders(cell)

    # Row 1: Names
    for i, name in enumerate(names):
        cell = sig_tbl.rows[1].cells[i]
        cell.width = Cm(5.67)
        cell_para(cell, name or "— ", size=10,
                  center=True)
        set_shading(cell, "E8F0FE")
        set_borders(cell)

    # Row 2: Signature lines
    for i in range(3):
        cell = sig_tbl.rows[2].cells[i]
        cell.width = Cm(5.67)
        cell_para(
            cell,
            "\n\nSignature: _______________________"
            "\n\nDate: ____________________________",
            size=9
        )
        set_borders(cell)

    doc.add_paragraph()

    # ── FOOTER NOTE ───────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = footer_p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), "F2F2F2")
    pPr.append(shd)
    fr = footer_p.add_run(
        f"  {data['company']} — Confidential HSE Document  "
        f"|  Generated: {date.today()}  "
        f"|  Emirate: {emirate}  "
        f"|  Ref: {data['report_no']}  "
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════
# INCIDENT TYPE OPTIONS
# ═══════════════════════════════════════════════════════════

if emirate == "Abu Dhabi":
    incident_types = [
        "Near Miss",
        "First Aid Injury (FAI)",
        "Medical Treatment Case (MTC)",
        "Restricted Work Case (RWC)",
        "Lost Time Injury (LTI)",
        "Permanent Partial Disability (PPD)",
        "Permanent Total Disability (PTD)",
        "Fatality",
        "Property Damage",
        "Environmental Incident",
        "Fire / Explosion",
        "Dangerous Occurrence",
    ]
    injury_classes = [
        "Near Miss — No injury",
        "First Aid Only",
        "Medical Treatment (no lost time)",
        "Restricted Work (modified duties)",
        "Lost Time Injury (>1 day away from work)",
        "Serious Injury",
        "Fatality",
        "Property / Equipment Damage Only",
    ]
else:
    incident_types = [
        "Near Miss",
        "Minor Injury (First Aid)",
        "Lost Time Injury (LTI — >3 days)",
        "Serious / Major Injury",
        "Fatality",
        "Property Damage",
        "Fire / Explosion",
        "Structural Collapse",
        "Environmental Incident",
        "Dangerous Occurrence",
    ]
    injury_classes = [
        "Minor Injury (First Aid only)",
        "Lost Time Injury (>3 days absence)",
        "Serious Injury — fracture/amputation/burns",
        "Fatality",
        "Property Damage Only",
        "No Injury (Near Miss)",
    ]

# ═══════════════════════════════════════════════════════════
# FORM
# ═══════════════════════════════════════════════════════════

with st.form("incident_form"):

    st.subheader("🔴 Section 1 — Incident Classification")
    c1, c2 = st.columns(2)
    with c1:
        incident_type = st.selectbox(
            "Incident Type", incident_types
        )
        incident_date = st.date_input(
            "Date of Incident", value=date.today()
        )
        report_no = st.text_input(
            "Report Number",
            value=(
                f"IR-{emirate[:2].upper()}-"
                f"{datetime.now().strftime('%Y%m%d-%H%M')}"
            )
        )
    with c2:
        injury_class = st.selectbox(
            "Injury Classification", injury_classes
        )
        incident_time = st.time_input(
            "Time of Incident", value=time(8, 0)
        )
        day_of_week = st.selectbox(
            "Day of Week",
            ["Monday", "Tuesday", "Wednesday",
             "Thursday", "Friday", "Saturday", "Sunday"]
        )

    st.divider()
    st.subheader("📋 Section 2 — Project Details")
    c3, c4 = st.columns(2)
    with c3:
        project_name = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria"
        )
        company = st.text_input("Company", value="ENGC")
        location = st.text_input(
            "Location / Site",
            placeholder="e.g. Plot C13, Zayed City"
        )
    with c4:
        reported_by = st.text_input(
            "Reported By",
            placeholder="Name and designation"
        )
        hse_officer = st.text_input(
            "HSE Officer",
            placeholder="e.g. Aftab Qamar"
        )
        report_date = st.date_input(
            "Report Date", value=date.today()
        )

    st.divider()
    st.subheader(
        "👷 Section 3 — Injured / Affected Person"
    )
    c5, c6 = st.columns(2)
    with c5:
        injured_name = st.text_input(
            "Full Name of Injured Person"
        )
        injured_age = st.text_input("Age")
        injured_nationality = st.text_input("Nationality")
        injured_trade = st.text_input(
            "Trade / Designation",
            placeholder="e.g. Carpenter, Electrician"
        )
        injured_employer = st.text_input(
            "Employer / Sub-contractor"
        )
    with c6:
        injured_experience = st.text_input(
            "Experience in Trade",
            placeholder="e.g. 3 years"
        )
        nature_of_injury = st.text_input(
            "Nature of Injury",
            placeholder="e.g. Laceration to right hand"
        )
        body_part = st.selectbox(
            "Body Part Affected",
            ["Head", "Eye", "Neck", "Shoulder",
             "Arm", "Hand", "Finger", "Back",
             "Chest", "Abdomen", "Leg", "Knee",
             "Foot", "Multiple", "None"]
        )
        medical_treatment = st.selectbox(
            "Medical Treatment Given",
            ["First Aid on Site",
             "Referred to Clinic",
             "Referred to Hospital",
             "Admitted to Hospital",
             "None Required"]
        )
        hospital = st.text_input(
            "Hospital / Clinic Name (if applicable)"
        )

    st.divider()
    st.subheader("📝 Section 4 — Incident Description")
    incident_description = st.text_area(
        "Describe what happened in detail "
        "(sequence of events)",
        height=150,
        placeholder=(
            "e.g. At approximately 09:30hrs, while the "
            "worker was carrying out excavation works at "
            "a depth of 1.8m, the trench wall collapsed..."
        )
    )

    st.divider()
    st.subheader("⚡ Section 5 — Immediate Causes")
    c7, c8 = st.columns(2)
    with c7:
        unsafe_act = st.text_area(
            "Unsafe Act",
            height=80,
            placeholder="e.g. Working in unshored trench"
        )
        activity_at_time = st.text_input(
            "Activity at Time of Incident",
            placeholder="e.g. Excavation works"
        )
    with c8:
        unsafe_condition = st.text_area(
            "Unsafe Condition",
            height=80,
            placeholder="e.g. No shoring, no edge protection"
        )
        equipment_involved = st.text_input(
            "Equipment / Material Involved",
            placeholder="e.g. Excavator, hand tools"
        )

    st.divider()
    st.subheader("🔍 Section 6 — Root Causes")
    c9, c10 = st.columns(2)
    with c9:
        root_personal = st.text_area(
            "Personal Factors", height=80,
            placeholder="e.g. Lack of training"
        )
        root_job = st.text_area(
            "Job / Task Factors", height=80,
            placeholder="e.g. No method statement"
        )
    with c10:
        root_management = st.text_area(
            "Management System Factors", height=80,
            placeholder="e.g. Permit to Work not issued"
        )

    st.divider()
    witnesses = st.text_area(
        "👁️ Witnesses",
        height=60,
        placeholder="e.g. Ahmed Ali — Foreman"
    )

    submitted = st.form_submit_button(
        "🚨 Generate Incident Report (.docx)",
        type="primary",
        use_container_width=True
    )

# ═══════════════════════════════════════════════════════════
# GENERATE REPORT
# ═══════════════════════════════════════════════════════════

if submitted:
    if not incident_description:
        st.error("Please provide an incident description.")
    else:
        with st.spinner(
            "Generating professional incident report..."
        ):
            try:
                client = anthropic.Anthropic(
                    api_key=os.getenv("ANTHROPIC_API_KEY")
                )

                reg_ref = (
                    "ADOSH-SF Version 4.0, ADOSH CoPs"
                    if emirate == "Abu Dhabi"
                    else "Dubai Municipality Code Art. 2.4.5"
                )

                ai_prompt = f"""You are a senior HSE 
investigation expert in {emirate}, UAE.

Analyse this incident:
TYPE: {incident_type}
ACTIVITY: {activity_at_time}
DESCRIPTION: {incident_description}
UNSAFE ACT: {unsafe_act}
UNSAFE CONDITION: {unsafe_condition}
ROOT CAUSES: Personal={root_personal}; 
Job={root_job}; Management={root_management}
REGULATION: {reg_ref}

Provide structured analysis with these EXACT sections
(use these exact headings):

## SEQUENCE OF EVENTS
(7 numbered steps as table rows, one per line,
format: Step N | Timing | Description)

## IMMEDIATE CAUSES
(bullet list of unsafe acts and conditions)

## ROOT CAUSE ANALYSIS (5-WHY METHOD)
(3 why-chains with root causes)

## IMMEDIATE ACTIONS
(table rows: No | Action | Responsible | Target | Ref)
Format each row as: IA-01 | action text | person | 
24hrs | CoP ref

## SHORT-TERM ACTIONS
(table rows same format, prefix ST-01, ST-02...)

## LONG-TERM ACTIONS
(table rows same format, prefix LT-01, LT-02...)

## LESSONS LEARNED
(5-8 bullet points, one per line)

## REGULATORY COMPLIANCE
(paragraph about {emirate} reporting obligations)

Keep each cell value under 15 words.
Be specific to {emirate} and cite CoP/Chapter numbers."""

                response = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=3000,
                    messages=[
                        {
                            "role": "user",
                            "content": ai_prompt
                        }
                    ]
                )

                ai_text = response.content[0].text

                data = {
                    "report_no":    report_no,
                    "incident_type": incident_type,
                    "injury_class": injury_class,
                    "incident_date": incident_date,
                    "incident_time": incident_time,
                    "day_of_week":  day_of_week,
                    "project_name": project_name,
                    "company":      company,
                    "location":     location,
                    "reported_by":  reported_by,
                    "hse_officer":  hse_officer,
                    "report_date":  report_date,
                    "injured_name": injured_name,
                    "injured_age":  injured_age,
                    "injured_nationality":
                        injured_nationality,
                    "injured_trade":    injured_trade,
                    "injured_employer": injured_employer,
                    "injured_experience":
                        injured_experience,
                    "nature_of_injury": nature_of_injury,
                    "body_part":    body_part,
                    "medical_treatment": medical_treatment,
                    "hospital":     hospital,
                    "incident_description":
                        incident_description,
                    "unsafe_act":   unsafe_act,
                    "unsafe_condition": unsafe_condition,
                    "activity_at_time": activity_at_time,
                    "equipment_involved":
                        equipment_involved,
                    "root_personal":    root_personal,
                    "root_job":         root_job,
                    "root_management":  root_management,
                    "witnesses":        witnesses,
                }

                docx_buf = build_incident_docx(
                    data, ai_text, emirate
                )

                st.success(
                    "✅ Professional Incident Report "
                    "generated!"
                )

                # Preview AI analysis
                with st.expander(
                    "📊 View AI Analysis (Preview)",
                    expanded=False
                ):
                    st.markdown(ai_text)

                fname = (
                    f"Incident_Report_"
                    f"{report_no}_"
                    f"{incident_date}.docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download Incident Report (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument."
                        "wordprocessingml.document"
                    )
                )

                if emirate == "Dubai":
                    st.error(
                        "⚠️ REMINDER: Submit written report "
                        "to Dubai Municipality within "
                        "**72 HOURS**. DM: 800900"
                    )

            except Exception as e:
                st.error(f"Error: {str(e)}")