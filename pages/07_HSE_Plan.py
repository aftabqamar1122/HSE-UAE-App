import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64

load_dotenv()

st.set_page_config(
    page_title="HSE Plan Generator",
    page_icon="📄"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning(
        "Please go to Home and select "
        "your Emirate first."
    )
    st.stop()

color = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(
    "📄 HSE Plan Generator — "
    + color + " " + emirate
)

if emirate == "Abu Dhabi":
    st.markdown(
        "**Reference: ADOSH CoP 53.1 — "
        "OSH Construction Management Plan | "
        "ADOSH-SF Version 4.0**"
    )
    st.info(
        "Generates a full OSH Construction "
        "Management Plan as required by "
        "ADOSH CoP 53.1 — must be approved "
        "before work commences."
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality "
        "Code Art. 2.3.2 — Safety Plan**"
    )
    st.info(
        "Generates a Construction Safety Plan "
        "as required by Dubai Municipality — "
        "must be approved by Engineer before "
        "construction commences."
    )

st.divider()


# ── DOCX HELPERS ──────────────────────────────────────────

def set_cell_color(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_borders(cell):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '1F3864')
        borders.append(b)
    tcPr.append(borders)


def add_heading(doc, text, level=1,
                color_hex="1F3864"):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    pPr.append(shd)
    run           = p.add_run("  " + text)
    run.bold      = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p


def add_info_row(tbl, label, value,
                 label_color="E8F0FE"):
    row = tbl.add_row()
    lc  = row.cells[0]
    vc  = row.cells[1]
    lc.width = Cm(5)
    vc.width = Cm(12)
    lc.text  = label
    lc.paragraphs[0].runs[0].bold      = True
    lc.paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_color(lc, label_color)
    set_borders(lc)
    vc.text = str(value) if value else "—"
    vc.paragraphs[0].runs[0].font.size = Pt(10)
    set_borders(vc)


def add_body_text(doc, text, size=10,
                  bold=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    run           = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold      = bold
    return p


def add_bullet(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def build_hse_plan_docx(meta, sections, emirate):
    doc = Document()

    # A4 Portrait
    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # ── COVER PAGE ────────────────────────────
    cover           = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = cover._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1F3864')
    pPr.append(shd)
    em_icon = "🔵" if emirate == "Abu Dhabi" else "🔴"
    r = cover.add_run(
        "\n\n\n"
        + em_icon + "  "
        + meta['company']
        + "\n\n"
    )
    r.bold           = True
    r.font.size      = Pt(22)
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc_title = (
        "OSH CONSTRUCTION MANAGEMENT PLAN"
        if emirate == "Abu Dhabi"
        else "CONSTRUCTION SAFETY PLAN"
    )

    cover2           = doc.add_paragraph()
    cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = cover2._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'),   'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'),  'C00000')
    pPr2.append(shd2)
    r2 = cover2.add_run(
        "\n  " + doc_title + "  \n\n"
        "  " + meta['project'] + "  \n\n"
    )
    r2.bold           = True
    r2.font.size      = Pt(18)
    r2.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_page_break()

    # ── DOCUMENT CONTROL TABLE ────────────────
    add_heading(
        doc,
        "DOCUMENT CONTROL INFORMATION",
        color_hex="1F3864"
    )
    ctrl_tbl       = doc.add_table(rows=0, cols=2)
    ctrl_tbl.style = 'Table Grid'

    if emirate == "Abu Dhabi":
        reg_ref = "ADOSH CoP 53.1 | ADOSH-SF v4.0"
    else:
        reg_ref = "Dubai Municipality Code Art.2.3.2"

    ctrl_rows = [
        ("Document Title",        doc_title),
        ("Project Name",          meta['project']),
        ("Project Location",      meta['location']),
        ("Client",                meta['client']),
        ("Consultant / Engineer", meta['consultant']),
        ("Main Contractor",       meta['company']),
        ("Contract No.",          meta['contract_no']),
        ("Document No.",          meta['doc_no']),
        ("Revision",              meta['revision']),
        ("Prepared By",           meta['prepared_by']),
        ("Approved By",           meta['approved_by']),
        ("Date",                  str(meta['doc_date'])),
        ("Regulatory Reference",  reg_ref),
    ]
    for lbl, val in ctrl_rows:
        add_info_row(ctrl_tbl, lbl, val)

    doc.add_paragraph()
    doc.add_page_break()

    # ── TABLE OF CONTENTS ─────────────────────
    add_heading(
        doc, "TABLE OF CONTENTS",
        color_hex="1F3864"
    )
    toc_items = [
        "1.0   Introduction & Project Overview",
        "2.0   Scope of Work",
        "3.0   HSE Policy Statement",
        "4.0   Regulatory & Legal Requirements",
        "5.0   HSE Organisation & Responsibilities",
        "6.0   Hazard Identification & Risk Assessment",
        "7.0   Emergency Response Plan",
        "8.0   Permit to Work System",
        "9.0   HSE Training & Competency",
        "10.0  Personal Protective Equipment",
        "11.0  Occupational Health & Worker Welfare",
        "12.0  Environmental Management",
        "13.0  Incident Reporting & Investigation",
        "14.0  HSE Inspections & Audits",
        "15.0  Sub-contractor Management",
        "16.0  Traffic & Logistics Management",
        "17.0  Communication & Consultation",
        "18.0  Document Control",
        "19.0  Performance Monitoring & KPIs",
        "20.0  Signatures & Approvals",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    # ── MAIN SECTIONS ─────────────────────────
    section_colors = {
        "1":  "1F3864", "2":  "1F3864",
        "3":  "C00000", "4":  "1F3864",
        "5":  "1F5864", "6":  "BF8F00",
        "7":  "C00000", "8":  "C00000",
        "9":  "1F3864", "10": "1F3864",
        "11": "1B5E20", "12": "1B5E20",
        "13": "C00000", "14": "1F3864",
        "15": "1F3864", "16": "1F3864",
        "17": "1F3864", "18": "1F3864",
        "19": "1F3864", "20": "1F3864",
    }

    for section in sections:
        num   = str(section.get("number", ""))
        title = section.get("title", "")
        body  = section.get("body", "")
        color = section_colors.get(
            num.split(".")[0], "1F3864"
        )
        add_heading(
            doc,
            num + "  " + title.upper(),
            color_hex=color
        )
        if isinstance(body, list):
            for item in body:
                item = str(item).strip()
                if item.startswith(("•", "-")):
                    add_bullet(
                        doc,
                        item.lstrip("•- ").strip()
                    )
                else:
                    add_body_text(doc, item)
        else:
            for para in str(body).split("\n"):
                para = para.strip()
                if not para:
                    continue
                if para.startswith(("•", "-")):
                    add_bullet(
                        doc,
                        para.lstrip("•- ").strip()
                    )
                elif para.startswith("**"):
                    add_body_text(
                        doc,
                        para.strip("*"),
                        bold=True
                    )
                else:
                    add_body_text(doc, para)
        doc.add_paragraph()

    # ── SIGNATURE BLOCK ───────────────────────
    add_heading(
        doc,
        "20.0  SIGNATURES & APPROVAL",
        color_hex="1F3864"
    )
    sig_tbl       = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    roles = [
        "Prepared By\n(HSE Officer)",
        "Reviewed By\n(Project Manager)",
        "Approved By\n(Projects Director)"
    ]
    for i, role in enumerate(roles):
        c = sig_tbl.rows[0].cells[i]
        c.text = role
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].runs[0].font.color.rgb = (
            RGBColor(255, 255, 255)
        )
        set_cell_color(c, "1F3864")
        set_borders(c)
    for i in range(3):
        c      = sig_tbl.rows[1].cells[i]
        c.text = "\n\nName: _______________\n"
        c.paragraphs[0].runs[0].font.size = Pt(10)
        set_borders(c)
        c      = sig_tbl.rows[2].cells[i]
        c.text = (
            "Date: _______________\n\n"
            "Signature: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(10)
        set_borders(c)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════
# MAIN UI — TWO TABS
# ════════════════════════════════════════════════════════

tab_gen, tab_upload = st.tabs([
    "⚡ Generate HSE Plan (AI)",
    "📤 Upload Your Own Format",
])


# ════════════════════════════════════════════════════════
# TAB 1 — AI GENERATOR
# ════════════════════════════════════════════════════════

with tab_gen:
    st.markdown(
        "### 🤖 AI-Powered HSE Plan Generator"
    )
    st.info(
        "Fill in your project details below. "
        "The AI will generate a complete, "
        "professional "
        + ("OSH Construction Management Plan"
           if emirate == "Abu Dhabi"
           else "Construction Safety Plan")
        + " with all 20 sections — "
        "ready for consultant submission."
    )

    st.subheader("📋 Project Information")
    c1, c2 = st.columns(2)
    with c1:
        project    = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="hp_proj"
        )
        location   = st.text_input(
            "Project Location",
            placeholder="e.g. Zayed City, Abu Dhabi",
            key="hp_loc"
        )
        client     = st.text_input(
            "Client",
            placeholder="e.g. Bloom District Properties",
            key="hp_client"
        )
        consultant = st.text_input(
            "Consultant / Engineer",
            placeholder="e.g. Dar Al-Handasah",
            key="hp_cons"
        )
    with c2:
        company     = st.text_input(
            "Main Contractor",
            value="ENGC",
            key="hp_co"
        )
        contract_no = st.text_input(
            "Contract Number",
            placeholder="e.g. BL-2024-001",
            key="hp_contractno"
        )
        doc_no      = st.text_input(
            "Document Number",
            placeholder="e.g. ENGC-HSE-PLAN-001",
            key="hp_docno"
        )
        revision    = st.selectbox(
            "Revision",
            ["Rev 00", "Rev 01",
             "Rev 02", "Rev 03"],
            key="hp_rev"
        )

    st.divider()
    c3, c4 = st.columns(2)
    with c3:
        prepared_by = st.text_input(
            "Prepared By",
            placeholder=(
                "e.g. Aftab Qamar — "
                "Senior HSSE Engineer"
            ),
            key="hp_prep"
        )
        approved_by = st.text_input(
            "Approved By",
            placeholder=(
                "e.g. Rami Kamal — "
                "Projects Director"
            ),
            key="hp_approve"
        )
    with c4:
        doc_date      = st.date_input(
            "Document Date",
            value=date.today(),
            key="hp_date"
        )
        project_value = st.text_input(
            "Approximate Project Value",
            placeholder="e.g. AED 50 Million",
            key="hp_value"
        )

    st.divider()
    st.subheader("🏗️ Project Scope Details")
    c5, c6 = st.columns(2)
    with c5:
        project_type = st.multiselect(
            "Project Type",
            [
                "Residential Villas",
                "Community Center",
                "Service Station",
                "Commercial Building",
                "Infrastructure Works",
                "Fit-out Works",
                "Demolition",
                "MEP Works",
            ],
            default=["Residential Villas"],
            key="hp_type"
        )
        peak_workers = st.number_input(
            "Peak Workforce (number of workers)",
            min_value=10,
            max_value=5000,
            value=150,
            key="hp_workers"
        )
    with c6:
        project_duration = st.text_input(
            "Project Duration",
            placeholder=(
                "e.g. 24 months "
                "(Jan 2025 – Dec 2026)"
            ),
            key="hp_duration"
        )
        high_risk = st.multiselect(
            "High-Risk Activities on Site",
            [
                "Excavation Works",
                "Working at Heights",
                "Scaffolding Erection",
                "Lifting Operations (Crane)",
                "Precast Installation",
                "Formwork & Falsework",
                "Hot Work (Welding/Cutting)",
                "Electrical Works",
                "Confined Space Entry",
                "Demolition Works",
                "Underground Services",
                "Traffic Management",
                "Chemical Handling",
                "Concrete Pumping",
            ],
            default=[
                "Excavation Works",
                "Working at Heights",
                "Lifting Operations (Crane)",
                "Scaffolding Erection",
            ],
            key="hp_risks"
        )

    st.divider()
    st.subheader("📌 Key Personnel")
    c7, c8 = st.columns(2)
    with c7:
        hse_manager = st.text_input(
            "Project HSE Manager / Officer",
            placeholder="e.g. Aftab Qamar",
            key="hp_hse"
        )
        pm_name     = st.text_input(
            "Project Manager",
            placeholder="e.g. Ahmad Abdelrahman",
            key="hp_pm"
        )
    with c8:
        pd_name      = st.text_input(
            "Projects Director",
            placeholder="e.g. Rami Kamal Yassin",
            key="hp_pd"
        )
        site_address = st.text_input(
            "Site Address / Plot Number",
            placeholder="e.g. Plot C13, Zayed City",
            key="hp_site"
        )

    special_notes = st.text_area(
        "Special Site Conditions / Notes (optional)",
        placeholder=(
            "e.g. Site adjacent to live road, "
            "groundwater expected at 2m depth..."
        ),
        height=80,
        key="hp_notes"
    )

    if st.button(
        "⚡ Generate Complete HSE Plan (.docx)",
        type="primary",
        use_container_width=True,
        key="hp_gen_btn"
    ):
        if not project:
            st.error("Please enter a project name.")
        else:
            prog = st.progress(
                0, text="Starting HSE Plan..."
            )
            try:
                client_api = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                if emirate == "Abu Dhabi":
                    reg_ref = (
                        "ADOSH-SF Version 4.0 "
                        "(July 2024), ADOSH CoP 53.1, "
                        "ADOSH CoP 53.0, "
                        "MOHRE Resolution 44/2022"
                    )
                    auth = "ADOSH / ADPHC"
                else:
                    reg_ref = (
                        "Dubai Municipality Code of "
                        "Construction Safety Practice, "
                        "Articles 2.3.2, 2.4.5, "
                        "Local Order 61/1991"
                    )
                    auth = "Dubai Municipality"

                activities_str = ", ".join(high_risk)
                type_str       = ", ".join(project_type)

                sections_to_generate = [
                    ("1.0",  "Introduction & Project Overview"),
                    ("2.0",  "Scope of Work"),
                    ("3.0",  "HSE Policy Statement"),
                    ("4.0",  "Regulatory & Legal Requirements"),
                    ("5.0",  "HSE Organisation & Responsibilities"),
                    ("6.0",  "Hazard Identification & Risk Assessment"),
                    ("7.0",  "Emergency Response Plan"),
                    ("8.0",  "Permit to Work System"),
                    ("9.0",  "HSE Training & Competency"),
                    ("10.0", "Personal Protective Equipment"),
                    ("11.0", "Occupational Health & Worker Welfare"),
                    ("12.0", "Environmental Management"),
                    ("13.0", "Incident Reporting & Investigation"),
                    ("14.0", "HSE Inspections & Audits"),
                    ("15.0", "Sub-contractor Management"),
                    ("16.0", "Traffic & Logistics Management"),
                    ("17.0", "Communication & Consultation"),
                    ("18.0", "Document Control"),
                    ("19.0", "Performance Monitoring & KPIs"),
                ]

                all_sections = []
                total        = len(sections_to_generate)

                for i, (num, title) in enumerate(
                    sections_to_generate
                ):
                    pct = int((i / total) * 88)
                    prog.progress(
                        pct,
                        text=(
                            "Generating Section "
                            + num + " — " + title
                            + "..."
                        )
                    )

                    plan_type = (
                        "OSH Construction Management Plan"
                        if emirate == "Abu Dhabi"
                        else "Construction Safety Plan"
                    )

                    prompt = (
                        "You are a senior HSE professional "
                        "writing a formal " + plan_type
                        + " for " + emirate + ", UAE.\n\n"
                        "PROJECT DETAILS:\n"
                        "- Project: " + project + "\n"
                        "- Location: " + location + "\n"
                        "- Client: " + client + "\n"
                        "- Consultant: " + consultant + "\n"
                        "- Main Contractor: " + company + "\n"
                        "- Project Type: " + type_str + "\n"
                        "- Duration: " + project_duration + "\n"
                        "- Peak Workers: "
                        + str(peak_workers) + "\n"
                        "- High-Risk Activities: "
                        + activities_str + "\n"
                        "- HSE Officer: " + hse_manager + "\n"
                        "- Project Manager: " + pm_name + "\n"
                        "- Projects Director: " + pd_name + "\n"
                        "- Special Notes: " + special_notes + "\n\n"
                        "REGULATORY REFERENCE: " + reg_ref + "\n"
                        "AUTHORITY: " + auth + "\n\n"
                        "Write SECTION " + num
                        + " — " + title.upper()
                        + " for this specific project.\n\n"
                        "Requirements:\n"
                        "1. Write professionally as a REAL "
                        "HSE Plan (not a template)\n"
                        "2. Reference specific "
                        + emirate + " regulations throughout\n"
                        "3. Use actual project details provided\n"
                        "4. Include specific procedures, "
                        "responsibilities, and requirements\n"
                        "5. For Abu Dhabi: cite specific "
                        "ADOSH CoP numbers. "
                        "For Dubai: cite DM Code chapters\n"
                        "6. Make it comprehensive — this will "
                        "be submitted to consultant/authority\n"
                        "7. Format with clear paragraphs "
                        "and bullet points where appropriate\n"
                        "8. Length: 250-400 words\n\n"
                        "Write ONLY the section content. "
                        "No section number or title in your "
                        "response. Start directly with content."
                    )

                    response = client_api.messages.create(
                        model="claude-sonnet-4-6",
                        max_tokens=1000,
                        messages=[{
                            "role": "user",
                            "content": prompt
                        }]
                    )

                    all_sections.append({
                        "number": num,
                        "title":  title,
                        "body":   response.content[0].text
                    })

                prog.progress(
                    92,
                    text="Building Word document..."
                )

                meta = {
                    "project":     project,
                    "location":    location,
                    "client":      client,
                    "consultant":  consultant,
                    "company":     company,
                    "contract_no": contract_no,
                    "doc_no":      doc_no,
                    "revision":    revision,
                    "prepared_by": prepared_by,
                    "approved_by": approved_by,
                    "doc_date":    doc_date,
                }

                docx_buf = build_hse_plan_docx(
                    meta, all_sections, emirate
                )

                prog.progress(
                    100,
                    text="✅ HSE Plan complete!"
                )

                st.success(
                    "✅ Complete HSE Plan generated! "
                    + str(len(all_sections))
                    + " sections — Ready for submission."
                )

                # Preview
                with st.expander(
                    "📋 Preview Generated Sections",
                    expanded=False
                ):
                    for s in all_sections:
                        st.markdown(
                            "### "
                            + s['number']
                            + " " + s['title']
                        )
                        st.markdown(s['body'])
                        st.divider()

                # Build filename safely
                proj_clean = project.replace(' ', '_')
                em_clean   = emirate.replace(' ', '_')
                fname = (
                    "HSE_Plan_"
                    + proj_clean + "_"
                    + em_clean + "_"
                    + str(doc_date) + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download Complete "
                        "HSE Plan (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="hp_download"
                )

                if emirate == "Abu Dhabi":
                    st.info(
                        "📋 **Next Steps (Abu Dhabi):**"
                        "\n\n1. Review and customize "
                        "for your specific site\n\n"
                        "2. Submit to Consultant "
                        "(Dar Al-Handasah) for review\n\n"
                        "3. Consultant submits to ADOSH\n\n"
                        "4. Keep approved copy on site\n\n"
                        "**Ref: ADOSH CoP 53.1**"
                    )
                else:
                    st.info(
                        "📋 **Next Steps (Dubai):**"
                        "\n\n1. Engineer reviews and "
                        "approves\n\n"
                        "2. Submit to Dubai Municipality\n\n"
                        "3. Safety meetings every 15 days "
                        "(DM Code Art. 2.2.8)\n\n"
                        "**Ref: DM Code Art. 2.3.2**"
                    )

            except Exception as e:
                prog.progress(0)
                st.error("Error: " + str(e))


# ════════════════════════════════════════════════════════
# TAB 2 — UPLOAD YOUR OWN FORMAT
# ════════════════════════════════════════════════════════

def build_filled_docx(sections_text, meta,
                       emirate, plan_name):
    """
    Build a properly formatted .docx from
    AI-generated filled content.
    Preserves professional look with navy headers.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # A4 Portrait
    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    def shd_para(p, fill):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        pPr.append(shd)

    def shd_cell(cell, fill):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        tcPr.append(shd)

    def cell_borders(cell):
        tc      = cell._tc
        tcPr    = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    '6')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '1F3864')
            borders.append(b)
        tcPr.append(borders)

    # ── COVER PAGE ────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(cover, '1F3864')
    r = cover.add_run(
        "\n\n  "
        + meta.get('company', 'ENGC')
        + "\n\n"
    )
    r.bold           = True
    r.font.size      = Pt(22)
    r.font.color.rgb = RGBColor(255, 255, 255)

    cover2 = doc.add_paragraph()
    cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(cover2, 'C00000')
    r2 = cover2.add_run(
        "\n  " + plan_name.upper() + "  \n\n"
        "  " + meta.get('project', '') + "  \n"
        "  " + meta.get('location', '') + "  \n\n"
    )
    r2.bold           = True
    r2.font.size      = Pt(18)
    r2.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_page_break()

    # ── DOCUMENT CONTROL ──────────────────────
    ctrl_p = doc.add_paragraph()
    shd_para(ctrl_p, '1F3864')
    cr = ctrl_p.add_run(
        "  DOCUMENT CONTROL INFORMATION"
    )
    cr.bold           = True
    cr.font.size      = Pt(12)
    cr.font.color.rgb = RGBColor(255, 255, 255)

    ctrl_tbl       = doc.add_table(rows=0, cols=2)
    ctrl_tbl.style = 'Table Grid'
    ctrl_rows = [
        ("Document Title",      plan_name),
        ("Project Name",        meta.get('project','')),
        ("Location",            meta.get('location','')),
        ("Client",              meta.get('client','')),
        ("Consultant",          meta.get('consultant','')),
        ("Main Contractor",     meta.get('company','ENGC')),
        ("HSE Officer",         meta.get('hse_officer','')),
        ("Project Manager",     meta.get('pm','')),
        ("Document No.",        meta.get('doc_no','')),
        ("Revision",            meta.get('revision','Rev 00')),
        ("Date",                str(meta.get('doc_date', date.today()))),
        ("Prepared By",         meta.get('hse_officer','')),
        ("Regulatory Ref.",
         "ADOSH-SF v4.0 | ADOSH CoP 53.1"
         if emirate == "Abu Dhabi"
         else "Dubai Municipality Code Art.2.3.2"),
    ]
    for lbl, val in ctrl_rows:
        row  = ctrl_tbl.add_row()
        lc   = row.cells[0]
        vc   = row.cells[1]
        lc.width = Cm(5)
        vc.width = Cm(12)
        lc.text  = lbl
        lc.paragraphs[0].runs[0].bold      = True
        lc.paragraphs[0].runs[0].font.size = Pt(10)
        shd_cell(lc, 'E8F0FE')
        cell_borders(lc)
        vc.text = str(val) if val else "—"
        vc.paragraphs[0].runs[0].font.size = Pt(10)
        cell_borders(vc)

    doc.add_page_break()

    # ── MAIN CONTENT FROM AI ──────────────────
    # Parse the AI response into sections
    lines = sections_text.split('\n')
    current_section = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_section:
                current_section.append('')
            continue

        # Detect section headings
        is_heading = False
        heading_fill = '1F3864'

        # Numbered sections like "1.0", "2.0"
        import re
        if re.match(
            r'^\d+\.?\d*\s+[A-Z]', stripped
        ) or re.match(
            r'^(SECTION|CHAPTER)\s+\d+',
            stripped.upper()
        ):
            is_heading = True
            # Color based on content
            low = stripped.lower()
            if any(w in low for w in [
                'emergency', 'incident',
                'accident', 'fire'
            ]):
                heading_fill = 'C00000'
            elif any(w in low for w in [
                'environment', 'waste'
            ]):
                heading_fill = '1B5E20'
            elif any(w in low for w in [
                'hazard', 'risk', 'permit'
            ]):
                heading_fill = 'BF8F00'

        if is_heading:
            # Add heading paragraph
            p   = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            shd_para(p, heading_fill)
            run           = p.add_run(
                "  " + stripped.upper()
            )
            run.bold      = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(
                255, 255, 255
            )

        elif stripped.startswith(('•', '-', '*',
                                   '–')):
            # Bullet point
            p   = doc.add_paragraph(
                style='List Bullet'
            )
            run = p.add_run(
                stripped.lstrip('•-*– ').strip()
            )
            run.font.size = Pt(10)

        elif re.match(r'^\d+[\.\)]\s', stripped):
            # Numbered list
            p   = doc.add_paragraph(
                style='List Number'
            )
            run = p.add_run(
                re.sub(r'^\d+[\.\)]\s+', '',
                       stripped)
            )
            run.font.size = Pt(10)

        elif (stripped.startswith('**')
              and stripped.endswith('**')):
            # Bold text
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.5)
            run      = p.add_run(
                stripped.strip('*')
            )
            run.bold      = True
            run.font.size = Pt(10)

        else:
            # Normal paragraph
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.5)
            run           = p.add_run(stripped)
            run.font.size = Pt(10)

    doc.add_page_break()

    # ── SIGNATURE BLOCK ───────────────────────
    sig_p = doc.add_paragraph()
    shd_para(sig_p, '1F3864')
    sr = sig_p.add_run(
        "  SIGNATURES & APPROVAL"
    )
    sr.bold           = True
    sr.font.size      = Pt(12)
    sr.font.color.rgb = RGBColor(255, 255, 255)

    sig_tbl       = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    roles = [
        "Prepared By\n(HSE Officer)",
        "Reviewed By\n(Project Manager)",
        "Approved By\n(Projects Director)"
    ]
    for i, role in enumerate(roles):
        c      = sig_tbl.rows[0].cells[i]
        c.text = role
        r      = c.paragraphs[0].runs[0]
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        shd_cell(c, '1F3864')
        cell_borders(c)

    for i in range(3):
        c      = sig_tbl.rows[1].cells[i]
        c.text = "\n\nName: _______________\n"
        c.paragraphs[0].runs[0].font.size = Pt(10)
        cell_borders(c)
        c      = sig_tbl.rows[2].cells[i]
        c.text = (
            "Date: _______________\n\n"
            "Signature: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(10)
        cell_borders(c)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


with tab_upload:
    st.markdown(
        "### 📤 Upload Your Own HSE Plan Format"
    )
    st.info(
        "Upload your company's existing HSE Plan "
        "or OSH Construction Management Plan "
        "(Word or PDF). The AI will read your "
        "EXACT format — headings, sections, tables, "
        "layout — and fill it with content for your "
        "specific project, keeping your company "
        "branding and structure. "
        "Downloads as a proper **.docx** file."
    )

    uploaded_file = st.file_uploader(
        "📎 Upload Your HSE Plan Template or Sample",
        type=["docx", "pdf"],
        help=(
            "Upload a blank template or a previously "
            "completed HSE Plan. Word (.docx) gives "
            "the best format matching. PDF also works."
        ),
        key="hp_upload"
    )

    if uploaded_file:
        st.success(
            "✅ File uploaded: **"
            + uploaded_file.name + "** ("
            + "{:,}".format(uploaded_file.size)
            + " bytes)"
        )

        st.divider()
        st.subheader("📋 Project Details to Fill In")

        u1, u2 = st.columns(2)
        with u1:
            u_project    = st.text_input(
                "Project Name",
                placeholder="e.g. Bloom Living Almeria",
                key="u_proj"
            )
            u_location   = st.text_input(
                "Project Location",
                placeholder="e.g. Zayed City, Abu Dhabi",
                key="u_loc"
            )
            u_client     = st.text_input(
                "Client",
                placeholder="e.g. Bloom District",
                key="u_client"
            )
            u_consultant = st.text_input(
                "Consultant",
                placeholder="e.g. Dar Al-Handasah",
                key="u_cons"
            )
        with u2:
            u_company  = st.text_input(
                "Main Contractor",
                value="ENGC",
                key="u_co"
            )
            u_hse      = st.text_input(
                "HSE Officer Name",
                placeholder="e.g. Aftab Qamar",
                key="u_hse"
            )
            u_pm       = st.text_input(
                "Project Manager",
                placeholder="e.g. Ahmad Abdelrahman",
                key="u_pm"
            )
            u_workers  = st.number_input(
                "Peak Workforce",
                min_value=10,
                max_value=5000,
                value=150,
                key="u_workers"
            )

        u_risks = st.multiselect(
            "High-Risk Activities on Site",
            [
                "Excavation Works",
                "Working at Heights",
                "Scaffolding Erection",
                "Lifting Operations (Crane)",
                "Precast Installation",
                "Formwork & Falsework",
                "Hot Work (Welding/Cutting)",
                "Electrical Works",
                "Confined Space Entry",
                "Demolition Works",
            ],
            default=[
                "Excavation Works",
                "Working at Heights",
            ],
            key="u_risks"
        )

        u_notes = st.text_area(
            "Special Conditions / Notes (optional)",
            placeholder=(
                "e.g. Site adjacent to live road, "
                "groundwater expected, night works..."
            ),
            height=70,
            key="u_notes"
        )

        if st.button(
            "⚡ Generate HSE Plan Using My Format",
            type="primary",
            use_container_width=True,
            key="u_gen_btn"
        ):
            if not u_project:
                st.error(
                    "Please enter a project name."
                )
            else:
                with st.spinner(
                    "Reading your format and "
                    "generating HSE Plan content..."
                ):
                    try:
                        client_api = (
                            anthropic.Anthropic(
                                api_key=os.getenv(
                                    "ANTHROPIC_API_KEY"
                                )
                            )
                        )

                        file_bytes    = (
                            uploaded_file.read()
                        )
                        risks_str     = (
                            ", ".join(u_risks)
                        )

                        if emirate == "Abu Dhabi":
                            reg = (
                                "ADOSH-SF v4.0, "
                                "ADOSH CoP 53.1, "
                                "cite specific CoP numbers"
                            )
                            plan_name = (
                                "OSH Construction "
                                "Management Plan"
                            )
                        else:
                            reg = (
                                "Dubai Municipality Code "
                                "Art. 2.3.2, "
                                "cite chapter numbers"
                            )
                            plan_name = (
                                "Construction Safety Plan"
                            )

                        fill_instruction = (
                            "This is a " + plan_name
                            + " template/format for "
                            + emirate + ", UAE.\n\n"
                            "Fill it with content for:\n"
                            "- Project: " + u_project + "\n"
                            "- Location: " + u_location + "\n"
                            "- Client: " + u_client + "\n"
                            "- Consultant: " + u_consultant + "\n"
                            "- Main Contractor: " + u_company + "\n"
                            "- HSE Officer: " + u_hse + "\n"
                            "- Project Manager: " + u_pm + "\n"
                            "- Peak Workers: "
                            + str(u_workers) + "\n"
                            "- High-Risk Activities: "
                            + risks_str + "\n"
                            "- Special Notes: " + u_notes + "\n"
                            "- Regulatory Reference: " + reg + "\n\n"
                            "IMPORTANT INSTRUCTIONS:\n"
                            "1. Keep the EXACT same headings, "
                            "section numbers, table structures, "
                            "and layout as in the uploaded document\n"
                            "2. Replace ALL placeholder text, "
                            "[brackets], and blank fields with "
                            "real, specific content\n"
                            "3. Reference specific "
                            + emirate
                            + " regulations throughout "
                            "(ADOSH CoP numbers or "
                            "DM Code chapters)\n"
                            "4. Make it professional and "
                            "ready for submission\n"
                            "5. Use the actual project details "
                            "provided above throughout\n\n"
                            "Write the complete filled document."
                        )

                        if uploaded_file.name\
                                .lower()\
                                .endswith('.pdf'):
                            # PDF — send as document
                            b64 = (
                                base64.standard_b64encode(
                                    file_bytes
                                ).decode()
                            )
                            response = (
                                client_api.messages
                                .create(
                                model="claude-sonnet-4-6",
                                max_tokens=4000,
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "document",
                                            "source": {
                                                "type": "base64",
                                                "media_type": "application/pdf",
                                                "data": b64,
                                            }
                                        },
                                        {
                                            "type": "text",
                                            "text": fill_instruction
                                        }
                                    ]
                                }]
                            ))

                        else:
                            # DOCX — extract text
                            doc_temp  = Document(
                                BytesIO(file_bytes)
                            )
                            extracted = "\n".join([
                                p.text
                                for p in doc_temp.paragraphs
                                if p.text.strip()
                            ])

                            # Also get table content
                            table_text = ""
                            for tbl in doc_temp.tables:
                                for row in tbl.rows:
                                    row_text = " | ".join([
                                        c.text.strip()
                                        for c in row.cells
                                        if c.text.strip()
                                    ])
                                    if row_text:
                                        table_text += (
                                            row_text + "\n"
                                        )

                            full_content = (
                                extracted
                                + ("\n\nTABLE CONTENT:\n"
                                   + table_text
                                   if table_text else "")
                            )

                            response = (
                                client_api.messages
                                .create(
                                model="claude-sonnet-4-6",
                                max_tokens=4000,
                                messages=[{
                                    "role": "user",
                                    "content": (
                                        "Here is the content of my "
                                        + plan_name
                                        + " template:\n\n"
                                        "---\n"
                                        + full_content
                                        + "\n---\n\n"
                                        + fill_instruction
                                    )
                                }]
                            ))

                        filled_content = (
                            response.content[0].text
                        )

                        st.success(
                            "✅ HSE Plan filled for "
                            + u_project
                            + " — building .docx..."
                        )

                        # Build proper docx
                        doc_meta = {
                            "project":     u_project,
                            "location":    u_location,
                            "client":      u_client,
                            "consultant":  u_consultant,
                            "company":     u_company,
                            "hse_officer": u_hse,
                            "pm":          u_pm,
                            "doc_no":      "",
                            "revision":    "Rev 00",
                            "doc_date":    date.today(),
                        }

                        filled_docx = build_filled_docx(
                            filled_content,
                            doc_meta,
                            emirate,
                            plan_name
                        )

                        proj_fn = u_project.replace(
                            ' ', '_'
                        )
                        em_fn = emirate.replace(
                            ' ', '_'
                        )
                        dl_fname = (
                            "HSE_Plan_"
                            + proj_fn + "_"
                            + em_fn + "_"
                            + str(date.today())
                            + ".docx"
                        )

                        st.download_button(
                            label=(
                                "⬇️ Download "
                                + plan_name
                                + " (.docx)"
                            ),
                            data=filled_docx,
                            file_name=dl_fname,
                            mime=(
                                "application/vnd"
                                ".openxmlformats-"
                                "officedocument"
                                ".wordprocessingml"
                                ".document"
                            ),
                            key="u_dl"
                        )

                        # Also show preview
                        with st.expander(
                            "📋 Preview Generated Content",
                            expanded=False
                        ):
                            st.markdown(filled_content)

                        st.info(
                            "✅ **Downloaded as .docx**"
                            " — professional Word "
                            "document with cover page, "
                            "document control table, "
                            "all sections, and signature "
                            "block.\n\n"
                            "📋 Review and have it "
                            "signed before submission."
                        )

                    except Exception as e:
                        st.error(
                            "Error: " + str(e)
                        )

    else:
        # Tips when no file uploaded
        st.markdown(
            "#### 💡 What You Can Upload:"
        )
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("""
**Accepted Formats:**
- ✅ Word Document (.docx) — Best
- ✅ PDF (.pdf)

**What to Upload:**
- Your company HSE Plan template
- A previously approved HSE Plan
- Consultant-provided HSE Plan format
- ADOSH or DM standard HSE Plan format
""")
        with col_b:
            st.markdown("""
**What the AI Will Do:**
- Read ALL your headings and sections
- Fill every section with real content
- Use correct Emirates regulations
- Include your project details throughout
- Reference specific CoP/Chapter numbers
- Keep your EXACT layout and structure
""")

        st.divider()
        st.markdown(
            "#### 📁 Don't Have a Template?"
        )
        st.info(
            "Use the **'Generate HSE Plan (AI)'** "
            "tab to create a complete, professional "
            + ("OSH Construction Management Plan"
               if emirate == "Abu Dhabi"
               else "Construction Safety Plan")
            + " from scratch — fully formatted as a "
            ".docx file, ready for submission."
        )