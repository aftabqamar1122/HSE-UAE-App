import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import json
import re

load_dotenv()

st.set_page_config(
    page_title="Method Statements",
    page_icon="📋"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning(
        "Please go to Home and select "
        "your Emirate first."
    )
    st.stop()

color = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(
    f"📋 Method Statement Generator "
    f"— {color} {emirate}"
)

if emirate == "Abu Dhabi":
    st.markdown(
        "**Reference: ADOSH-SF v4.0 | "
        "ADOSH CoP 53.1 | Activity-Specific CoPs**"
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality "
        "Code of Construction Safety Practice**"
    )

st.divider()


# ── DOCX HELPERS ──────────────────────────────────────────

def set_cell_color(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_borders(cell, color="1F3864"):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def section_header(doc, text, fill="1F3864"):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)
    run      = p.add_run("  " + text.upper())
    run.bold = True
    run.font.size      = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)


def add_info_row(tbl, label, value,
                 lbl_color="E8F0FE"):
    row = tbl.add_row()
    lc  = row.cells[0]
    vc  = row.cells[1]
    lc.width = Cm(5)
    vc.width = Cm(12)
    lc.text  = label
    p        = lc.paragraphs[0]
    p.runs[0].bold      = True
    p.runs[0].font.size = Pt(10)
    set_cell_color(lc, lbl_color)
    set_borders(lc)
    vc.text = str(value) if value else "—"
    vc.paragraphs[0].runs[0].font.size = Pt(10)
    set_borders(vc)


def add_body(doc, text, size=10,
             bold=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run      = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold


def add_bullet(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(str(text))
    run.font.size = Pt(size)


def make_table_header(tbl, headers, fill,
                      sizes=None):
    """Add colored header row to a table."""
    row   = tbl.rows[0]
    cells = row.cells
    for i, (cell, hdr) in enumerate(
        zip(cells, headers)
    ):
        cell.text = hdr
        run       = cell.paragraphs[0].runs[0]
        run.bold  = True
        run.font.size      = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_color(cell, fill)
        set_borders(cell)
        if sizes and i < len(sizes):
            cell.width = sizes[i]


def build_ms_docx(meta, content, emirate):
    doc = Document()

    # A4 Portrait
    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # ── HEADER BAR ────────────────────────────
    hdr           = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = hdr._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1F3864')
    pPr.append(shd)
    r           = hdr.add_run(
        "\n  " + meta['company']
        + "  —  METHOD STATEMENT\n  "
        + meta['activity'] + "  \n"
    )
    r.bold           = True
    r.font.size      = Pt(16)
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── DOCUMENT INFO TABLE ───────────────────
    section_header(
        doc, "1. DOCUMENT INFORMATION",
        fill="1F3864"
    )
    info_tbl       = doc.add_table(rows=0, cols=2)
    info_tbl.style = 'Table Grid'

    if emirate == "Abu Dhabi":
        reg_ref = (
            "ADOSH-SF v4.0 | "
            + meta.get('cop_ref', '')
        )
    else:
        reg_ref = (
            "Dubai Municipality Code | "
            + meta.get('cop_ref', '')
        )

    info_rows = [
        ("Method Statement Title", meta['activity']),
        ("Project",               meta['project']),
        ("Location",              meta['location']),
        ("Client",                meta['client']),
        ("Consultant",            meta['consultant']),
        ("Main Contractor",       meta['company']),
        ("Document No.",          meta['doc_no']),
        ("Revision",              meta['revision']),
        ("Date",                  str(meta['ms_date'])),
        ("Prepared By",           meta['prepared_by']),
        ("Reviewed By",           meta['reviewed_by']),
        ("Approved By",           meta['approved_by']),
        ("Regulatory Reference",  reg_ref),
    ]
    for lbl, val in info_rows:
        add_info_row(info_tbl, lbl, val)

    doc.add_paragraph()

    # ── SCOPE ─────────────────────────────────
    section_header(
        doc, "2. SCOPE OF WORK",
        fill="1F3864"
    )
    add_body(doc, content.get(
        'scope',
        'Refer to project specifications.'
    ))
    doc.add_paragraph()

    # ── PERSONNEL ─────────────────────────────
    section_header(
        doc, "3. PERSONNEL & RESPONSIBILITIES",
        fill="1F3864"
    )
    resp_tbl       = doc.add_table(rows=1, cols=3)
    resp_tbl.style = 'Table Grid'
    make_table_header(
        resp_tbl,
        ["Role", "Name / Title", "Responsibility"],
        "1F3864",
        [Cm(4), Cm(4), Cm(9)]
    )
    for person in content.get('personnel', []):
        row  = resp_tbl.add_row()
        vals = [
            person.get('role', ''),
            person.get('name', ''),
            person.get('responsibility', '')
        ]
        for i, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            set_borders(c)

    doc.add_paragraph()

    # ── EQUIPMENT ─────────────────────────────
    section_header(
        doc,
        "4. PLANT, EQUIPMENT & MATERIALS",
        fill="1F5864"
    )
    eq_tbl       = doc.add_table(rows=1, cols=3)
    eq_tbl.style = 'Table Grid'
    make_table_header(
        eq_tbl,
        ["Item", "Specification / Type",
         "Quantity"],
        "1F5864",
        [Cm(5), Cm(8), Cm(4)]
    )
    for i, item in enumerate(
        content.get('equipment', [])
    ):
        row  = eq_tbl.add_row()
        vals = [
            item.get('item', ''),
            item.get('spec', ''),
            item.get('qty', '')
        ]
        for c, v in zip(row.cells, vals):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_color(c, "EEF6FF")
            set_borders(c)

    doc.add_paragraph()

    # ── PPE ───────────────────────────────────
    section_header(
        doc,
        "5. PERSONAL PROTECTIVE EQUIPMENT (PPE)",
        fill="1B5E20"
    )
    ppe_tbl       = doc.add_table(rows=1, cols=3)
    ppe_tbl.style = 'Table Grid'
    make_table_header(
        ppe_tbl,
        ["PPE Item", "Standard / Reference",
         "Mandatory For"],
        "1B5E20",
        [Cm(5), Cm(6), Cm(6)]
    )
    for i, item in enumerate(
        content.get('ppe', [])
    ):
        row  = ppe_tbl.add_row()
        vals = [
            item.get('item', ''),
            item.get('standard', ''),
            item.get(
                'mandatory_for', 'All workers'
            )
        ]
        for c, v in zip(row.cells, vals):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_color(c, "E8F5E9")
            set_borders(c)

    doc.add_paragraph()

    # ── SEQUENCE OF WORK ──────────────────────
    section_header(
        doc,
        "6. SEQUENCE OF WORK / METHOD",
        fill="BF8F00"
    )
    seq_tbl       = doc.add_table(rows=1, cols=3)
    seq_tbl.style = 'Table Grid'
    make_table_header(
        seq_tbl,
        ["Step", "Activity / Task",
         "Key Safety Requirements"],
        "BF8F00",
        [Cm(1.5), Cm(7), Cm(8.5)]
    )
    for i, step in enumerate(
        content.get('sequence', [])
    ):
        row  = seq_tbl.add_row()
        vals = [
            str(i + 1),
            step.get('task', ''),
            step.get('safety_req', '')
        ]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                set_cell_color(c, "FFF9E6")
            set_borders(c)

    doc.add_paragraph()

    # ── HAZARDS ───────────────────────────────
    section_header(
        doc,
        "7. HAZARD IDENTIFICATION & RISK SUMMARY",
        fill="C00000"
    )
    haz_tbl       = doc.add_table(rows=1, cols=4)
    haz_tbl.style = 'Table Grid'
    make_table_header(
        haz_tbl,
        ["Hazard", "Risk",
         "Controls", "Residual Risk"],
        "C00000",
        [Cm(3.5), Cm(3), Cm(7), Cm(3.5)]
    )
    risk_colors = {
        "LOW":      "00AF50",
        "MODERATE": "FFFF00",
        "HIGH":     "FFC000",
        "EXTREME":  "FF0000",
    }
    for i, haz in enumerate(
        content.get('hazards', [])
    ):
        row = haz_tbl.add_row()
        rr  = haz.get('residual_risk', 'LOW')
        vals = [
            haz.get('hazard', ''),
            haz.get('risk', ''),
            haz.get('controls', ''),
            rr
        ]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 3:
                set_cell_color(
                    c,
                    risk_colors.get(
                        rr.upper(), "00AF50"
                    )
                )
                c.paragraphs[0].runs[0].bold = True
            elif i % 2 == 0:
                set_cell_color(c, "FFF2F2")
            set_borders(c)

    doc.add_paragraph()

    # ── EMERGENCY ─────────────────────────────
    section_header(
        doc, "8. EMERGENCY PROCEDURES",
        fill="C00000"
    )
    if emirate == "Abu Dhabi":
        prep_name = meta.get(
            'prepared_by', 'HSE Officer'
        )
        em_lines = [
            "STOP WORK immediately — raise alarm",
            "Call: Ambulance 998 | Police 999 "
            "| Civil Defence 997",
            "Notify HSE Officer: " + prep_name,
            "Evacuate to designated Muster Point",
            "Do not re-enter until area is safe",
            "Report to ADOSH per ADOSH-SF v4.0",
        ]
    else:
        prep_name = meta.get(
            'prepared_by', 'HSE Officer'
        )
        em_lines = [
            "STOP WORK immediately — raise alarm",
            "Call: Police 999 | Civil Defence 997 "
            "| Ambulance 998 | DM 800900",
            "Notify HSE Officer: " + prep_name,
            "Evacuate to designated Muster Point",
            "Do not re-enter until area is safe",
            "Report to DM within 72 hrs "
            "(DM Code Art. 2.4.5)",
        ]
    for line in em_lines:
        add_bullet(doc, line)

    doc.add_paragraph()

    # ── PERMITS ───────────────────────────────
    section_header(
        doc, "9. PERMITS REQUIRED",
        fill="4A148C"
    )
    permits = content.get('permits', [])
    if permits:
        ptw_tbl       = doc.add_table(
            rows=1, cols=3
        )
        ptw_tbl.style = 'Table Grid'
        make_table_header(
            ptw_tbl,
            ["Permit Type", "Required When",
             "Reference"],
            "4A148C",
            [Cm(5), Cm(7), Cm(5)]
        )
        for i, p in enumerate(permits):
            row  = ptw_tbl.add_row()
            vals = [
                p.get('type', ''),
                p.get('when', ''),
                p.get('ref', '')
            ]
            for c, v in zip(row.cells, vals):
                c.text = v
                c.paragraphs[0].runs[0].font.size\
                    = Pt(9)
                if i % 2 == 0:
                    set_cell_color(c, "F3E5F5")
                set_borders(c)
    else:
        add_body(doc, "No permits required "
                      "for this activity.")

    doc.add_paragraph()

    # ── CHECKLIST ─────────────────────────────
    section_header(
        doc, "10. PRE-WORK INSPECTION CHECKLIST",
        fill="0D47A1"
    )
    chk_tbl       = doc.add_table(rows=1, cols=3)
    chk_tbl.style = 'Table Grid'
    make_table_header(
        chk_tbl,
        ["Inspection Item",
         "Checked By", "Status"],
        "0D47A1",
        [Cm(10), Cm(4), Cm(3)]
    )
    for i, item in enumerate(
        content.get('checklist', [])
    ):
        row  = chk_tbl.add_row()
        vals = [item, "HSE Officer",
                "☐ Yes  ☐ No"]
        for c, v in zip(row.cells, vals):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_color(c, "E8F0FE")
            set_borders(c)

    doc.add_paragraph()

    # ── REFERENCES ────────────────────────────
    section_header(
        doc, "11. REFERENCES & STANDARDS",
        fill="1F3864"
    )
    for ref in content.get('references', []):
        add_bullet(doc, ref)

    doc.add_paragraph()

    # ── SIGNATURES ────────────────────────────
    section_header(
        doc, "12. SIGNATURES & APPROVAL",
        fill="1F3864"
    )
    sig_tbl       = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    roles = [
        "Prepared By\n(HSE Officer)",
        "Reviewed By\n(Project Manager)",
        "Approved By\n(Projects Director)"
    ]
    for i, role in enumerate(roles):
        c           = sig_tbl.rows[0].cells[i]
        c.text      = role
        run         = c.paragraphs[0].runs[0]
        run.bold    = True
        run.font.size      = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_color(c, "1F3864")
        set_borders(c)

    for i in range(3):
        c      = sig_tbl.rows[1].cells[i]
        c.text = "\n\nName: _______________\n"
        c.paragraphs[0].runs[0].font.size = Pt(10)
        set_borders(c)

        c      = sig_tbl.rows[2].cells[i]
        c.text = (
            "Date: _______________\n\n"
            "Signature: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(10)
        set_borders(c)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════
# ACTIVITIES LIST
# ════════════════════════════════════════════════════════

activities = {
    "🏗️ Civil & Structural": [
        "Excavation & Earthworks",
        "Dewatering Operations",
        "Foundation Works & Piling",
        "Concrete Pouring & Pumping",
        "Formwork & Falsework (Shuttering)",
        "Reinforcement Steel Fixing",
        "Precast Concrete Installation",
        "Masonry & Blockwork",
        "Steel Erection & Structural Steel",
        "Demolition Works",
        "Underground Construction",
        "Waterproofing Works",
    ],
    "🏠 Finishing Works": [
        "Plastering & Rendering",
        "Tiling & Stone Works",
        "Painting & Decorating",
        "False Ceiling Installation",
        "Flooring Installation",
        "Glazing & Curtain Wall Installation",
        "Carpentry & Joinery Works",
    ],
    "⚡ MEP Works": [
        "Electrical Installation Works",
        "Plumbing & Drainage Works",
        "HVAC Installation",
        "Fire Fighting System Installation",
        "Cable Pulling & Termination",
        "Generator Installation",
        "Solar Panel Installation",
    ],
    "🔧 Temporary Works": [
        "Scaffolding Erection & Dismantling",
        "Hoarding & Fencing Installation",
        "Site Accommodation Setup",
        "Crane Erection & Dismantling",
        "Traffic Diversion & Management",
        "Temporary Electrical Installation",
    ],
    "⚠️ High Risk Activities": [
        "Working at Heights (>1.8m)",
        "Confined Space Entry",
        "Hot Work (Welding & Cutting)",
        "Abrasive Blasting",
        "Chemical Handling & Storage",
        "Radiography / NDT Works",
        "LOTO (Lockout-Tagout)",
        "Lifting Operations (Crane & Rigging)",
        "Overhead Line Work",
        "Underground Services Works",
        "Spray Painting & Finishing",
    ],
    "🚛 Logistics & Transport": [
        "Material Delivery & Unloading",
        "Heavy Equipment Mobilization",
        "Material Handling & Storage",
        "Waste Management & Disposal",
    ],
}


# ════════════════════════════════════════════════════════
# MAIN UI — TWO TABS
# ════════════════════════════════════════════════════════

tab_gen, tab_upload = st.tabs([
    "⚡ Generate Method Statement (AI)",
    "📤 Upload Your Own Format",
])


# ════════════════════════════════════════════════════════
# TAB 1 — AI GENERATOR
# ════════════════════════════════════════════════════════

with tab_gen:
    st.markdown("### 🤖 AI-Powered Method Statement")
    st.info(
        "Select any activity. The AI generates a "
        "complete, professional Method Statement "
        "with " + emirate + " regulatory references "
        "— ready to submit to your consultant."
    )

    cat      = st.selectbox(
        "Select Category",
        list(activities.keys()),
        key="ms_category"
    )
    activity = st.selectbox(
        "Select Activity",
        activities[cat],
        key="ms_activity"
    )
    custom = st.text_input(
        "Or type a custom activity "
        "(leave blank to use selection above)",
        placeholder="e.g. Pile Cap Construction",
        key="ms_custom"
    )
    final_activity = custom if custom else activity

    st.divider()
    st.subheader("📋 Project Information")

    c1, c2 = st.columns(2)
    with c1:
        ms_project = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="ms_proj"
        )
        ms_location = st.text_input(
            "Location",
            placeholder="e.g. Zayed City, Abu Dhabi",
            key="ms_loc"
        )
        ms_client = st.text_input(
            "Client",
            placeholder="e.g. Bloom District",
            key="ms_client"
        )
        ms_consultant = st.text_input(
            "Consultant",
            placeholder="e.g. Dar Al-Handasah",
            key="ms_cons"
        )
    with c2:
        ms_company = st.text_input(
            "Main Contractor",
            value="ENGC",
            key="ms_co"
        )
        ms_doc_no = st.text_input(
            "Document No.",
            placeholder="e.g. ENGC-MS-EXC-001",
            key="ms_docno"
        )
        ms_rev = st.selectbox(
            "Revision",
            ["Rev 00", "Rev 01",
             "Rev 02", "Rev 03"],
            key="ms_rev"
        )
        ms_date = st.date_input(
            "Date",
            value=date.today(),
            key="ms_date"
        )

    c3, c4 = st.columns(2)
    with c3:
        ms_prep = st.text_input(
            "Prepared By",
            placeholder=(
                "e.g. Aftab Qamar — "
                "Senior HSSE Engineer"
            ),
            key="ms_prep"
        )
        ms_review = st.text_input(
            "Reviewed By",
            placeholder=(
                "e.g. Ahmad Abdelrahman "
                "— Project Manager"
            ),
            key="ms_review"
        )
    with c4:
        ms_approve = st.text_input(
            "Approved By",
            placeholder=(
                "e.g. Rami Kamal — "
                "Projects Director"
            ),
            key="ms_approve"
        )
        ms_workers = st.number_input(
            "No. of Workers for this Activity",
            min_value=1,
            max_value=500,
            value=20,
            key="ms_workers"
        )

    ms_special = st.text_area(
        "Special Conditions / Notes (optional)",
        placeholder=(
            "e.g. Adjacent to live road, "
            "night works required..."
        ),
        height=70,
        key="ms_special"
    )

    if st.button(
        "⚡ Generate Method Statement (.docx)",
        type="primary",
        use_container_width=True,
        key="ms_gen_btn"
    ):
        if not ms_project:
            st.error("Please enter a project name.")
        else:
            prog = st.progress(
                0, text="Starting..."
            )
            try:
                client_api = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                reg = (
                    "ADOSH-SF v4.0, relevant ADOSH "
                    "CoPs (cite specific CoP numbers)"
                    if emirate == "Abu Dhabi"
                    else
                    "Dubai Municipality Code of "
                    "Construction Safety Practice "
                    "(cite specific chapters)"
                )

                prog.progress(
                    15,
                    text="AI generating content..."
                )

                prepared_name = (
                    ms_prep if ms_prep
                    else "Aftab Qamar"
                )

                prompt = (
                    "You are a senior HSE engineer "
                    "in " + emirate + ", UAE.\n\n"
                    "Generate a complete Method "
                    "Statement for:\n"
                    "ACTIVITY: " + final_activity + "\n"
                    "PROJECT: " + ms_project + "\n"
                    "LOCATION: " + ms_location + "\n"
                    "CLIENT: " + ms_client + "\n"
                    "CONSULTANT: " + ms_consultant + "\n"
                    "CONTRACTOR: " + ms_company + "\n"
                    "WORKERS: " + str(ms_workers) + "\n"
                    "SPECIAL: " + ms_special + "\n"
                    "REGULATION: " + reg + "\n\n"
                    "Return ONLY valid JSON. "
                    "No markdown. No text before or "
                    "after. Start with { end with }\n\n"
                    '{\n'
                    '  "cop_ref": "specific CoP or '
                    'Chapter ref",\n'
                    '  "scope": "2-3 sentence scope '
                    'paragraph specific to activity",\n'
                    '  "personnel": [\n'
                    '    {"role": "Site Engineer", '
                    '"name": "TBD", '
                    '"responsibility": "Supervise"},\n'
                    '    {"role": "HSE Officer", '
                    '"name": "' + prepared_name + '", '
                    '"responsibility": "HSE oversight"},\n'
                    '    {"role": "Foreman", '
                    '"name": "TBD", '
                    '"responsibility": "Direct workers"},\n'
                    '    {"role": "Competent Worker", '
                    '"name": "TBD", '
                    '"responsibility": "Execute works"}\n'
                    '  ],\n'
                    '  "equipment": [\n'
                    '    {"item": "name", '
                    '"spec": "type/capacity", '
                    '"qty": "1"},\n'
                    '    ... 5-8 items relevant to '
                    + final_activity + '\n'
                    '  ],\n'
                    '  "ppe": [\n'
                    '    {"item": "Safety Helmet", '
                    '"standard": "EN 397", '
                    '"mandatory_for": "All"},\n'
                    '    {"item": "Safety Footwear", '
                    '"standard": "EN ISO 20345", '
                    '"mandatory_for": "All"},\n'
                    '    {"item": "High-Vis Vest", '
                    '"standard": "EN ISO 20471", '
                    '"mandatory_for": "All"},\n'
                    '    ... 5-8 more PPE specific to '
                    + final_activity + '\n'
                    '  ],\n'
                    '  "sequence": [\n'
                    '    {"task": "Pre-work briefing", '
                    '"safety_req": "Toolbox talk. '
                    'PTW if required."},\n'
                    '    ... 8-12 steps for '
                    + final_activity + '\n'
                    '  ],\n'
                    '  "hazards": [\n'
                    '    {"hazard": "specific hazard", '
                    '"risk": "injury type", '
                    '"controls": "controls with CoP", '
                    '"initial_risk": "HIGH", '
                    '"residual_risk": "LOW"},\n'
                    '    ... 6-10 hazards specific to '
                    + final_activity + '\n'
                    '  ],\n'
                    '  "permits": [\n'
                    '    {"type": "permit name", '
                    '"when": "when needed", '
                    '"ref": "CoP ref"}\n'
                    '  ],\n'
                    '  "checklist": [\n'
                    '    "All personnel briefed",\n'
                    '    "PTW obtained if required",\n'
                    '    ... 10-15 checklist items\n'
                    '  ],\n'
                    '  "references": [\n'
                    '    "ADOSH-SF Version 4.0 '
                    '(July 2024)",\n'
                    '    ... 5-8 references\n'
                    '  ]\n'
                    '}\n\n'
                    "Make all content SPECIFIC to "
                    + final_activity
                    + " in " + emirate + ". "
                    "Return ONLY the JSON object."
                )

                response = client_api.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=4000,
                    messages=[{
                        "role": "user",
                        "content": prompt
                    }]
                )

                prog.progress(
                    70,
                    text="Parsing response..."
                )

                raw = (
                    response.content[0].text.strip()
                )
                raw = re.sub(
                    r'```json|```', '', raw
                ).strip()
                start = raw.find('{')
                end   = raw.rfind('}') + 1
                if start == -1 or end == 0:
                    raise ValueError(
                        "No JSON found in response"
                    )
                raw = raw[start:end]

                content_data = json.loads(raw)

                prog.progress(
                    85,
                    text="Building Word document..."
                )

                meta = {
                    "activity":    final_activity,
                    "project":     ms_project,
                    "location":    ms_location,
                    "client":      ms_client,
                    "consultant":  ms_consultant,
                    "company":     ms_company,
                    "doc_no":      ms_doc_no,
                    "revision":    ms_rev,
                    "ms_date":     ms_date,
                    "prepared_by": ms_prep,
                    "reviewed_by": ms_review,
                    "approved_by": ms_approve,
                    "cop_ref": content_data.get(
                        'cop_ref', ''
                    ),
                }

                docx_buf = build_ms_docx(
                    meta, content_data, emirate
                )

                prog.progress(
                    100,
                    text="✅ Method Statement ready!"
                )

                st.success(
                    "✅ Method Statement for '"
                    + final_activity
                    + "' generated successfully!"
                )

                # Build filename safely
                act_clean  = final_activity.replace(
                    ' ', '_'
                ).replace('/', '-')
                proj_clean = ms_project.replace(
                    ' ', '_'
                )
                fname = (
                    "MS_"
                    + act_clean
                    + "_"
                    + proj_clean
                    + "_"
                    + str(ms_date)
                    + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download Method Statement"
                        " — "
                        + final_activity
                        + " (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="ms_download"
                )

                if emirate == "Abu Dhabi":
                    st.info(
                        "📋 **Next Steps (Abu Dhabi):**"
                        "\n\n1. Review and get HSE "
                        "Officer sign-off\n\n"
                        "2. Submit to Consultant "
                        "(Dar Al-Handasah) for approval"
                        "\n\n3. Brief all workers before"
                        " work commences\n\n"
                        "4. Keep signed copy on site\n\n"
                        "Ref: ADOSH CoP 53.1"
                    )
                else:
                    st.info(
                        "📋 **Next Steps (Dubai):**"
                        "\n\n1. Engineer reviews and "
                        "approves\n\n"
                        "2. Submit to Dubai Municipality"
                        " if required\n\n"
                        "3. Brief all workers before "
                        "work commences\n\n"
                        "Ref: DM Code Art. 2.3.2"
                    )

            except json.JSONDecodeError as e:
                prog.progress(0)
                st.error(
                    "JSON parse error: "
                    + str(e)
                    + "\n\nPlease try again — "
                    "AI response was too complex."
                )
            except Exception as e:
                prog.progress(0)
                st.error("Error: " + str(e))


# ════════════════════════════════════════════════════════
# TAB 2 — UPLOAD YOUR OWN FORMAT
# ════════════════════════════════════════════════════════

with tab_upload:
    st.markdown("### 📤 Upload Your Own Format")
    st.info(
        "Upload your company's Method Statement "
        "template (Word or PDF). The AI will fill "
        "it with content for your selected activity "
        "— keeping your EXACT format, layout, "
        "and branding."
    )

    uploaded_file = st.file_uploader(
        "Upload Your Method Statement "
        "Template or Sample",
        type=["docx", "pdf"],
        help=(
            "Upload a blank template or a previously "
            "completed Method Statement to use as "
            "the format reference."
        ),
        key="ms_upload"
    )

    if uploaded_file:
        st.success(
            "✅ File uploaded: **"
            + uploaded_file.name
            + "** ("
            + str(uploaded_file.size)
            + " bytes)"
        )

        st.divider()
        st.subheader("📋 Fill In Details")

        u_col1, u_col2 = st.columns(2)
        with u_col1:
            u_activity = st.text_input(
                "Activity for this Method Statement",
                placeholder="e.g. Excavation Works",
                key="upload_activity"
            )
            u_project = st.text_input(
                "Project Name",
                placeholder=(
                    "e.g. Bloom Living Almeria"
                ),
                key="upload_project"
            )
        with u_col2:
            u_company = st.text_input(
                "Company Name",
                value="ENGC",
                key="upload_company"
            )
            u_location = st.text_input(
                "Location",
                placeholder="e.g. Zayed City, Abu Dhabi",
                key="upload_location"
            )

        u_special = st.text_area(
            "Special Notes / Requirements",
            placeholder=(
                "e.g. Night works, confined space, "
                "adjacent to live road..."
            ),
            height=70,
            key="upload_special"
        )

        if st.button(
            "⚡ Generate Using My Uploaded Format",
            type="primary",
            use_container_width=True,
            key="upload_gen_btn"
        ):
            if not u_activity:
                st.error(
                    "Please enter an activity name."
                )
            else:
                with st.spinner(
                    "Reading your format and "
                    "generating content..."
                ):
                    try:
                        client_api = (
                            anthropic.Anthropic(
                                api_key=os.getenv(
                                    "ANTHROPIC_API_KEY"
                                )
                            )
                        )

                        file_bytes = uploaded_file.read()

                        if uploaded_file.name\
                                .lower()\
                                .endswith('.pdf'):
                            # PDF — send as base64
                            b64 = (
                                base64.standard_b64encode(
                                    file_bytes
                                ).decode()
                            )
                            response = (
                                client_api.messages
                                .create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "document",
                                            "source": {
                                                "type": "base64",
                                                "media_type": "application/pdf",
                                                "data": b64,
                                            }
                                        },
                                        {
                                            "type": "text",
                                            "text": (
                                                "This is a Method Statement "
                                                "template/format. "
                                                "Fill it with content for:\n"
                                                "Activity: " + u_activity + "\n"
                                                "Project: " + u_project + "\n"
                                                "Company: " + u_company + "\n"
                                                "Location: " + u_location + "\n"
                                                "Emirate: " + emirate + "\n"
                                                "Special: " + u_special + "\n\n"
                                                "Keep the EXACT same structure, "
                                                "headings, and sections. "
                                                "Fill with specific "
                                                + emirate
                                                + " regulatory references. "
                                                "Write the complete filled "
                                                "document."
                                            )
                                        }
                                    ]
                                }]
                            ))

                        else:
                            # DOCX — extract text
                            doc_temp = Document(
                                BytesIO(file_bytes)
                            )
                            extracted = "\n".join([
                                p.text
                                for p in doc_temp.paragraphs
                                if p.text.strip()
                            ])

                            response = (
                                client_api.messages
                                .create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": (
                                        "Here is a Method Statement "
                                        "template/format:\n\n"
                                        + extracted
                                        + "\n\n---\n\n"
                                        "Now fill this EXACT format "
                                        "with content for:\n"
                                        "Activity: " + u_activity + "\n"
                                        "Project: " + u_project + "\n"
                                        "Company: " + u_company + "\n"
                                        "Location: " + u_location + "\n"
                                        "Emirate: " + emirate + "\n"
                                        "Special Notes: " + u_special + "\n\n"
                                        "Keep the EXACT same headings, "
                                        "structure, and section order. "
                                        "Fill every section with specific "
                                        + emirate
                                        + " regulatory content. "
                                        "Reference specific ADOSH CoP "
                                        "numbers or Dubai Code chapters. "
                                        "Write the complete filled document."
                                    )
                                }]
                            ))

                        st.success(
                            "✅ Your format has been "
                            "filled with content for "
                            + u_activity + "!"
                        )

                        filled_content = (
                            response.content[0].text
                        )

                        st.markdown(
                            "### 📄 Generated Content "
                            "(in your format):"
                        )
                        st.markdown(filled_content)

                        # Build filename safely
                        act_fn  = u_activity.replace(
                            ' ', '_'
                        )
                        proj_fn = u_project.replace(
                            ' ', '_'
                        )
                        dl_fname = (
                            "MS_"
                            + act_fn
                            + "_"
                            + proj_fn
                            + "_filled.txt"
                        )

                        st.download_button(
                            label=(
                                "⬇️ Download as Text "
                                "(.txt)"
                            ),
                            data=filled_content,
                            file_name=dl_fname,
                            mime="text/plain",
                            key="upload_dl"
                        )

                        st.info(
                            "💡 **Tip:** Copy the "
                            "generated content above "
                            "and paste it into your "
                            "original Word template "
                            "to get a perfectly "
                            "formatted document."
                        )

                    except Exception as e:
                        st.error(
                            "Error: " + str(e)
                        )

    else:
        st.markdown("#### 💡 Tips for Best Results:")
        col_t1, col_t2 = st.columns(2)
        with col_t1:
            st.markdown("""
**Supported Formats:**
- ✅ Word Document (.docx)
- ✅ PDF (.pdf)

**What to Upload:**
- A blank Method Statement template
- A previously completed example
- Your company's standard format
""")
        with col_t2:
            st.markdown("""
**What the AI Does:**
- Reads your headings and structure
- Fills every section with content
- Uses correct Emirates regulations
- Keeps your exact layout

**Best Result:**
Upload a .docx file for best
format matching
""")
        st.divider()
        st.info(
            "📋 **Don't have a template?**\n\n"
            "Use the **'Generate Method Statement "
            "(AI)'** tab to create a professional "
            "one from scratch — ready for immediate "
            "use and submission."
        )