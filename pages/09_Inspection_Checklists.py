import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64

load_dotenv()

st.set_page_config(
    page_title="Inspection Checklists",
    page_icon="✅"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning(
        "Please go to Home and select "
        "your Emirate first."
    )
    st.stop()

color = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(
    "✅ Inspection Checklists — "
    + color + " " + emirate
)

if emirate == "Abu Dhabi":
    st.markdown(
        "**Reference: ADOSH-SF v4.0 | All 54 CoPs "
        "| MOHRE Resolution 44/2022**"
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality Code "
        "| All 23 Chapters**"
    )

st.divider()


# ════════════════════════════════════════════════════════
# CHECKLIST CATEGORIES — COMPREHENSIVE
# ════════════════════════════════════════════════════════

CHECKLIST_CATEGORIES = {

    "🏗️ SITE SETUP & ENTRY": {
        "color": "1F3864",
        "items": [
            "Site Entry & Access Control",
            "Security & Visitor Management",
            "Site Layout & Signage",
            "Hoarding & Fencing Inspection",
            "Site Office Setup",
            "Welfare Facilities (Toilets/Changing)",
            "Waste Segregation Area Setup",
            "Emergency Muster Point Setup",
            "First Aid Room / Station",
            "Fire Point & Fire Extinguishers",
        ]
    },

    "🌡️ WELFARE & WORKER HEALTH": {
        "color": "1B5E20",
        "items": [
            "Rest Area Inspection",
            "Water Cooler / Drinking Water",
            "Shaded Break Area",
            "Toilets & Sanitation Facilities",
            "Canteen / Food Storage Area",
            "First Aid Kit Inspection",
            "Heat Stress Prevention Checklist",
            "Worker Camp / Accommodation",
            "Medical Fitness Records",
            "PPE Storage & Issue Area",
        ]
    },

    "⚡ ELECTRICAL & POWER": {
        "color": "BF8F00",
        "items": [
            "Distribution Board (DB) Inspection",
            "Temporary Electrical Installation",
            "Generator Inspection",
            "Cable Management & Routing",
            "LOTO (Lockout-Tagout) Checklist",
            "Earthing & Bonding Verification",
            "RCD / ELCB Testing",
            "Portable Tools & Equipment",
            "Lighting Inspection (Site & Night)",
            "Electrical Panel / MDB Inspection",
        ]
    },

    "🏗️ CIVIL & STRUCTURAL WORKS": {
        "color": "1F5864",
        "items": [
            "Excavation Inspection",
            "Shoring & Trench Support",
            "Formwork & Falsework Inspection",
            "Reinforcement / Rebar Inspection",
            "Concrete Pouring Pre-Pour Check",
            "Post-Pour Inspection",
            "Precast Installation Checklist",
            "Masonry & Blockwork",
            "Foundation Works Inspection",
            "Underground Services Check",
        ]
    },

    "🏗️ SCAFFOLDING & HEIGHTS": {
        "color": "C00000",
        "items": [
            "Scaffolding Erection Inspection",
            "Scaffold Weekly Inspection",
            "Working at Heights Checklist",
            "Edge Protection & Guardrails",
            "Ladder Inspection",
            "Mobile Elevated Work Platform (MEWP)",
            "Rope Access Checklist",
            "Fall Arrest / Harness Inspection",
            "Safety Net Inspection",
            "Scaffolding Dismantling Checklist",
        ]
    },

    "🏗️ LIFTING OPERATIONS": {
        "color": "4A148C",
        "items": [
            "Crane Pre-Lift Checklist",
            "Daily Crane Inspection",
            "Rigging & Lifting Gear Inspection",
            "Lifting Plan Review Checklist",
            "Load Chart Verification",
            "Exclusion Zone Setup Checklist",
            "Telehandler / Forklift Inspection",
            "Man-Basket / Personnel Basket",
            "Tower Crane Inspection",
            "Boom Truck / Aerial Platform",
        ]
    },

    "🔥 HOT WORK & FIRE SAFETY": {
        "color": "C00000",
        "items": [
            "Hot Work Permit Checklist",
            "Fire Watch Checklist",
            "Welding & Cutting Area Setup",
            "Fire Extinguisher Inspection",
            "Fire Alarm System Check",
            "Flammable Materials Storage",
            "LPG / Gas Cylinder Storage",
            "Combustible Material Clearance",
            "Fire Emergency Drill Checklist",
            "Sprinkler / Suppression System",
        ]
    },

    "⚠️ CONFINED SPACE": {
        "color": "BF360C",
        "items": [
            "Confined Space Pre-Entry Checklist",
            "Atmospheric Testing Record",
            "Ventilation Equipment Check",
            "Rescue Equipment Inspection",
            "Attendant / Standby Person Checklist",
            "Confined Space Permit Review",
            "Communication System Check",
            "Emergency Extraction Equipment",
            "Worker Medical Fitness for CS",
            "Post-Exit Inspection Checklist",
        ]
    },

    "🚛 PLANT, EQUIPMENT & VEHICLES": {
        "color": "37474F",
        "items": [
            "Daily Plant & Equipment Checklist",
            "Heavy Vehicle Pre-Start Inspection",
            "Concrete Mixer Inspection",
            "Concrete Pump Inspection",
            "Compactor / Roller Inspection",
            "Excavator Daily Inspection",
            "Dump Truck / Tipper Inspection",
            "Mobile Plant Inspection",
            "Equipment Maintenance Record Check",
            "Third-Party Certification Check",
        ]
    },

    "🔧 MEP WORKS": {
        "color": "1565C0",
        "items": [
            "Plumbing Installation Inspection",
            "HVAC Equipment Checklist",
            "Fire Fighting System Inspection",
            "Electrical Conduit & Tray",
            "Cable Pulling Pre-Work Check",
            "Duct Installation Inspection",
            "Pressure Testing Checklist",
            "Flushing & Commissioning",
            "Solar Panel Installation Check",
            "BMS / Control Panel Inspection",
        ]
    },

    "🌿 ENVIRONMENTAL": {
        "color": "1B5E20",
        "items": [
            "Waste Management Inspection",
            "Hazardous Chemical Storage",
            "Spill Prevention & Containment",
            "Dust Control Inspection",
            "Noise Level Monitoring Checklist",
            "Water Pollution Prevention",
            "Dewatering Operations Check",
            "Concrete Washout Area",
            "Fuel & Oil Storage Inspection",
            "Environmental Incident Near Miss",
        ]
    },

    "📦 MATERIAL & STORAGE": {
        "color": "5C4033",
        "items": [
            "Material Storage Area Inspection",
            "Chemical / Hazmat Storage",
            "Rebar & Steel Storage",
            "Timber & Formwork Storage",
            "Fuel Tank Inspection",
            "Cement / Aggregate Storage",
            "Paint & Solvent Storage",
            "Gas Cylinder Storage",
            "MSDS / SDS Availability Check",
            "Material Delivery Checklist",
        ]
    },

    "📋 DAILY & PERIODIC INSPECTIONS": {
        "color": "1F3864",
        "items": [
            "Daily Site HSE Inspection",
            "Weekly HSE Inspection",
            "Monthly HSE Audit",
            "Pre-Task Risk Assessment (Take 5)",
            "Morning Site Inspection",
            "End-of-Day Site Inspection",
            "Toolbox Talk Attendance Record",
            "Induction Completion Record",
            "Near Miss Reporting Form",
            "Observation & Intervention Card",
        ]
    },

    "🏁 PROJECT CLOSE-OUT": {
        "color": "424242",
        "items": [
            "Site Demobilization Checklist",
            "Scaffold Dismantling Sign-Off",
            "Temporary Power Disconnection",
            "Waste Clearance & Disposal",
            "Hazardous Material Removal",
            "Final Site Inspection",
            "Handover Inspection Checklist",
            "Documentation Close-Out",
            "Lessons Learned Register",
            "HSE Close-Out Report Checklist",
        ]
    },
}


# ════════════════════════════════════════════════════════
# DOCX BUILDER
# ════════════════════════════════════════════════════════

def shd_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)


def shd_cell(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def cell_borders(cell, color="1F3864"):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def build_checklist_docx(
    checklist_name,
    checklist_items,
    meta,
    emirate,
    header_color="1F3864"
):
    """Build a professional A4 checklist .docx"""
    doc = Document()

    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    # ── TOP HEADER BANNER ─────────────────────
    hdr           = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(hdr, '1F3864')
    r = hdr.add_run(
        "\n  "
        + meta.get('company', 'ENGC')
        + "  \n"
    )
    r.bold           = True
    r.font.size      = Pt(16)
    r.font.color.rgb = RGBColor(255, 255, 255)

    sub_hdr           = doc.add_paragraph()
    sub_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(sub_hdr, header_color)
    sr = sub_hdr.add_run(
        "  ✅  "
        + checklist_name.upper()
        + "  ✅  "
    )
    sr.bold           = True
    sr.font.size      = Pt(14)
    sr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── PROJECT INFO TABLE ────────────────────
    info_tbl       = doc.add_table(rows=0, cols=4)
    info_tbl.style = 'Table Grid'

    def add_info_pair(tbl, lbl1, val1,
                       lbl2, val2):
        row = tbl.add_row()
        cells = row.cells
        cells[0].text = lbl1
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        shd_cell(cells[0], 'E8F0FE')
        cell_borders(cells[0])

        cells[1].text = str(val1) if val1 else "—"
        cells[1].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        cell_borders(cells[1])

        cells[2].text = lbl2
        cells[2].paragraphs[0].runs[0].bold = True
        cells[2].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        shd_cell(cells[2], 'E8F0FE')
        cell_borders(cells[2])

        cells[3].text = str(val2) if val2 else "—"
        cells[3].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        cell_borders(cells[3])

    add_info_pair(
        info_tbl,
        "Project", meta.get('project', ''),
        "Date", str(meta.get('insp_date',
                             date.today()))
    )
    add_info_pair(
        info_tbl,
        "Location / Area",
        meta.get('location', ''),
        "Shift", meta.get('shift', 'Day Shift')
    )
    add_info_pair(
        info_tbl,
        "Inspected By",
        meta.get('inspector', ''),
        "Designation",
        meta.get('designation', 'HSE Officer')
    )
    add_info_pair(
        info_tbl,
        "Contractor",
        meta.get('company', 'ENGC'),
        "Ref No.",
        meta.get('ref_no', '')
    )
    add_info_pair(
        info_tbl,
        "Permit No.",
        meta.get('permit_no', ''),
        "Weather",
        meta.get('weather', '')
    )

    doc.add_paragraph()

    # ── LEGEND ────────────────────────────────
    leg_p = doc.add_paragraph()
    shd_para(leg_p, '37474F')
    lr = leg_p.add_run(
        "  INSPECTION LEGEND:   "
        "✅ = Satisfactory   "
        "❌ = Non-Conformance   "
        "⚠️ = Needs Attention   "
        "N/A = Not Applicable"
    )
    lr.bold           = True
    lr.font.size      = Pt(9)
    lr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── CHECKLIST TABLE ───────────────────────
    # Header row
    hdr_p = doc.add_paragraph()
    shd_para(hdr_p, header_color)
    hdr_r = hdr_p.add_run(
        "  INSPECTION ITEMS"
    )
    hdr_r.bold           = True
    hdr_r.font.size      = Pt(11)
    hdr_r.font.color.rgb = RGBColor(255, 255, 255)

    chk_tbl       = doc.add_table(
        rows=1, cols=5
    )
    chk_tbl.style = 'Table Grid'

    # Column headers
    col_headers = [
        ("No.", Cm(1.0)),
        ("Inspection Item / Checklist Point",
         Cm(7.5)),
        ("Status\n✅ ❌ ⚠️ N/A", Cm(2.5)),
        ("Observations / Findings", Cm(4.5)),
        ("Action Required / By Whom / Date",
         Cm(3.5)),
    ]
    for i, (hd, wd) in enumerate(col_headers):
        c = chk_tbl.rows[0].cells[i]
        c.width = wd
        c.text  = hd
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, header_color)
        cell_borders(c)
        c.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    # Data rows
    for idx, item in enumerate(checklist_items):
        row   = chk_tbl.add_row()
        cells = row.cells

        # No.
        cells[0].width = Cm(1.0)
        cells[0].text  = str(idx + 1)
        cells[0].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        cells[0].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
        if idx % 2 == 0:
            shd_cell(cells[0], 'F5F5F5')
        cell_borders(cells[0])

        # Inspection Item
        cells[1].width = Cm(7.5)
        p = cells[1].paragraphs[0]
        if isinstance(item, dict):
            item_text   = item.get('item', '')
            item_detail = item.get('detail', '')
        else:
            item_text   = str(item)
            item_detail = ''
        run = p.add_run(item_text)
        run.font.size = Pt(9)
        run.bold = True
        if item_detail:
            p2  = cells[1].add_paragraph()
            r2  = p2.add_run(item_detail)
            r2.font.size = Pt(8)
            r2.font.color.rgb = RGBColor(
                100, 100, 100
            )
        if idx % 2 == 0:
            shd_cell(cells[1], 'F5F5F5')
        cell_borders(cells[1])

        # Status checkbox
        cells[2].width = Cm(2.5)
        cells[2].text  = (
            "☐ ✅  ☐ ❌\n☐ ⚠️  ☐ N/A"
        )
        cells[2].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        cells[2].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
        if idx % 2 == 0:
            shd_cell(cells[2], 'F5F5F5')
        cell_borders(cells[2])

        # Observations
        cells[3].width = Cm(4.5)
        cells[3].text  = ""
        if idx % 2 == 0:
            shd_cell(cells[3], 'F5F5F5')
        cell_borders(cells[3])

        # Action required
        cells[4].width = Cm(3.5)
        cells[4].text  = ""
        if idx % 2 == 0:
            shd_cell(cells[4], 'F5F5F5')
        cell_borders(cells[4])

    doc.add_paragraph()

    # ── SUMMARY BOX ───────────────────────────
    sum_p = doc.add_paragraph()
    shd_para(sum_p, '1F3864')
    sum_r = sum_p.add_run(
        "  INSPECTION SUMMARY"
    )
    sum_r.bold           = True
    sum_r.font.size      = Pt(11)
    sum_r.font.color.rgb = RGBColor(255, 255, 255)

    sum_tbl       = doc.add_table(rows=2, cols=4)
    sum_tbl.style = 'Table Grid'
    sum_hdrs = [
        "Total Items",
        "Satisfactory ✅",
        "Non-Conformances ❌",
        "Actions Required"
    ]
    sum_fills = ['1B5E20','1B5E20','C00000','BF8F00']
    for i, (h, f) in enumerate(
        zip(sum_hdrs, sum_fills)
    ):
        c = sum_tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, f)
        cell_borders(c)
        c.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    for i, val in enumerate([
        str(len(checklist_items)), "", "", ""
    ]):
        c = sum_tbl.rows[1].cells[i]
        c.text = val
        c.paragraphs[0].runs[0].font.size = Pt(14)
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
        cell_borders(c)

    doc.add_paragraph()

    # ── OVERALL RATING ────────────────────────
    rate_tbl       = doc.add_table(rows=1, cols=4)
    rate_tbl.style = 'Table Grid'
    ratings = [
        ("SATISFACTORY", "1B5E20"),
        ("NEEDS IMPROVEMENT", "BF8F00"),
        ("UNSATISFACTORY", "C00000"),
        ("SITE STOPPED", "FF0000"),
    ]
    for i, (rating, fill) in enumerate(ratings):
        c = rate_tbl.rows[0].cells[i]
        c.text = "☐  " + rating
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, fill)
        cell_borders(c)
        c.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    doc.add_paragraph()

    # ── REMARKS ───────────────────────────────
    rem_p = doc.add_paragraph()
    shd_para(rem_p, '37474F')
    rem_r = rem_p.add_run(
        "  ADDITIONAL REMARKS / OBSERVATIONS:"
    )
    rem_r.bold           = True
    rem_r.font.size      = Pt(10)
    rem_r.font.color.rgb = RGBColor(255, 255, 255)

    for _ in range(4):
        rp  = doc.add_paragraph()
        rr  = rp.add_run(
            "_" * 100
        )
        rr.font.size = Pt(10)
        rp.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ── REGULATORY NOTE ───────────────────────
    if emirate == "Abu Dhabi":
        reg_note = (
            "This checklist is prepared in "
            "accordance with ADOSH-SF Version 4.0 "
            "(July 2024) and applicable ADOSH Codes "
            "of Practice. All non-conformances must "
            "be closed out within the agreed "
            "timeframe and reported as required."
        )
    else:
        reg_note = (
            "This checklist is prepared in "
            "accordance with Dubai Municipality Code "
            "of Construction Safety Practice. All "
            "non-conformances must be reported to "
            "the site safety officer and closed "
            "within the agreed timeframe."
        )

    rn_p = doc.add_paragraph()
    rn_p.paragraph_format.space_before = Pt(4)
    rn_r = rn_p.add_run(
        "⚠️  REGULATORY NOTE: " + reg_note
    )
    rn_r.font.size  = Pt(8)
    rn_r.font.italic = True
    rn_r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── SIGNATURE BLOCK ───────────────────────
    sig_p = doc.add_paragraph()
    shd_para(sig_p, '1F3864')
    sig_r = sig_p.add_run(
        "  SIGNATURES"
    )
    sig_r.bold           = True
    sig_r.font.size      = Pt(11)
    sig_r.font.color.rgb = RGBColor(255, 255, 255)

    sig_tbl       = doc.add_table(rows=2, cols=3)
    sig_tbl.style = 'Table Grid'
    sig_roles = [
        "Inspector / HSE Officer",
        "Site Engineer / Foreman",
        "Project Manager"
    ]
    for i, role in enumerate(sig_roles):
        c = sig_tbl.rows[0].cells[i]
        c.text = role
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, '1F3864')
        cell_borders(c)

    for i in range(3):
        c      = sig_tbl.rows[1].cells[i]
        c.text = (
            "\nName: _______________\n\n"
            "Signature: _______________\n\n"
            "Date: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_multi_checklist_docx(
    selected_checklists,
    all_items_map,
    meta,
    emirate
):
    """
    Build one combined .docx with multiple
    checklists, each on a new page.
    """
    doc = Document()

    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    for cl_idx, (cat_name, cl_name) in enumerate(
        selected_checklists
    ):
        cat_data = CHECKLIST_CATEGORIES.get(
            cat_name, {}
        )
        hdr_color = cat_data.get('color', '1F3864')
        items     = all_items_map.get(cl_name, [])

        if cl_idx > 0:
            doc.add_page_break()

        # Category banner
        cat_p = doc.add_paragraph()
        cat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd_para(cat_p, '1F3864')
        cr = cat_p.add_run(
            "  "
            + meta.get('company', 'ENGC')
            + "  |  "
            + emirate.upper()
            + "  "
        )
        cr.bold           = True
        cr.font.size      = Pt(13)
        cr.font.color.rgb = RGBColor(255, 255, 255)

        # Checklist name banner
        cl_p           = doc.add_paragraph()
        cl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd_para(cl_p, hdr_color)
        clr = cl_p.add_run(
            "  ✅  "
            + cl_name.upper()
            + "  ✅  "
        )
        clr.bold           = True
        clr.font.size      = Pt(12)
        clr.font.color.rgb = RGBColor(255, 255, 255)

        doc.add_paragraph()

        # Project info
        info_tbl       = doc.add_table(
            rows=0, cols=4
        )
        info_tbl.style = 'Table Grid'
        info_pairs = [
            ("Project", meta.get('project', ''),
             "Date",
             str(meta.get('insp_date',
                          date.today()))),
            ("Location", meta.get('location', ''),
             "Inspector",
             meta.get('inspector', '')),
            ("Shift", meta.get('shift', 'Day'),
             "Ref No.", meta.get('ref_no', '')),
        ]
        for l1, v1, l2, v2 in info_pairs:
            row   = info_tbl.add_row()
            cells = row.cells

            cells[0].text = l1
            cells[0].paragraphs[0].runs[0]\
                .bold = True
            cells[0].paragraphs[0].runs[0]\
                .font.size = Pt(8)
            shd_cell(cells[0], 'E8F0FE')
            cell_borders(cells[0])

            cells[1].text = str(v1)
            cells[1].paragraphs[0].runs[0]\
                .font.size = Pt(8)
            cell_borders(cells[1])

            cells[2].text = l2
            cells[2].paragraphs[0].runs[0]\
                .bold = True
            cells[2].paragraphs[0].runs[0]\
                .font.size = Pt(8)
            shd_cell(cells[2], 'E8F0FE')
            cell_borders(cells[2])

            cells[3].text = str(v2)
            cells[3].paragraphs[0].runs[0]\
                .font.size = Pt(8)
            cell_borders(cells[3])

        doc.add_paragraph()

        # Legend
        leg_p = doc.add_paragraph()
        shd_para(leg_p, '37474F')
        lr = leg_p.add_run(
            "  ✅ Satisfactory  "
            "❌ Non-Conformance  "
            "⚠️ Needs Attention  "
            "N/A = Not Applicable"
        )
        lr.bold           = True
        lr.font.size      = Pt(8)
        lr.font.color.rgb = RGBColor(255, 255, 255)

        doc.add_paragraph()

        # Checklist table
        chk_tbl       = doc.add_table(
            rows=1, cols=5
        )
        chk_tbl.style = 'Table Grid'
        col_headers = [
            ("No.", Cm(0.8)),
            ("Inspection Item", Cm(7.0)),
            ("Status", Cm(2.5)),
            ("Observations", Cm(4.5)),
            ("Action / By / Date", Cm(4.2)),
        ]
        for i, (hd, wd) in enumerate(col_headers):
            c       = chk_tbl.rows[0].cells[i]
            c.width = wd
            c.text  = hd
            c.paragraphs[0].runs[0].bold      = True
            c.paragraphs[0].runs[0].font.size = Pt(8)
            c.paragraphs[0].runs[0]\
                .font.color.rgb = RGBColor(
                255, 255, 255
            )
            shd_cell(c, hdr_color)
            cell_borders(c)
            c.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

        for idx, item in enumerate(items):
            row   = chk_tbl.add_row()
            cells = row.cells

            cells[0].width = Cm(0.8)
            cells[0].text  = str(idx + 1)
            cells[0].paragraphs[0].runs[0]\
                .font.size = Pt(8)
            cells[0].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )
            if idx % 2 == 0:
                shd_cell(cells[0], 'F5F5F5')
            cell_borders(cells[0])

            cells[1].width = Cm(7.0)
            p_item = cells[1].paragraphs[0]
            if isinstance(item, dict):
                it  = item.get('item', str(item))
                det = item.get('detail', '')
            else:
                it  = str(item)
                det = ''
            run      = p_item.add_run(it)
            run.bold      = True
            run.font.size = Pt(8)
            if det:
                p2  = cells[1].add_paragraph()
                r2  = p2.add_run(det)
                r2.font.size      = Pt(7)
                r2.font.color.rgb = RGBColor(
                    100, 100, 100
                )
            if idx % 2 == 0:
                shd_cell(cells[1], 'F5F5F5')
            cell_borders(cells[1])

            cells[2].width = Cm(2.5)
            cells[2].text  = (
                "☐ ✅  ☐ ❌\n"
                "☐ ⚠️  ☐ N/A"
            )
            cells[2].paragraphs[0].runs[0]\
                .font.size = Pt(8)
            cells[2].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )
            if idx % 2 == 0:
                shd_cell(cells[2], 'F5F5F5')
            cell_borders(cells[2])

            cells[3].width = Cm(4.5)
            cells[3].text  = ""
            if idx % 2 == 0:
                shd_cell(cells[3], 'F5F5F5')
            cell_borders(cells[3])

            cells[4].width = Cm(4.2)
            cells[4].text  = ""
            if idx % 2 == 0:
                shd_cell(cells[4], 'F5F5F5')
            cell_borders(cells[4])

        doc.add_paragraph()

        # Signature row
        sig_tbl       = doc.add_table(rows=2, cols=3)
        sig_tbl.style = 'Table Grid'
        for i, role in enumerate([
            "Inspector / HSE Officer",
            "Site Engineer / Foreman",
            "Project Manager"
        ]):
            c = sig_tbl.rows[0].cells[i]
            c.text = role
            c.paragraphs[0].runs[0].bold      = True
            c.paragraphs[0].runs[0].font.size = Pt(8)
            c.paragraphs[0].runs[0]\
                .font.color.rgb = RGBColor(
                255, 255, 255
            )
            shd_cell(c, '1F3864')
            cell_borders(c)
        for i in range(3):
            c = sig_tbl.rows[1].cells[i]
            c.text = (
                "\nName: _______________\n"
                "Signature: _______________\n"
                "Date: _______________"
            )
            c.paragraphs[0].runs[0].font.size = Pt(8)
            cell_borders(c)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════
# MAIN UI — THREE TABS
# ════════════════════════════════════════════════════════

tab_single, tab_multi, tab_ai, tab_upload = st.tabs([
    "✅ Single Checklist",
    "📦 Multiple Checklists (Pack)",
    "🤖 AI Custom Checklist",
    "📤 Upload Your Format",
])


# ════════════════════════════════════════════════════════
# TAB 1 — SINGLE CHECKLIST
# ════════════════════════════════════════════════════════

with tab_single:
    st.markdown(
        "### ✅ Generate a Single Inspection "
        "Checklist"
    )
    st.info(
        "Select any checklist from the full "
        "library — from site entry to project "
        "close-out. Downloads as a professional "
        ".docx file ready for site use."
    )

    # Category + checklist selection
    cat_name = st.selectbox(
        "Select Category",
        list(CHECKLIST_CATEGORIES.keys()),
        key="single_cat"
    )

    cat_data = CHECKLIST_CATEGORIES[cat_name]
    cl_name  = st.selectbox(
        "Select Checklist",
        cat_data['items'],
        key="single_cl"
    )

    st.divider()
    st.subheader("📋 Inspection Details")

    s1, s2 = st.columns(2)
    with s1:
        s_project  = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="s_proj"
        )
        s_location = st.text_input(
            "Location / Area on Site",
            placeholder="e.g. Block A — Level 2",
            key="s_loc"
        )
        s_company  = st.text_input(
            "Company / Contractor",
            value="ENGC",
            key="s_co"
        )
    with s2:
        s_inspector = st.text_input(
            "Inspector Name",
            placeholder="e.g. Aftab Qamar",
            key="s_insp"
        )
        s_desig = st.text_input(
            "Designation",
            value="Senior HSSE Engineer",
            key="s_desig"
        )
        s_date = st.date_input(
            "Inspection Date",
            value=date.today(),
            key="s_date"
        )

    s3, s4 = st.columns(2)
    with s3:
        s_shift = st.selectbox(
            "Shift",
            ["Day Shift", "Night Shift",
             "Morning", "Afternoon"],
            key="s_shift"
        )
        s_ref = st.text_input(
            "Reference No.",
            placeholder="e.g. ENGC-CL-001",
            key="s_ref"
        )
    with s4:
        s_permit = st.text_input(
            "Permit No. (if applicable)",
            placeholder="e.g. PTW-2024-001",
            key="s_permit"
        )
        s_weather = st.selectbox(
            "Weather Conditions",
            ["Clear / Sunny", "Partly Cloudy",
             "Hot (>40°C)", "Windy",
             "Dusty / Sandstorm", "Humid"],
            key="s_weather"
        )

    if st.button(
        "⚡ Generate Checklist (.docx)",
        type="primary",
        use_container_width=True,
        key="single_gen_btn"
    ):
        if not s_project:
            st.error(
                "Please enter a project name."
            )
        else:
            with st.spinner(
                "Building inspection checklist..."
            ):
                meta = {
                    "project":     s_project,
                    "location":    s_location,
                    "company":     s_company,
                    "inspector":   s_inspector,
                    "designation": s_desig,
                    "insp_date":   s_date,
                    "shift":       s_shift,
                    "ref_no":      s_ref,
                    "permit_no":   s_permit,
                    "weather":     s_weather,
                }
                docx_buf = build_checklist_docx(
                    cl_name,
                    cat_data['items'],
                    meta,
                    emirate,
                    cat_data['color']
                )

                proj_fn = s_project.replace(
                    ' ', '_'
                )
                cl_fn   = cl_name.replace(
                    ' ', '_'
                ).replace('/', '-')
                fname = (
                    "CL_"
                    + cl_fn + "_"
                    + proj_fn + "_"
                    + str(s_date) + ".docx"
                )

                st.success(
                    "✅ Checklist ready — "
                    + str(len(cat_data['items']))
                    + " inspection items!"
                )
                st.download_button(
                    label=(
                        "⬇️ Download: "
                        + cl_name + " (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="single_dl"
                )


# ════════════════════════════════════════════════════════
# TAB 2 — MULTIPLE CHECKLISTS (PACK)
# ════════════════════════════════════════════════════════

with tab_multi:
    st.markdown(
        "### 📦 Generate Multiple Checklists "
        "in One Document"
    )
    st.info(
        "Select multiple checklists to combine "
        "into one professional pack — each "
        "checklist on its own page. Perfect for "
        "weekly inspection packs, activity packs, "
        "or complete site audit packages."
    )

    # Quick packs
    st.markdown("#### ⚡ Quick Packs")
    quick_packs = {
        "🌅 Morning Site Inspection Pack": [
            ("📋 DAILY & PERIODIC INSPECTIONS",
             "Morning Site Inspection"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Water Cooler / Drinking Water"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Rest Area Inspection"),
            ("⚡ ELECTRICAL & POWER",
             "Distribution Board (DB) Inspection"),
            ("🔥 HOT WORK & FIRE SAFETY",
             "Fire Extinguisher Inspection"),
        ],
        "🏗️ Excavation Activity Pack": [
            ("🏗️ CIVIL & STRUCTURAL WORKS",
             "Excavation Inspection"),
            ("🏗️ CIVIL & STRUCTURAL WORKS",
             "Shoring & Trench Support"),
            ("🏗️ LIFTING OPERATIONS",
             "Crane Pre-Lift Checklist"),
            ("⚡ ELECTRICAL & POWER",
             "Temporary Electrical Installation"),
            ("🌿 ENVIRONMENTAL",
             "Dewatering Operations Check"),
        ],
        "🏗️ Height Works Pack": [
            ("🏗️ SCAFFOLDING & HEIGHTS",
             "Scaffolding Erection Inspection"),
            ("🏗️ SCAFFOLDING & HEIGHTS",
             "Working at Heights Checklist"),
            ("🏗️ SCAFFOLDING & HEIGHTS",
             "Fall Arrest / Harness Inspection"),
            ("🏗️ SCAFFOLDING & HEIGHTS",
             "Edge Protection & Guardrails"),
            ("📋 DAILY & PERIODIC INSPECTIONS",
             "Pre-Task Risk Assessment (Take 5)"),
        ],
        "⚡ Electrical & Generator Pack": [
            ("⚡ ELECTRICAL & POWER",
             "Generator Inspection"),
            ("⚡ ELECTRICAL & POWER",
             "Distribution Board (DB) Inspection"),
            ("⚡ ELECTRICAL & POWER",
             "Temporary Electrical Installation"),
            ("⚡ ELECTRICAL & POWER",
             "LOTO (Lockout-Tagout) Checklist"),
            ("⚡ ELECTRICAL & POWER",
             "Cable Management & Routing"),
        ],
        "🌡️ Welfare & Heat Stress Pack": [
            ("🌡️ WELFARE & WORKER HEALTH",
             "Heat Stress Prevention Checklist"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Water Cooler / Drinking Water"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Shaded Break Area"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Rest Area Inspection"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Toilets & Sanitation Facilities"),
        ],
        "📋 Weekly Full Site Audit Pack": [
            ("📋 DAILY & PERIODIC INSPECTIONS",
             "Weekly HSE Inspection"),
            ("🌡️ WELFARE & WORKER HEALTH",
             "Rest Area Inspection"),
            ("⚡ ELECTRICAL & POWER",
             "Distribution Board (DB) Inspection"),
            ("🔥 HOT WORK & FIRE SAFETY",
             "Fire Extinguisher Inspection"),
            ("🏗️ SCAFFOLDING & HEIGHTS",
             "Scaffold Weekly Inspection"),
            ("🚛 PLANT, EQUIPMENT & VEHICLES",
             "Daily Plant & Equipment Checklist"),
            ("🌿 ENVIRONMENTAL",
             "Waste Management Inspection"),
        ],
    }

    selected_pack = st.selectbox(
        "Choose a Quick Pack",
        ["— Select a Pack —"]
        + list(quick_packs.keys()),
        key="quick_pack"
    )

    st.markdown("**— OR —**")
    st.markdown(
        "#### 🎯 Build Your Own Custom Pack"
    )

    # Multi-select from all categories
    all_checklists_flat = []
    for cat_n, cat_d in (
        CHECKLIST_CATEGORIES.items()
    ):
        for item in cat_d['items']:
            all_checklists_flat.append(
                cat_n + "  ›  " + item
            )

    custom_selection = st.multiselect(
        "Select Individual Checklists",
        all_checklists_flat,
        help=(
            "Choose any checklists to combine "
            "into one document pack"
        ),
        key="custom_multi"
    )

    st.divider()
    st.subheader("📋 Project Details")

    m1, m2 = st.columns(2)
    with m1:
        m_project = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="m_proj"
        )
        m_location = st.text_input(
            "Site Location / Area",
            placeholder="e.g. Zayed City, Abu Dhabi",
            key="m_loc"
        )
        m_company = st.text_input(
            "Company",
            value="ENGC",
            key="m_co"
        )
    with m2:
        m_inspector = st.text_input(
            "Inspector Name",
            placeholder="e.g. Aftab Qamar",
            key="m_insp"
        )
        m_date = st.date_input(
            "Inspection Date",
            value=date.today(),
            key="m_date"
        )
        m_shift = st.selectbox(
            "Shift",
            ["Day Shift", "Night Shift",
             "Morning", "Afternoon"],
            key="m_shift"
        )

    m_ref = st.text_input(
        "Reference / Pack Number",
        placeholder="e.g. ENGC-WEEKLY-W01-2024",
        key="m_ref"
    )

    if st.button(
        "⚡ Generate Checklist Pack (.docx)",
        type="primary",
        use_container_width=True,
        key="multi_gen_btn"
    ):
        if not m_project:
            st.error(
                "Please enter a project name."
            )
        else:
            # Determine which checklists to use
            final_selection = []

            if (selected_pack != "— Select a Pack —"
                    and selected_pack
                    in quick_packs):
                final_selection = (
                    quick_packs[selected_pack]
                )

            # Add custom selections
            for sel in custom_selection:
                parts = sel.split("  ›  ")
                if len(parts) == 2:
                    final_selection.append(
                        (parts[0], parts[1])
                    )

            # Deduplicate
            seen = set()
            unique_selection = []
            for item in final_selection:
                key = str(item)
                if key not in seen:
                    seen.add(key)
                    unique_selection.append(item)

            if not unique_selection:
                st.error(
                    "Please select a Quick Pack "
                    "or choose individual "
                    "checklists above."
                )
            else:
                with st.spinner(
                    "Building checklist pack "
                    "with "
                    + str(len(unique_selection))
                    + " checklists..."
                ):
                    meta = {
                        "project":   m_project,
                        "location":  m_location,
                        "company":   m_company,
                        "inspector": m_inspector,
                        "insp_date": m_date,
                        "shift":     m_shift,
                        "ref_no":    m_ref,
                    }

                    # Build items map
                    items_map = {}
                    for (cat_n, cl_n) in (
                        unique_selection
                    ):
                        cat_d = (
                            CHECKLIST_CATEGORIES
                            .get(cat_n, {})
                        )
                        items_map[cl_n] = (
                            cat_d.get('items', [])
                        )

                    docx_buf = (
                        build_multi_checklist_docx(
                            unique_selection,
                            items_map,
                            meta,
                            emirate
                        )
                    )

                    proj_fn = m_project.replace(
                        ' ', '_'
                    )
                    pack_name = (
                        selected_pack
                        .split(' ')[0]
                        .replace('—','Pack')
                        if selected_pack !=
                        "— Select a Pack —"
                        else "Custom"
                    )
                    fname = (
                        "CL_Pack_"
                        + proj_fn + "_"
                        + str(m_date) + ".docx"
                    )

                    st.success(
                        "✅ Checklist Pack ready — "
                        + str(len(unique_selection))
                        + " checklists included!"
                    )
                    st.download_button(
                        label=(
                            "⬇️ Download Checklist "
                            "Pack — "
                            + str(len(
                                unique_selection
                            ))
                            + " Checklists (.docx)"
                        ),
                        data=docx_buf,
                        file_name=fname,
                        mime=(
                            "application/vnd"
                            ".openxmlformats-"
                            "officedocument"
                            ".wordprocessingml"
                            ".document"
                        ),
                        key="multi_dl"
                    )


# ════════════════════════════════════════════════════════
# TAB 3 — AI CUSTOM CHECKLIST
# ════════════════════════════════════════════════════════

with tab_ai:
    st.markdown(
        "### 🤖 AI Custom Checklist Generator"
    )
    st.info(
        "Describe any specific activity, equipment, "
        "or area and the AI will generate a "
        "detailed, Emirates-specific inspection "
        "checklist with all relevant checks — "
        "including regulatory references."
    )

    ai_activity = st.text_input(
        "What do you want to inspect?",
        placeholder=(
            "e.g. Temporary generator set on site, "
            "water cooler station, concrete batching "
            "plant, fuel storage area, labour camp..."
        ),
        key="ai_activity"
    )

    ai_col1, ai_col2 = st.columns(2)
    with ai_col1:
        ai_project = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="ai_proj"
        )
        ai_company = st.text_input(
            "Company",
            value="ENGC",
            key="ai_co"
        )
        ai_inspector = st.text_input(
            "Inspector",
            placeholder="e.g. Aftab Qamar",
            key="ai_insp"
        )
    with ai_col2:
        ai_location = st.text_input(
            "Location / Area",
            placeholder="e.g. Site North, Level 3",
            key="ai_loc"
        )
        ai_date = st.date_input(
            "Date",
            value=date.today(),
            key="ai_date"
        )
        ai_items = st.slider(
            "Number of Checklist Items",
            min_value=10,
            max_value=50,
            value=25,
            key="ai_items"
        )

    ai_notes = st.text_area(
        "Special Requirements / Notes (optional)",
        placeholder=(
            "e.g. Include heat stress checks, "
            "night work conditions, ADOSH CoP refs..."
        ),
        height=70,
        key="ai_notes"
    )

    if st.button(
        "⚡ Generate AI Custom Checklist (.docx)",
        type="primary",
        use_container_width=True,
        key="ai_gen_btn"
    ):
        if not ai_activity or not ai_project:
            st.error(
                "Please enter the activity "
                "and project name."
            )
        else:
            prog = st.progress(
                0, text="AI generating checklist..."
            )
            try:
                client_api = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                if emirate == "Abu Dhabi":
                    reg = (
                        "ADOSH-SF v4.0, cite specific "
                        "ADOSH CoP numbers for each "
                        "check item"
                    )
                else:
                    reg = (
                        "Dubai Municipality Code, "
                        "cite specific chapter "
                        "numbers for each check item"
                    )

                prompt = (
                    "You are a senior HSE engineer "
                    "in " + emirate + ", UAE.\n\n"
                    "Generate a detailed inspection "
                    "checklist for:\n"
                    "ACTIVITY/ITEM: " + ai_activity + "\n"
                    "PROJECT: " + ai_project + "\n"
                    "LOCATION: " + ai_location + "\n"
                    "COMPANY: " + ai_company + "\n"
                    "SPECIAL NOTES: " + ai_notes + "\n"
                    "REGULATION: " + reg + "\n"
                    "NUMBER OF ITEMS: "
                    + str(ai_items) + "\n\n"
                    "Return ONLY a JSON array. "
                    "No markdown. No text. "
                    "Start with [ end with ]\n\n"
                    "Each item must be:\n"
                    '{\n'
                    '  "item": "Short checklist '
                    'point title",\n'
                    '  "detail": "Specific '
                    'requirement — ref: CoP XX.X '
                    'or DM Code Ch.XX"\n'
                    '}\n\n'
                    "Make items SPECIFIC to "
                    + ai_activity
                    + " in " + emirate + ".\n"
                    "Cover: pre-work checks, "
                    "during-work checks, "
                    "equipment condition, "
                    "safety measures, "
                    "regulatory compliance, "
                    "environmental checks, "
                    "PPE requirements, "
                    "emergency preparedness.\n"
                    "Return ONLY the JSON array."
                )

                prog.progress(
                    20,
                    text="AI generating items..."
                )

                response = client_api.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=3000,
                    messages=[{
                        "role": "user",
                        "content": prompt
                    }]
                )

                prog.progress(
                    65, text="Parsing items..."
                )

                import json
                import re

                raw = (
                    response.content[0].text.strip()
                )
                raw = re.sub(
                    r'```json|```', '', raw
                ).strip()
                start = raw.find('[')
                end   = raw.rfind(']') + 1
                if start == -1 or end == 0:
                    raise ValueError(
                        "No JSON array in response"
                    )
                raw   = raw[start:end]
                items = json.loads(raw)

                prog.progress(
                    80,
                    text="Building checklist doc..."
                )

                meta = {
                    "project":     ai_project,
                    "location":    ai_location,
                    "company":     ai_company,
                    "inspector":   ai_inspector,
                    "designation": "HSSE Engineer",
                    "insp_date":   ai_date,
                    "shift":       "Day Shift",
                    "ref_no":      "",
                    "permit_no":   "",
                    "weather":     "",
                }

                docx_buf = build_checklist_docx(
                    ai_activity,
                    items,
                    meta,
                    emirate,
                    "1F3864"
                )

                prog.progress(
                    100,
                    text="✅ Checklist ready!"
                )

                st.success(
                    "✅ Custom checklist generated "
                    "with "
                    + str(len(items))
                    + " inspection items!"
                )

                act_fn  = ai_activity.replace(
                    ' ', '_'
                ).replace('/', '-')
                proj_fn = ai_project.replace(
                    ' ', '_'
                )
                fname = (
                    "CL_AI_"
                    + act_fn + "_"
                    + proj_fn + "_"
                    + str(ai_date) + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download: "
                        + ai_activity
                        + " Checklist (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="ai_dl"
                )

                # Preview
                with st.expander(
                    "📋 Preview Checklist Items",
                    expanded=False
                ):
                    for i, item in enumerate(items):
                        if isinstance(item, dict):
                            st.markdown(
                                "**"
                                + str(i+1) + ". "
                                + item.get(
                                    'item', ''
                                )
                                + "**"
                            )
                            if item.get('detail'):
                                st.caption(
                                    item['detail']
                                )
                        else:
                            st.markdown(
                                str(i+1)
                                + ". " + str(item)
                            )

            except Exception as e:
                prog.progress(0)
                st.error("Error: " + str(e))


# ════════════════════════════════════════════════════════
# TAB 4 — UPLOAD YOUR FORMAT
# ════════════════════════════════════════════════════════

with tab_upload:
    st.markdown(
        "### 📤 Upload Your Own Checklist Format"
    )
    st.info(
        "Upload your company's existing checklist "
        "(Word or PDF). The AI will fill it with "
        "inspection items specific to your activity "
        "and project — keeping your EXACT format, "
        "headers, company logo area, and branding. "
        "Downloads as a proper **.docx** file."
    )

    up_file = st.file_uploader(
        "📎 Upload Your Checklist Template or Sample",
        type=["docx", "pdf"],
        help=(
            "Upload a blank checklist template or "
            "a completed sample. Word (.docx) works "
            "best. PDF also supported."
        ),
        key="up_cl_file"
    )

    if up_file:
        st.success(
            "✅ File uploaded: **"
            + up_file.name + "** ("
            + "{:,}".format(up_file.size)
            + " bytes)"
        )

        st.divider()
        st.subheader("📋 Fill In Details")

        up1, up2 = st.columns(2)
        with up1:
            up_activity = st.text_input(
                "What is this checklist for?",
                placeholder=(
                    "e.g. Excavation, Generator, "
                    "Scaffolding, Site Entry..."
                ),
                key="up_act"
            )
            up_project = st.text_input(
                "Project Name",
                placeholder=(
                    "e.g. Bloom Living Almeria"
                ),
                key="up_proj"
            )
            up_company = st.text_input(
                "Company",
                value="ENGC",
                key="up_co"
            )
        with up2:
            up_location = st.text_input(
                "Location / Area",
                placeholder=(
                    "e.g. Zayed City, Abu Dhabi"
                ),
                key="up_loc"
            )
            up_inspector = st.text_input(
                "Inspector Name",
                placeholder="e.g. Aftab Qamar",
                key="up_insp"
            )
            up_date = st.date_input(
                "Date",
                value=date.today(),
                key="up_date"
            )

        up_notes = st.text_area(
            "Special Requirements (optional)",
            placeholder=(
                "e.g. Include heat stress checks, "
                "night work conditions, ADOSH refs..."
            ),
            height=70,
            key="up_notes"
        )

        if st.button(
            "⚡ Fill My Checklist Format",
            type="primary",
            use_container_width=True,
            key="up_gen_btn"
        ):
            if not up_activity or not up_project:
                st.error(
                    "Please enter the activity "
                    "and project name."
                )
            else:
                with st.spinner(
                    "Reading your format and "
                    "generating inspection items..."
                ):
                    try:
                        client_api = (
                            anthropic.Anthropic(
                                api_key=os.getenv(
                                    "ANTHROPIC_API_KEY"
                                )
                            )
                        )

                        file_bytes = up_file.read()

                        if emirate == "Abu Dhabi":
                            reg = (
                                "ADOSH-SF v4.0 — "
                                "cite specific ADOSH "
                                "CoP numbers"
                            )
                        else:
                            reg = (
                                "Dubai Municipality "
                                "Code — cite specific "
                                "chapter numbers"
                            )

                        fill_instr = (
                            "This is an inspection "
                            "checklist template for "
                            + emirate + ", UAE.\n\n"
                            "Fill it for:\n"
                            "Activity: " + up_activity + "\n"
                            "Project: " + up_project + "\n"
                            "Location: " + up_location + "\n"
                            "Company: " + up_company + "\n"
                            "Inspector: " + up_inspector + "\n"
                            "Date: " + str(up_date) + "\n"
                            "Notes: " + up_notes + "\n"
                            "Regulation: " + reg + "\n\n"
                            "INSTRUCTIONS:\n"
                            "1. Keep the EXACT same "
                            "structure and layout\n"
                            "2. Fill all blank/placeholder "
                            "fields with real content\n"
                            "3. Add inspection items "
                            "specific to " + up_activity + "\n"
                            "4. Reference specific "
                            + emirate
                            + " regulations\n"
                            "5. Write the complete filled "
                            "checklist"
                        )

                        if up_file.name\
                                .lower()\
                                .endswith('.pdf'):
                            b64 = (
                                base64.standard_b64encode(
                                    file_bytes
                                ).decode()
                            )
                            response = (
                                client_api.messages
                                .create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "document",
                                            "source": {
                                                "type": "base64",
                                                "media_type": "application/pdf",
                                                "data": b64,
                                            }
                                        },
                                        {
                                            "type": "text",
                                            "text": fill_instr
                                        }
                                    ]
                                }]
                            ))
                        else:
                            doc_temp = Document(
                                BytesIO(file_bytes)
                            )
                            extracted = "\n".join([
                                p.text
                                for p in doc_temp.paragraphs
                                if p.text.strip()
                            ])
                            table_text = ""
                            for tbl in doc_temp.tables:
                                for row in tbl.rows:
                                    row_text = " | ".join([
                                        c.text.strip()
                                        for c in row.cells
                                        if c.text.strip()
                                    ])
                                    if row_text:
                                        table_text += (
                                            row_text + "\n"
                                        )
                            full = (
                                extracted
                                + ("\n\nTABLE:\n"
                                   + table_text
                                   if table_text else "")
                            )
                            response = (
                                client_api.messages
                                .create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": (
                                        "My checklist "
                                        "template:\n\n---\n"
                                        + full
                                        + "\n---\n\n"
                                        + fill_instr
                                    )
                                }]
                            ))

                        filled = (
                            response.content[0].text
                        )

                        st.success(
                            "✅ Checklist filled for "
                            + up_activity + "!"
                        )

                        # Show preview
                        with st.expander(
                            "📋 Preview Generated "
                            "Checklist",
                            expanded=True
                        ):
                            st.markdown(filled)

                        # Download as txt
                        act_fn  = up_activity.replace(
                            ' ', '_'
                        )
                        proj_fn = up_project.replace(
                            ' ', '_'
                        )
                        fname = (
                            "CL_"
                            + act_fn + "_"
                            + proj_fn + "_"
                            + str(up_date)
                            + "_filled.txt"
                        )

                        st.download_button(
                            label=(
                                "⬇️ Download Filled "
                                "Checklist (.txt)"
                            ),
                            data=filled,
                            file_name=fname,
                            mime="text/plain",
                            key="up_dl"
                        )

                        st.info(
                            "💡 **Tip:** Copy the "
                            "content above and paste "
                            "it into your original "
                            "Word template for the "
                            "final formatted version."
                        )

                    except Exception as e:
                        st.error(
                            "Error: " + str(e)
                        )

    else:
        c_a, c_b = st.columns(2)
        with c_a:
            st.markdown("""
**✅ Accepted Formats:**
- Word Document (.docx) — Best
- PDF (.pdf)

**📋 What to Upload:**
- Your company checklist template
- A consultant-approved checklist
- A completed sample checklist
- ADOSH or DM standard format
""")
        with c_b:
            st.markdown("""
**🤖 What the AI Does:**
- Reads your headings & structure
- Fills every row/field with content
- Uses correct Emirates regulations
- References specific CoP/chapter numbers
- Keeps your EXACT layout
""")
        st.divider()
        st.info(
            "📋 **Don't have a template?**\n\n"
            "Use the **Single Checklist** tab "
            "or **AI Custom Checklist** tab to "
            "generate professional checklists "
            "from scratch — ready for immediate "
            "site use."
        )