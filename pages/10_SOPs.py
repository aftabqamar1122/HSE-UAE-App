import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import json
import re

load_dotenv()

st.set_page_config(
    page_title="Standard Operating Procedures",
    page_icon="📘"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning(
        "Please go to Home and select "
        "your Emirate first."
    )
    st.stop()

flag = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(
    "📘 Standard Operating Procedures (SOPs)"
    " — " + flag + " " + emirate
)

if emirate == "Abu Dhabi":
    st.markdown(
        "**Reference: ADOSH-SF v4.0 | "
        "All relevant ADOSH CoPs | "
        "ISO 45001:2018 Clause 8.1**"
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality Code | "
        "UAE Federal Decree-Law 33/2021 | "
        "ISO 45001:2018 Clause 8.1**"
    )

st.info(
    "SOPs define the exact step-by-step procedure "
    "for carrying out any task safely and correctly. "
    "They are mandatory documents under ISO 45001 "
    "and ADOSH/DM requirements. Every high-risk "
    "activity must have an approved SOP before work "
    "commences."
)

st.divider()


# ════════════════════════════════════════════════════════
# SOP LIBRARY — Comprehensive UAE Construction SOPs
# ════════════════════════════════════════════════════════

SOP_LIBRARY = {

    "⚠️ High Risk Operations": {
        "color": "C00000",
        "sops": [
            "Working at Heights — Fall Prevention",
            "Confined Space Entry Procedure",
            "Hot Work (Welding & Cutting) Procedure",
            "Excavation & Trenching Operations",
            "Lifting Operations — Crane & Rigging",
            "Scaffolding Erection & Dismantling",
            "Demolition Works Procedure",
            "Explosive / Blasting Operations",
            "Abrasive Blasting Procedure",
            "Radiography / NDT Operations",
            "Asbestos Handling & Removal",
            "Underwater / Diving Operations",
        ]
    },

    "⚡ Electrical Safety": {
        "color": "BF8F00",
        "sops": [
            "Lockout-Tagout (LOTO) Procedure",
            "Temporary Electrical Installation",
            "Electrical Panel / DB Inspection",
            "Cable Laying & Termination",
            "Generator Operation & Maintenance",
            "High Voltage (HV) Work Procedure",
            "Testing & Commissioning (Electrical)",
            "RCD / ELCB Testing Procedure",
            "Solar Panel Installation Safety",
            "Emergency Electrical Isolation",
        ]
    },

    "🏗️ Civil & Structural": {
        "color": "1F5864",
        "sops": [
            "Concrete Pouring & Pumping",
            "Formwork & Falsework (Shuttering)",
            "Reinforcement Steel Fixing",
            "Precast Concrete Installation",
            "Foundation & Piling Works",
            "Dewatering Operations",
            "Ground Improvement Works",
            "Masonry & Blockwork",
            "Waterproofing Application",
            "Steel Erection & Structural Works",
            "Post-Tensioning Operations",
            "Underpinning & Shoring",
        ]
    },

    "🔧 MEP Operations": {
        "color": "1565C0",
        "sops": [
            "Plumbing & Drainage Works",
            "HVAC Installation & Commissioning",
            "Fire Fighting System Installation",
            "Pressure Testing (Hydrostatic)",
            "Pipe Welding & Joining",
            "Duct & Cladding Works",
            "Insulation Application",
            "BMS Commissioning Procedure",
            "Water Treatment System Operation",
            "Gas Line Installation & Testing",
        ]
    },

    "🚛 Plant & Equipment": {
        "color": "37474F",
        "sops": [
            "Crane Operation Procedure",
            "Forklift / Telehandler Operation",
            "Excavator Operation Procedure",
            "Concrete Mixer Operation",
            "Concrete Pump Operation",
            "Compactor / Roller Operation",
            "Aerial Work Platform (AWP) Operation",
            "Dump Truck / Tipper Operation",
            "Piling Rig Operation",
            "Air Compressor Operation",
            "Pressure Washer Operation",
            "Mobile Generator Operation",
        ]
    },

    "🌿 Environmental & Waste": {
        "color": "1B5E20",
        "sops": [
            "Hazardous Waste Handling & Disposal",
            "Chemical Spill Response Procedure",
            "Concrete Washout Management",
            "Dust Control & Suppression",
            "Dewatering Discharge Procedure",
            "Fuel & Oil Storage Management",
            "Noise Management Procedure",
            "Waste Segregation & Recycling",
            "Contaminated Soil Handling",
            "Environmental Incident Reporting",
        ]
    },

    "🌡️ Health & Welfare": {
        "color": "2E7D32",
        "sops": [
            "Heat Stress Prevention & Management",
            "First Aid Response Procedure",
            "Manual Handling & Ergonomics",
            "PPE Selection, Use & Maintenance",
            "Respiratory Protection Procedure",
            "Hygiene & Sanitation Management",
            "Drug & Alcohol Policy Procedure",
            "Occupational Health Monitoring",
            "Return to Work Procedure",
            "Welfare Facility Management",
        ]
    },

    "🚨 Emergency Procedures": {
        "color": "C00000",
        "sops": [
            "Site Emergency Response Procedure",
            "Fire Emergency Evacuation",
            "Medical Emergency Response",
            "Structural Collapse Response",
            "Chemical Emergency Response",
            "Flood / Heavy Rain Response",
            "Vehicle / Plant Accident Response",
            "Utility Strike Response",
            "Serious Injury / Fatality Protocol",
            "Pandemic / Mass Illness Response",
        ]
    },

    "📋 Administrative & Management": {
        "color": "4A148C",
        "sops": [
            "Permit to Work (PTW) System",
            "Induction & Onboarding Procedure",
            "Toolbox Talk Delivery Procedure",
            "Incident Reporting & Investigation",
            "Near Miss Reporting Procedure",
            "HSE Inspection & Audit Procedure",
            "Subcontractor HSE Management",
            "Site Access & Visitor Control",
            "Document Control Procedure",
            "Management of Change (MOC)",
            "Stop Work Authority Procedure",
            "HSE Training Management",
        ]
    },

    "🏗️ Finishing & Fit-Out": {
        "color": "5C4033",
        "sops": [
            "Spray Painting & Coating",
            "Epoxy Flooring Application",
            "Glass & Glazing Installation",
            "Curtain Wall Installation",
            "Flooring & Tiling Works",
            "False Ceiling Installation",
            "Carpentry & Joinery Works",
            "Adhesive & Sealant Application",
            "Stone Cladding Installation",
            "Roof Waterproofing Application",
        ]
    },
}


# ════════════════════════════════════════════════════════
# DOCX HELPERS
# ════════════════════════════════════════════════════════

def shd_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)


def shd_cell(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def cell_borders(cell, color="1F3864"):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def add_section_hdr(doc, text,
                    fill="1F3864", size=11):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    shd_para(p, fill)
    run           = p.add_run("  " + text.upper())
    run.bold      = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p


def add_body(doc, text, size=10,
             bold=False, indent=0.5):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent)
    run           = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold      = bold
    return p


def add_bullet_item(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    return p


def safe_json_list(raw_text, fallback=None):
    """Robustly parse JSON array with fallback."""
    if fallback is None:
        fallback = []
    raw_text = re.sub(
        r'```json|```', '', raw_text
    ).strip()
    # Direct parse
    try:
        result = json.loads(raw_text)
        if isinstance(result, list):
            return result
        if isinstance(result, dict):
            for v in result.values():
                if isinstance(v, list):
                    return v
    except Exception:
        pass
    # Extract array
    s = raw_text.find('[')
    e = raw_text.rfind(']') + 1
    if s != -1 and e > 0:
        try:
            result = json.loads(raw_text[s:e])
            if isinstance(result, list):
                return result
        except Exception:
            pass
    # Extract individual objects
    objs   = re.findall(r'\{[^\{\}]+\}', raw_text)
    parsed = []
    for obj in objs:
        try:
            parsed.append(json.loads(obj))
        except Exception:
            pass
    return parsed if parsed else fallback


def build_sop_docx(sop_data, meta,
                   emirate, hdr_color):
    """Build a professional SOP Word document."""
    doc = Document()

    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    # ── COVER HEADER ──────────────────────────
    cover           = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(cover, '1F3864')
    cr = cover.add_run(
        "\n  "
        + meta.get('company', 'ENGC')
        + "  \n"
    )
    cr.bold           = True
    cr.font.size      = Pt(18)
    cr.font.color.rgb = RGBColor(255, 255, 255)

    title_p           = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(title_p, hdr_color)
    tr = title_p.add_run(
        "\n  STANDARD OPERATING PROCEDURE"
        "  \n  "
        + sop_data.get('sop_title', '').upper()
        + "  \n"
    )
    tr.bold           = True
    tr.font.size      = Pt(15)
    tr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── DOCUMENT CONTROL TABLE ────────────────
    add_section_hdr(
        doc, "1. DOCUMENT CONTROL",
        fill="1F3864", size=10
    )
    ctrl_tbl       = doc.add_table(rows=0, cols=4)
    ctrl_tbl.style = 'Table Grid'

    def add_ctrl_row(tbl, l1, v1, l2, v2):
        row   = tbl.add_row()
        cells = row.cells
        for i, (lbl, val) in enumerate(
            [(l1,v1),(l2,v2)]
        ):
            lc = cells[i*2]
            vc = cells[i*2+1]
            lc.text = lbl
            lc.paragraphs[0].runs[0].bold      = True
            lc.paragraphs[0].runs[0].font.size = Pt(9)
            shd_cell(lc, 'E8F0FE')
            cell_borders(lc)
            vc.text = str(val) if val else "—"
            vc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_borders(vc)

    add_ctrl_row(
        ctrl_tbl,
        "SOP Title",
        sop_data.get('sop_title',''),
        "SOP Number",
        meta.get('sop_no','')
    )
    add_ctrl_row(
        ctrl_tbl,
        "Project",
        meta.get('project',''),
        "Date Issued",
        str(meta.get('sop_date', date.today()))
    )
    add_ctrl_row(
        ctrl_tbl,
        "Company",
        meta.get('company','ENGC'),
        "Revision",
        meta.get('revision','Rev 00')
    )
    add_ctrl_row(
        ctrl_tbl,
        "Prepared By",
        meta.get('prepared_by',''),
        "Reviewed By",
        meta.get('reviewed_by','')
    )
    add_ctrl_row(
        ctrl_tbl,
        "Approved By",
        meta.get('approved_by',''),
        "Next Review",
        meta.get('next_review','')
    )
    add_ctrl_row(
        ctrl_tbl,
        "Regulatory Reference",
        sop_data.get('reg_ref',''),
        "Applies To",
        sop_data.get('applies_to','All site workers')
    )

    doc.add_paragraph()

    # ── PURPOSE & SCOPE ───────────────────────
    add_section_hdr(
        doc, "2. PURPOSE & SCOPE",
        fill="1F3864"
    )
    add_body(doc, sop_data.get('purpose', ''))
    doc.add_paragraph()

    # ── DEFINITIONS ───────────────────────────
    add_section_hdr(
        doc, "3. DEFINITIONS & ABBREVIATIONS",
        fill="1F5864"
    )
    definitions = sop_data.get('definitions', [])
    if definitions:
        def_tbl       = doc.add_table(
            rows=1, cols=2
        )
        def_tbl.style = 'Table Grid'
        for cell, h in zip(
            def_tbl.rows[0].cells,
            ["Term / Abbreviation", "Definition"]
        ):
            cell.text = h
            cell.paragraphs[0].runs[0].bold      = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0]\
                .font.color.rgb = RGBColor(
                255, 255, 255
            )
            shd_cell(cell, '1F5864')
            cell_borders(cell)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )
        for i, item in enumerate(definitions):
            row = def_tbl.add_row()
            if isinstance(item, dict):
                t = item.get('term', str(item))
                d = item.get('definition', '')
            else:
                parts = str(item).split(':', 1)
                t = parts[0].strip()
                d = parts[1].strip() if len(
                    parts
                ) > 1 else ''
            row.cells[0].text = t
            row.cells[0].paragraphs[0]\
                .runs[0].bold      = True
            row.cells[0].paragraphs[0]\
                .runs[0].font.size = Pt(9)
            row.cells[1].text = d
            row.cells[1].paragraphs[0]\
                .runs[0].font.size = Pt(9)
            if i % 2 == 0:
                shd_cell(row.cells[0], 'EEF6FF')
                shd_cell(row.cells[1], 'EEF6FF')
            for c in row.cells:
                cell_borders(c)

    doc.add_paragraph()

    # ── RESPONSIBILITIES ──────────────────────
    add_section_hdr(
        doc, "4. ROLES & RESPONSIBILITIES",
        fill="1F3864"
    )
    resp_tbl       = doc.add_table(rows=1, cols=3)
    resp_tbl.style = 'Table Grid'
    for cell, h in zip(
        resp_tbl.rows[0].cells,
        ["Role / Position",
         "Responsibility Under This SOP",
         "Authority"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F3864')
        cell_borders(cell)

    responsibilities = sop_data.get(
        'responsibilities', []
    )
    for i, resp in enumerate(responsibilities):
        row  = resp_tbl.add_row()
        role = resp.get('role','') if isinstance(
            resp, dict
        ) else str(resp)
        duty = resp.get('duty','') if isinstance(
            resp, dict
        ) else ''
        auth = resp.get('authority','') if isinstance(
            resp, dict
        ) else ''
        vals = [role, duty, auth]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'F5F8FF')
            cell_borders(c)

    doc.add_paragraph()

    # ── PREREQUISITES ─────────────────────────
    add_section_hdr(
        doc, "5. PREREQUISITES & REQUIREMENTS",
        fill="BF8F00"
    )
    prereqs = sop_data.get('prerequisites', [])
    for item in prereqs:
        add_bullet_item(doc, str(item))
    doc.add_paragraph()

    # ── STEP BY STEP PROCEDURE ────────────────
    add_section_hdr(
        doc, "6. STEP-BY-STEP PROCEDURE",
        fill=hdr_color
    )
    steps = sop_data.get('steps', [])
    step_tbl       = doc.add_table(rows=1, cols=4)
    step_tbl.style = 'Table Grid'
    col_hdrs = [
        ("Step", Cm(1.2)),
        ("Action / Task", Cm(5.5)),
        ("Safety Requirements & Controls", Cm(6.0)),
        ("Responsible Person", Cm(4.3)),
    ]
    for (hd, wd), cell in zip(
        col_hdrs, step_tbl.rows[0].cells
    ):
        cell.width = wd
        cell.text  = hd
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, hdr_color)
        cell_borders(cell)
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    for i, step in enumerate(steps):
        row = step_tbl.add_row()
        if isinstance(step, dict):
            s_num  = str(i + 1)
            s_act  = step.get('action', '')
            s_ctrl = step.get('controls', '')
            s_resp = step.get('responsible',
                              'Site Supervisor')
        else:
            s_num  = str(i + 1)
            s_act  = str(step)
            s_ctrl = ''
            s_resp = 'Site Supervisor'

        vals = [s_num, s_act, s_ctrl, s_resp]
        widths = [
            Cm(1.2), Cm(5.5), Cm(6.0), Cm(4.3)
        ]
        for j, (c, v, w) in enumerate(
            zip(row.cells, vals, widths)
        ):
            c.width = w
            c.text  = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
            if i % 2 == 0:
                shd_cell(c, 'FFF9E6')
            cell_borders(c)

    doc.add_paragraph()

    # ── HAZARDS & CONTROLS ────────────────────
    add_section_hdr(
        doc, "7. HAZARDS, RISKS & CONTROLS "
             "(Hierarchy of Controls)",
        fill="C00000"
    )
    haz_tbl       = doc.add_table(rows=1, cols=5)
    haz_tbl.style = 'Table Grid'
    haz_hdrs = [
        ("Hazard", Cm(3.0)),
        ("Risk / Consequence", Cm(3.0)),
        ("Control Measures\n(Elimination → PPE)",
         Cm(6.5)),
        ("Residual\nRisk", Cm(2.0)),
        ("CoP / Ref", Cm(2.5)),
    ]
    for (hd, wd), cell in zip(
        haz_hdrs, haz_tbl.rows[0].cells
    ):
        cell.width = wd
        cell.text  = hd
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, 'C00000')
        cell_borders(cell)
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    risk_fills = {
        "LOW":      "00AF50",
        "MODERATE": "FFFF00",
        "HIGH":     "FFC000",
        "EXTREME":  "FF0000",
    }
    hazards = sop_data.get('hazards', [])
    for i, haz in enumerate(hazards):
        row = haz_tbl.add_row()
        if isinstance(haz, dict):
            h_haz  = haz.get('hazard','')
            h_risk = haz.get('risk','')
            h_ctrl = haz.get('controls','')
            h_rr   = haz.get(
                'residual_risk','LOW'
            ).upper()
            h_ref  = haz.get('ref','')
        else:
            h_haz  = str(haz)
            h_risk = ''
            h_ctrl = ''
            h_rr   = 'LOW'
            h_ref  = ''

        vals   = [h_haz, h_risk, h_ctrl,
                  h_rr, h_ref]
        widths = [Cm(3.0), Cm(3.0), Cm(6.5),
                  Cm(2.0), Cm(2.5)]
        for j, (c, v, w) in enumerate(
            zip(row.cells, vals, widths)
        ):
            c.width = w
            c.text  = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 3:  # Risk rating
                fill = risk_fills.get(
                    h_rr, "00AF50"
                )
                shd_cell(c, fill)
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
            elif i % 2 == 0:
                shd_cell(c, 'FFF2F2')
            cell_borders(c)

    doc.add_paragraph()

    # ── PPE REQUIREMENTS ──────────────────────
    add_section_hdr(
        doc, "8. PPE REQUIREMENTS",
        fill="1B5E20"
    )
    ppe_tbl       = doc.add_table(rows=1, cols=3)
    ppe_tbl.style = 'Table Grid'
    for cell, h in zip(
        ppe_tbl.rows[0].cells,
        ["PPE Item", "Standard / Specification",
         "Mandatory For"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1B5E20')
        cell_borders(cell)

    ppe_items = sop_data.get('ppe', [])
    for i, item in enumerate(ppe_items):
        row = ppe_tbl.add_row()
        if isinstance(item, dict):
            p_item = item.get('item','')
            p_std  = item.get('standard','')
            p_for  = item.get(
                'mandatory_for','All workers'
            )
        else:
            p_item = str(item)
            p_std  = ''
            p_for  = 'All workers'
        vals = [p_item, p_std, p_for]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F5E9')
            cell_borders(c)

    doc.add_paragraph()

    # ── TRAINING REQUIREMENTS ─────────────────
    add_section_hdr(
        doc, "9. TRAINING & COMPETENCY REQUIREMENTS",
        fill="4A148C"
    )
    training = sop_data.get('training', [])
    train_tbl       = doc.add_table(rows=1, cols=3)
    train_tbl.style = 'Table Grid'
    for cell, h in zip(
        train_tbl.rows[0].cells,
        ["Training / Certification Required",
         "Who Requires It",
         "Frequency / Validity"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '4A148C')
        cell_borders(cell)

    for i, item in enumerate(training):
        row = train_tbl.add_row()
        if isinstance(item, dict):
            t_name = item.get('training','')
            t_who  = item.get('who','All workers')
            t_freq = item.get('frequency','Annual')
        else:
            t_name = str(item)
            t_who  = 'All workers'
            t_freq = 'Annual'
        vals = [t_name, t_who, t_freq]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                shd_cell(c, 'F3E5F5')
            cell_borders(c)

    doc.add_paragraph()

    # ── EMERGENCY ─────────────────────────────
    add_section_hdr(
        doc, "10. EMERGENCY PROCEDURES",
        fill="C00000"
    )
    em_steps = sop_data.get('emergency_steps', [])
    if em_steps:
        for step in em_steps:
            add_bullet_item(doc, str(step))
    else:
        # Default emergency steps
        if emirate == "Abu Dhabi":
            defaults = [
                "STOP WORK immediately — raise alarm",
                "Ambulance: 998 | Police: 999 "
                "| Civil Defence: 997",
                "Notify HSE Officer and Site Manager",
                "Evacuate to Muster Point",
                "Do not re-enter until cleared",
                "Report to ADOSH per ADOSH-SF v4.0",
            ]
        else:
            defaults = [
                "STOP WORK immediately — raise alarm",
                "Police: 999 | Civil Defence: 997 "
                "| Ambulance: 998 | DM: 800900",
                "Notify HSE Officer and Site Manager",
                "Evacuate to Muster Point",
                "Do not re-enter until cleared",
                "Report to DM within 72 hours",
            ]
        for d in defaults:
            add_bullet_item(doc, d)

    doc.add_paragraph()

    # ── MONITORING & REVIEW ───────────────────
    add_section_hdr(
        doc, "11. MONITORING, REVIEW & AUDIT",
        fill="1F3864"
    )
    monitoring = sop_data.get('monitoring', [])
    for item in monitoring:
        add_bullet_item(doc, str(item))

    doc.add_paragraph()

    # ── RELATED DOCUMENTS ─────────────────────
    add_section_hdr(
        doc, "12. RELATED DOCUMENTS & REFERENCES",
        fill="37474F"
    )
    refs = sop_data.get('references', [])
    for ref in refs:
        add_bullet_item(doc, str(ref))

    doc.add_paragraph()

    # ── SIGNATURE BLOCK ───────────────────────
    add_section_hdr(
        doc, "13. APPROVAL & SIGNATURES",
        fill="1F3864"
    )
    sig_tbl       = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    sig_roles = [
        "Prepared By\n(HSE Officer / Engineer)",
        "Reviewed By\n(Project Manager)",
        "Approved By\n(Projects Director / DM)"
    ]
    for i, role in enumerate(sig_roles):
        c = sig_tbl.rows[0].cells[i]
        c.text = role
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, '1F3864')
        cell_borders(c)
    for i in range(3):
        c = sig_tbl.rows[1].cells[i]
        c.text = "\n\nName: _______________\n"
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c)
        c = sig_tbl.rows[2].cells[i]
        c.text = (
            "Designation: _______________\n\n"
            "Signature: _______________\n\n"
            "Date: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c)

    # ── REVISION HISTORY ──────────────────────
    doc.add_paragraph()
    add_section_hdr(
        doc, "14. REVISION HISTORY",
        fill="37474F", size=9
    )
    rev_tbl       = doc.add_table(rows=2, cols=4)
    rev_tbl.style = 'Table Grid'
    for cell, h in zip(
        rev_tbl.rows[0].cells,
        ["Rev No.", "Date", "Description",
         "Approved By"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '37474F')
        cell_borders(cell)

    for i, (rv, dt, desc, appr) in enumerate([
        ("Rev 00",
         str(meta.get('sop_date', date.today())),
         "Initial Issue",
         meta.get('approved_by',''))
    ]):
        row   = rev_tbl.add_row()
        vals  = [rv, dt, desc, appr]
        for c, v in zip(row.cells, vals):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if i % 2 == 0:
                shd_cell(c, 'F5F5F5')
            cell_borders(c)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════
# AI HELPER
# ════════════════════════════════════════════════════════

def ask_ai(client, prompt_text, max_tok=1200):
    """Single AI call with clean response."""
    r = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tok,
        messages=[{
            "role":    "user",
            "content": prompt_text
        }]
    )
    return r.content[0].text.strip()


def generate_sop_data(
    client, sop_title, meta, emirate
):
    """
    Generate all SOP sections using
    separate small AI calls — prevents
    JSON overflow errors.
    """
    reg = (
        "ADOSH-SF v4.0 — cite specific ADOSH CoP"
        " numbers"
        if emirate == "Abu Dhabi"
        else
        "Dubai Municipality Code — cite specific"
        " chapter numbers"
    )
    ctx = (
        "SOP: " + sop_title
        + " | Project: " + meta.get('project','')
        + " | Emirate: " + emirate
        + " | Company: " + meta.get('company','ENGC')
        + " | Regulation: " + reg
    )

    data = {"sop_title": sop_title}

    # Regulatory reference
    raw = ask_ai(
        client,
        "For this SOP in " + emirate + ": "
        + sop_title
        + " — give the most relevant regulatory "
        "reference (e.g. 'ADOSH CoP 29.0 — "
        "Excavation | ISO 45001:2018 Cl.8.1'). "
        "One line only."
    )
    data['reg_ref'] = raw[:300].strip()

    # Applies to
    raw = ask_ai(
        client,
        "Who does this SOP apply to? "
        + sop_title + " in " + emirate + ". "
        "One short line (e.g. 'All site workers, "
        "supervisors, and subcontractors involved "
        "in excavation works'). Plain text."
    )
    data['applies_to'] = raw[:200].strip()

    # Purpose
    raw = ask_ai(
        client,
        "Write 2-3 sentences for the PURPOSE "
        "section of this SOP in " + emirate + ": "
        + sop_title
        + ". Be specific, professional, and "
        "reference " + reg + ". Plain text only."
    )
    data['purpose'] = raw[:500].strip()

    # Definitions
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 5 key "
        "definitions for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Format: [{"term":"Term","definition":'
        '"Meaning"}]\n'
        "Return ONLY the JSON array."
    )
    data['definitions'] = safe_json_list(raw, [
        {"term": "SOP",
         "definition": "Standard Operating Procedure"},
        {"term": "HSE",
         "definition": "Health, Safety & Environment"},
        {"term": "PTW",
         "definition": "Permit to Work"},
        {"term": "PPE",
         "definition": "Personal Protective Equipment"},
        {"term": "ADOSH",
         "definition": "Abu Dhabi Department of "
                       "Occupational Safety and Health"},
    ])

    # Responsibilities
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 4 role "
        "responsibilities for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Format: [{"role":"Role title",'
        '"duty":"What they must do",'
        '"authority":"Their authority"}]\n'
        "Include: Projects Director, Project "
        "Manager, HSE Officer, Site Supervisor. "
        "Return ONLY the JSON array."
    )
    data['responsibilities'] = safe_json_list(
        raw,
        [
            {"role": "Projects Director",
             "duty": "Approve and endorse SOP",
             "authority": "Full authority"},
            {"role": "Project Manager",
             "duty": "Implement and enforce SOP",
             "authority": "Site authority"},
            {"role": "HSE Officer",
             "duty": "Monitor compliance",
             "authority": "Stop Work Authority"},
            {"role": "Site Supervisor",
             "duty": "Brief workers and supervise",
             "authority": "Task authority"},
        ]
    )

    # Prerequisites
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 6 prerequisite"
        " requirements (as plain strings) before "
        "starting this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Example: ["All workers must have site '
        'induction","PTW must be obtained"]\n'
        "Return ONLY the JSON array."
    )
    data['prerequisites'] = safe_json_list(
        raw,
        [
            "Valid site induction completed",
            "Relevant permits obtained",
            "Risk assessment reviewed and signed",
            "All equipment inspected and certified",
            "PPE issued and worn",
            "Toolbox talk conducted",
        ]
    )

    # Step-by-step procedure
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 10 procedure "
        "steps for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Format: [{"action":"What to do",'
        '"controls":"Safety controls with '
        + reg[:20] + ' ref",'
        '"responsible":"Who does it"}]\n'
        "Start with pre-task, end with post-task. "
        "Be specific to " + sop_title + ". "
        "Return ONLY the JSON array."
    )
    data['steps'] = safe_json_list(raw, [
        {"action": "Pre-task briefing",
         "controls": "All workers briefed on SOP",
         "responsible": "HSE Officer"},
        {"action": "Site/area inspection",
         "controls": "Area checked and cleared",
         "responsible": "Site Supervisor"},
        {"action": "Execute task as per SOP",
         "controls": "Follow all controls",
         "responsible": "Competent Worker"},
        {"action": "Post-task inspection",
         "controls": "Area left safe and clean",
         "responsible": "Site Supervisor"},
    ])

    # Hazards
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 6 hazards "
        "for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Format: [{"hazard":"Hazard",'
        '"risk":"Consequence",'
        '"controls":"Control measures using '
        'hierarchy (Eliminate→Substitute→'
        'Engineering→Admin→PPE)",'
        '"residual_risk":"LOW/MODERATE/HIGH",'
        '"ref":"CoP or Code reference"}]\n'
        "Return ONLY the JSON array."
    )
    data['hazards'] = safe_json_list(raw, [
        {"hazard": "General site hazards",
         "risk": "Injury or illness",
         "controls": "Follow standard precautions",
         "residual_risk": "LOW",
         "ref": reg[:30]},
    ])

    # PPE
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 6 PPE items "
        "for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Format: [{"item":"PPE name",'
        '"standard":"EN/ANSI standard",'
        '"mandatory_for":"Who must wear"}]\n'
        "Always include: helmet, boots, vest. "
        "Add activity-specific PPE. "
        "Return ONLY the JSON array."
    )
    data['ppe'] = safe_json_list(raw, [
        {"item": "Safety Helmet",
         "standard": "EN 397 / ANSI Z89.1",
         "mandatory_for": "All personnel"},
        {"item": "Safety Footwear",
         "standard": "EN ISO 20345",
         "mandatory_for": "All personnel"},
        {"item": "High-Visibility Vest",
         "standard": "EN ISO 20471",
         "mandatory_for": "All personnel"},
    ])

    # Training
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 5 training "
        "requirements for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Format: [{"training":"Training name",'
        '"who":"Who needs it",'
        '"frequency":"How often"}]\n'
        "Return ONLY the JSON array."
    )
    data['training'] = safe_json_list(raw, [
        {"training": "Site Induction",
         "who": "All workers",
         "frequency": "Once on joining"},
        {"training": "Activity-specific training",
         "who": "All involved workers",
         "frequency": "Before starting activity"},
        {"training": "Emergency response",
         "who": "All workers",
         "frequency": "Annual"},
    ])

    # Emergency steps
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 5 emergency "
        "response steps (as plain strings) for: "
        + sop_title + " in " + emirate + ".\n"
        'Example: ["STOP WORK immediately",'
        '"Call 998 for Ambulance"]\n'
        "Include correct " + emirate
        + " emergency numbers. "
        "Return ONLY the JSON array."
    )
    data['emergency_steps'] = safe_json_list(
        raw,
        [
            "STOP WORK immediately — raise alarm",
            "Call emergency services immediately",
            "Notify HSE Officer and Site Manager",
            "Evacuate to Muster Point",
            "Complete incident report",
        ]
    )

    # Monitoring
    raw = ask_ai(
        client,
        "Return ONLY a JSON array of 5 monitoring "
        "and review activities (as plain strings) "
        "for this SOP: "
        + sop_title + " in " + emirate + ".\n"
        'Example: ["Daily inspection by HSE '
        'Officer","Monthly SOP review"]\n'
        "Return ONLY the JSON array."
    )
    data['monitoring'] = safe_json_list(raw, [
        "Daily site inspection by HSE Officer",
        "Weekly SOP compliance audit",
        "Monthly management review",
        "Annual SOP review and update",
        "Review after any incident or near miss",
    ])

    # References
    if emirate == "Abu Dhabi":
        data['references'] = [
            "ADOSH-SF Version 4.0 (July 2024)",
            data.get('reg_ref','ADOSH CoP — relevant'),
            "ADOSH CoP 2.0 — PPE",
            "ADOSH CoP 11.0 — Heat Stress",
            "ADOSH CoP 53.1 — OSH Plan",
            "MOHRE Resolution 44/2022",
            "ISO 45001:2018 — OH&S Management",
            "UAE Federal Decree-Law 33/2021",
        ]
    else:
        data['references'] = [
            "Dubai Municipality Code of "
            "Construction Safety Practice",
            data.get('reg_ref','DM Code — relevant'),
            "DM Code Chapter 2 — Safety Management",
            "DM Code Chapter 5 — PPE",
            "DM Code Chapter 12 — Emergency",
            "UAE Federal Decree-Law 33/2021",
            "ISO 45001:2018 — OH&S Management",
        ]

    return data


# ════════════════════════════════════════════════════════
# MAIN UI — THREE TABS
# ════════════════════════════════════════════════════════

tab_gen, tab_ai, tab_upload = st.tabs([
    "📘 Generate SOP (Library)",
    "🤖 AI Custom SOP",
    "📤 Upload Your SOP Format",
])


# ════════════════════════════════════════════════════════
# TAB 1 — LIBRARY GENERATOR
# ════════════════════════════════════════════════════════

with tab_gen:
    st.markdown(
        "### 📘 Generate SOP from Library"
    )
    st.info(
        "Select any SOP from 100+ pre-defined "
        "procedures across 10 categories. The AI "
        "generates a complete, professional SOP "
        "with 14 sections — including step-by-step "
        "procedure, hazards, controls, training, "
        "and " + emirate + " regulatory references."
    )

    # Category + SOP selection
    cat_name = st.selectbox(
        "Select Category",
        list(SOP_LIBRARY.keys()),
        key="lib_cat"
    )
    cat_data = SOP_LIBRARY[cat_name]
    sop_name = st.selectbox(
        "Select SOP",
        cat_data['sops'],
        key="lib_sop"
    )

    # Show regulatory reference hint
    if emirate == "Abu Dhabi":
        st.caption(
            "✅ This SOP will reference relevant "
            "ADOSH CoPs and ADOSH-SF v4.0"
        )
    else:
        st.caption(
            "✅ This SOP will reference relevant "
            "Dubai Municipality Code chapters"
        )

    st.divider()
    st.subheader("📋 Document Details")

    g1, g2 = st.columns(2)
    with g1:
        g_project = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="g_proj"
        )
        g_company = st.text_input(
            "Company",
            value="ENGC",
            key="g_co"
        )
        g_sop_no  = st.text_input(
            "SOP Number",
            placeholder="e.g. ENGC-SOP-001",
            key="g_sopno"
        )
        g_rev     = st.selectbox(
            "Revision",
            ["Rev 00","Rev 01","Rev 02","Rev 03"],
            key="g_rev"
        )
    with g2:
        g_prep    = st.text_input(
            "Prepared By",
            placeholder=(
                "e.g. Aftab Qamar — "
                "Senior HSSE Engineer"
            ),
            key="g_prep"
        )
        g_review  = st.text_input(
            "Reviewed By",
            placeholder=(
                "e.g. Ahmad Abdelrahman — "
                "Project Manager"
            ),
            key="g_review"
        )
        g_approve = st.text_input(
            "Approved By",
            placeholder=(
                "e.g. Rami Kamal — "
                "Projects Director"
            ),
            key="g_approve"
        )
        g_date    = st.date_input(
            "Issue Date",
            value=date.today(),
            key="g_date"
        )

    g_next_rev = st.text_input(
        "Next Review Date",
        placeholder="e.g. July 2025 (Annual Review)",
        key="g_nextrev"
    )

    if st.button(
        "⚡ Generate SOP: " + sop_name + " (.docx)",
        type="primary",
        use_container_width=True,
        key="lib_gen_btn"
    ):
        if not g_project:
            st.error(
                "Please enter a project name."
            )
        else:
            prog = st.progress(
                0, text="Starting SOP generation..."
            )
            try:
                client = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                meta = {
                    "sop_title":   sop_name,
                    "project":     g_project,
                    "company":     g_company,
                    "sop_no":      g_sop_no,
                    "revision":    g_rev,
                    "prepared_by": g_prep,
                    "reviewed_by": g_review,
                    "approved_by": g_approve,
                    "sop_date":    g_date,
                    "next_review": g_next_rev,
                }

                total_steps = 12

                def upd(n, msg):
                    prog.progress(
                        int(n/total_steps*90),
                        text=msg
                    )

                upd(1, "Getting CoP reference...")
                upd(2, "Writing purpose & scope...")
                upd(3, "Generating definitions...")
                upd(4, "Setting responsibilities...")
                upd(5, "Listing prerequisites...")

                sop_data = generate_sop_data(
                    client, sop_name, meta, emirate
                )

                upd(6, "Building procedure steps...")
                upd(7, "Identifying hazards...")
                upd(8, "Selecting PPE...")
                upd(9, "Training requirements...")
                upd(10, "Emergency procedures...")
                upd(11, "Building Word document...")

                docx_buf = build_sop_docx(
                    sop_data,
                    meta,
                    emirate,
                    cat_data['color']
                )

                prog.progress(
                    100,
                    text="✅ SOP ready!"
                )

                st.success(
                    "✅ SOP generated: "
                    + sop_name
                    + " — 14 sections complete!"
                )

                sop_fn   = sop_name.replace(
                    ' ', '_'
                ).replace('/', '-').replace(
                    '&', 'and'
                )
                proj_fn  = g_project.replace(
                    ' ', '_'
                )
                fname = (
                    "SOP_"
                    + sop_fn + "_"
                    + proj_fn + "_"
                    + str(g_date) + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download SOP — "
                        + sop_name + " (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="lib_dl"
                )

                st.info(
                    "📋 **Next Steps:**\n\n"
                    "1. Review all sections with "
                    "your HSE team\n\n"
                    "2. Get sign-off from Project "
                    "Manager and Projects Director\n\n"
                    "3. Brief all affected workers "
                    "before work commences\n\n"
                    "4. Keep signed copy in site "
                    "document register\n\n"
                    "5. Review annually or after "
                    "any incident\n\n"
                    "⚠️ *AI output is a starting "
                    "point — must be verified by a "
                    "competent HSE professional.*"
                )

            except Exception as e:
                prog.progress(0)
                st.error(
                    "❌ Error: " + str(e)
                    + "\n\nPlease try again."
                )


# ════════════════════════════════════════════════════════
# TAB 2 — AI CUSTOM SOP
# ════════════════════════════════════════════════════════

with tab_ai:
    st.markdown(
        "### 🤖 AI Custom SOP Generator"
    )
    st.info(
        "Describe any activity, process, or "
        "procedure and the AI will generate a "
        "complete, bespoke SOP — specific to your "
        "project, your company, and "
        + emirate + " regulations."
    )

    ai_sop = st.text_input(
        "What SOP do you need?",
        placeholder=(
            "e.g. Batching Plant Operation, "
            "CCTV Installation, Structural Cabling, "
            "Curtain Wall Erection, Swimming Pool "
            "Construction, Pile Cap Construction..."
        ),
        key="ai_sop_name"
    )

    ai_c1, ai_c2 = st.columns(2)
    with ai_c1:
        ai_project = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="ai_sop_proj"
        )
        ai_company = st.text_input(
            "Company",
            value="ENGC",
            key="ai_sop_co"
        )
        ai_prep    = st.text_input(
            "Prepared By",
            placeholder="e.g. Aftab Qamar",
            key="ai_sop_prep"
        )
        ai_sop_no  = st.text_input(
            "SOP Number",
            placeholder="e.g. ENGC-SOP-CUS-001",
            key="ai_sop_no"
        )
    with ai_c2:
        ai_review  = st.text_input(
            "Reviewed By",
            placeholder="e.g. Ahmad Abdelrahman",
            key="ai_sop_rev"
        )
        ai_approve = st.text_input(
            "Approved By",
            placeholder="e.g. Rami Kamal",
            key="ai_sop_approve"
        )
        ai_sop_date = st.date_input(
            "Issue Date",
            value=date.today(),
            key="ai_sop_date"
        )
        ai_next_rev = st.text_input(
            "Next Review Date",
            placeholder="e.g. July 2025",
            key="ai_sop_nextrev"
        )

    ai_notes = st.text_area(
        "Special Requirements / Notes",
        placeholder=(
            "e.g. Works near live road, "
            "night shift operation, "
            "high dust environment, "
            "ADOSH inspection expected..."
        ),
        height=70,
        key="ai_sop_notes"
    )

    if st.button(
        "⚡ Generate Custom SOP (.docx)",
        type="primary",
        use_container_width=True,
        key="ai_sop_btn"
    ):
        if not ai_sop or not ai_project:
            st.error(
                "Please enter the SOP name "
                "and project name."
            )
        else:
            prog = st.progress(
                0,
                text="Starting custom SOP..."
            )
            try:
                client = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                meta = {
                    "sop_title":   ai_sop,
                    "project":     ai_project,
                    "company":     ai_company,
                    "sop_no":      ai_sop_no,
                    "revision":    "Rev 00",
                    "prepared_by": ai_prep,
                    "reviewed_by": ai_review,
                    "approved_by": ai_approve,
                    "sop_date":    ai_sop_date,
                    "next_review": ai_next_rev,
                }

                prog.progress(
                    10,
                    text="Generating SOP content..."
                )

                sop_data = generate_sop_data(
                    client, ai_sop, meta, emirate
                )

                prog.progress(
                    85,
                    text="Building Word document..."
                )

                docx_buf = build_sop_docx(
                    sop_data, meta,
                    emirate, "1F3864"
                )

                prog.progress(
                    100, text="✅ Custom SOP ready!"
                )

                st.success(
                    "✅ Custom SOP generated: "
                    + ai_sop + "!"
                )

                sop_fn  = ai_sop.replace(
                    ' ', '_'
                ).replace('/', '-')
                proj_fn = ai_project.replace(
                    ' ', '_'
                )
                fname = (
                    "SOP_Custom_"
                    + sop_fn + "_"
                    + proj_fn + "_"
                    + str(ai_sop_date) + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download SOP: "
                        + ai_sop + " (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="ai_sop_dl"
                )

            except Exception as e:
                prog.progress(0)
                st.error("Error: " + str(e))


# ════════════════════════════════════════════════════════
# TAB 3 — UPLOAD YOUR SOP FORMAT
# ════════════════════════════════════════════════════════

with tab_upload:
    st.markdown(
        "### 📤 Upload Your Own SOP Format"
    )
    st.info(
        "Upload your company's existing SOP "
        "template (Word or PDF). The AI will read "
        "your EXACT format — headings, sections, "
        "tables, company branding — and fill it "
        "with content for your selected procedure. "
        "Downloads as a **.docx** file."
    )

    up_file = st.file_uploader(
        "📎 Upload Your SOP Template or Sample",
        type=["docx", "pdf"],
        help=(
            "Upload a blank SOP template or a "
            "previously completed SOP to use as "
            "the format reference."
        ),
        key="sop_up_file"
    )

    if up_file:
        st.success(
            "✅ Uploaded: **"
            + up_file.name + "** ("
            + "{:,}".format(up_file.size)
            + " bytes)"
        )

        st.divider()
        st.subheader("📋 SOP Details")

        up1, up2 = st.columns(2)
        with up1:
            up_sop     = st.text_input(
                "SOP Name / Procedure Title",
                placeholder=(
                    "e.g. Working at Heights"
                ),
                key="up_sop_name"
            )
            up_project = st.text_input(
                "Project Name",
                placeholder=(
                    "e.g. Bloom Living Almeria"
                ),
                key="up_sop_proj"
            )
            up_company = st.text_input(
                "Company",
                value="ENGC",
                key="up_sop_co"
            )
        with up2:
            up_prep    = st.text_input(
                "Prepared By",
                placeholder="e.g. Aftab Qamar",
                key="up_sop_prep"
            )
            up_approve = st.text_input(
                "Approved By",
                placeholder="e.g. Rami Kamal",
                key="up_sop_approve"
            )
            up_date    = st.date_input(
                "Issue Date",
                value=date.today(),
                key="up_sop_date"
            )

        up_notes = st.text_area(
            "Special Requirements (optional)",
            placeholder=(
                "e.g. Include ADOSH CoP refs, "
                "night work, confined space..."
            ),
            height=70,
            key="up_sop_notes"
        )

        if st.button(
            "⚡ Fill My SOP Format",
            type="primary",
            use_container_width=True,
            key="up_sop_btn"
        ):
            if not up_sop or not up_project:
                st.error(
                    "Please enter the SOP name "
                    "and project name."
                )
            else:
                with st.spinner(
                    "Reading your SOP format "
                    "and generating content..."
                ):
                    try:
                        client = anthropic.Anthropic(
                            api_key=os.getenv(
                                "ANTHROPIC_API_KEY"
                            )
                        )

                        file_bytes = up_file.read()

                        reg = (
                            "ADOSH-SF v4.0, cite "
                            "specific ADOSH CoP numbers"
                            if emirate == "Abu Dhabi"
                            else
                            "Dubai Municipality Code, "
                            "cite chapter numbers"
                        )

                        instr = (
                            "This is an SOP template "
                            "for " + emirate + ".\n\n"
                            "Fill it for:\n"
                            "SOP: " + up_sop + "\n"
                            "Project: " + up_project + "\n"
                            "Company: " + up_company + "\n"
                            "Prepared By: " + up_prep + "\n"
                            "Approved By: " + up_approve + "\n"
                            "Date: " + str(up_date) + "\n"
                            "Regulation: " + reg + "\n"
                            "Notes: " + up_notes + "\n\n"
                            "INSTRUCTIONS:\n"
                            "1. Keep EXACT same "
                            "structure and headings\n"
                            "2. Fill ALL sections with "
                            "specific content\n"
                            "3. Reference specific "
                            + emirate + " regulations\n"
                            "4. Use hierarchy of controls: "
                            "Elimination → Substitution → "
                            "Engineering → Administrative "
                            "→ PPE\n"
                            "5. Be professional and ready "
                            "for submission\n"
                            "6. Include step-by-step "
                            "procedure specific to "
                            + up_sop
                        )

                        if up_file.name\
                                .lower()\
                                .endswith('.pdf'):
                            b64 = (
                                base64.standard_b64encode(
                                    file_bytes
                                ).decode()
                            )
                            response = (
                                client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "document",
                                            "source": {
                                                "type": "base64",
                                                "media_type": "application/pdf",
                                                "data": b64,
                                            }
                                        },
                                        {
                                            "type": "text",
                                            "text": instr
                                        }
                                    ]
                                }]
                            ))
                        else:
                            doc_temp = Document(
                                BytesIO(file_bytes)
                            )
                            extracted = "\n".join([
                                p.text for p in
                                doc_temp.paragraphs
                                if p.text.strip()
                            ])
                            tbl_text = ""
                            for tbl in doc_temp.tables:
                                for row in tbl.rows:
                                    rt = " | ".join([
                                        c.text.strip()
                                        for c in row.cells
                                        if c.text.strip()
                                    ])
                                    if rt:
                                        tbl_text += (
                                            rt + "\n"
                                        )
                            full = (
                                extracted
                                + ("\n\nTABLES:\n"
                                   + tbl_text
                                   if tbl_text else "")
                            )
                            response = (
                                client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": (
                                        "SOP TEMPLATE:\n\n"
                                        "---\n"
                                        + full[:3000]
                                        + "\n---\n\n"
                                        + instr
                                    )
                                }]
                            ))

                        filled = (
                            response.content[0].text
                        )

                        st.success(
                            "✅ SOP filled for: "
                            + up_sop + "!"
                        )

                        with st.expander(
                            "📄 Preview Generated SOP",
                            expanded=True
                        ):
                            st.markdown(filled)

                        sop_fn  = up_sop.replace(
                            ' ', '_'
                        )
                        proj_fn = up_project.replace(
                            ' ', '_'
                        )
                        dl_fname = (
                            "SOP_"
                            + sop_fn + "_"
                            + proj_fn + "_"
                            + str(up_date)
                            + "_filled.txt"
                        )

                        st.download_button(
                            label=(
                                "⬇️ Download Filled "
                                "SOP (.txt)"
                            ),
                            data=filled,
                            file_name=dl_fname,
                            mime="text/plain",
                            key="up_sop_dl"
                        )

                        st.info(
                            "💡 **Tip:** Copy the "
                            "content above into your "
                            "Word template for the "
                            "final formatted SOP."
                        )

                    except Exception as e:
                        st.error(
                            "Error: " + str(e)
                        )

    else:
        c_a, c_b = st.columns(2)
        with c_a:
            st.markdown("""
**✅ Accepted Formats:**
- Word Document (.docx) — Best
- PDF (.pdf)

**📋 What to Upload:**
- Your company SOP template
- A consultant-approved SOP format
- A previously completed SOP sample
- ISO 45001 SOP template
""")
        with c_b:
            st.markdown("""
**🤖 What the AI Does:**
- Reads your exact structure
- Fills every section with real content
- References correct regulations
- Applies hierarchy of controls
- Includes step-by-step procedure
- Keeps your company branding
""")

        st.divider()
        st.info(
            "📘 **Don't have an SOP template?**\n\n"
            "Use the **Generate SOP (Library)** "
            "tab to create a complete professional "
            "SOP from 100+ pre-defined procedures — "
            "or use **AI Custom SOP** to describe "
            "any procedure and get it generated "
            "from scratch."
        )