import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import json
import re

load_dotenv()

st.set_page_config(
    page_title="Emergency Response Plan",
    page_icon="🚨"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning(
        "Please go to Home and select "
        "your Emirate first."
    )
    st.stop()

flag = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(
    "🚨 Emergency Response Plan — "
    + flag + " " + emirate
)

if emirate == "Abu Dhabi":
    st.markdown(
        "**Reference: ADOSH CoP 44.0 — Emergency "
        "Management | ADOSH CoP 53.1 | "
        "UAE Civil Defence Standards | "
        "ADOSH-SF v4.0**"
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality Code "
        "Chapter 12 — Emergency Procedures | "
        "Dubai Civil Defence (DCD) Requirements | "
        "UAE Federal Law No. 33/2021**"
    )

st.info(
    "🚨 **Critical Requirement:** Every "
    "construction project in UAE must have an "
    "approved Emergency Response Plan (ERP) "
    "before work commences. The ERP must be "
    "reviewed with all workers during induction, "
    "posted at prominent site locations, and "
    "drills must be conducted at least quarterly. "
    "\n\n"
    "⚠️ *Always verify this plan with a "
    "competent HSE professional and obtain "
    "consultant/authority approval before use.*"
)

st.divider()


# ════════════════════════════════════════════════════════
# EMERGENCY NUMBERS — EMIRATES SPECIFIC
# ════════════════════════════════════════════════════════

EMERGENCY_NUMBERS = {
    "Abu Dhabi": {
        "Police":          "999",
        "Ambulance":       "998",
        "Civil Defence":   "997",
        "ADNOC Emergency": "800 2332",
        "Etisalat Faults": "101",
        "ADDC Emergency":  "800 2332",
        "Gas Emergency":   "800 2332",
        "ADSSC Emergency": "800 1 (Sewerage)",
        "ADOSH Hotline":   "800 60",
        "Poison Control":  "+971 2 5037800",
    },
    "Dubai": {
        "Police":                  "999",
        "Ambulance":               "998",
        "Civil Defence (Fire)":    "997",
        "Dubai Municipality":      "800 900",
        "DEWA Emergency":          "991",
        "Gas Emergency (EMGAS)":   "800 36427",
        "Dubai Police HSE":        "800 4353",
        "DM Environment Dept":     "800 900",
        "Poison Control (DHA)":    "800 342",
        "Coast Guard (Maritime)":  "800 4354",
    }
}

# ════════════════════════════════════════════════════════
# EMERGENCY SCENARIOS
# ════════════════════════════════════════════════════════

SCENARIOS = {
    "🔥 Fire & Explosion": {
        "color": "FF0000",
        "cop": (
            "ADOSH CoP 19.0 — Fire Safety"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 10 — Fire Safety"
        ),
    },
    "🚑 Medical Emergency / Serious Injury": {
        "color": "C00000",
        "cop": (
            "ADOSH CoP 3.0 — First Aid"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 12 — Emergency"
        ),
    },
    "🏗️ Structural Collapse": {
        "color": "5D4037",
        "cop": (
            "ADOSH CoP 23.0 — Heights | "
            "CoP 40.0 — Formwork"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 6 — Structures"
        ),
    },
    "⚡ Electrical Emergency": {
        "color": "F57F17",
        "cop": (
            "ADOSH CoP 16.0 — Electrical Safety"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 9 — Electrical"
        ),
    },
    "☠️ Hazardous Substance Release": {
        "color": "1B5E20",
        "cop": (
            "ADOSH CoP 7.0 — Chemicals"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 11 — Chemicals"
        ),
    },
    "⛏️ Excavation Collapse / Cave-In": {
        "color": "795548",
        "cop": (
            "ADOSH CoP 29.0 — Excavation"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 7 — Excavation"
        ),
    },
    "🏗️ Crane / Lifting Failure": {
        "color": "4A148C",
        "cop": (
            "ADOSH CoP 34.0 — Lifting"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 8 — Lifting"
        ),
    },
    "🌊 Flood / Heavy Rain": {
        "color": "0D47A1",
        "cop": (
            "ADOSH CoP 53.1 — OSH Plan"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 12 — Emergency"
        ),
    },
    "🌪️ Severe Dust Storm / Sandstorm": {
        "color": "FF6F00",
        "cop": (
            "ADOSH CoP 11.0 — Heat & Dust | "
            "ADOSH CoP 53.1"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 3 — Environment"
        ),
    },
    "🚗 Vehicle / Plant Accident": {
        "color": "BF360C",
        "cop": (
            "ADOSH CoP 30.0 — Traffic"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 13 — Traffic"
        ),
    },
    "🌡️ Heat Stroke / Heat Emergency": {
        "color": "E65100",
        "cop": (
            "ADOSH CoP 11.0 — Heat Stress | "
            "MOHRE Res. 44/2022"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 4 — Welfare"
        ),
    },
    "⚠️ Confined Space Rescue": {
        "color": "37474F",
        "cop": (
            "ADOSH CoP 27.0 — Confined Spaces"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 7 — Confined Space"
        ),
    },
    "💀 Fatality / Major Incident": {
        "color": "000000",
        "cop": (
            "ADOSH-SF v4.0 Section 8 | "
            "ADOSH CoP 53.1"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 12 | "
                 "UAE Federal Law 33/2021"
        ),
    },
    "☢️ Gas Leak / Utility Strike": {
        "color": "558B2F",
        "cop": (
            "ADOSH CoP 16.0 | CoP 29.0"
            if emirate == "Abu Dhabi"
            else "DM Code Chapter 9 | "
                 "DEWA Emergency Protocol"
        ),
    },
}


# ════════════════════════════════════════════════════════
# DOCX HELPERS
# ════════════════════════════════════════════════════════

def shd_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)


def shd_cell(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def cell_borders(cell, color="C00000"):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def section_hdr(doc, text, fill="C00000",
                size=11):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    shd_para(p, fill)
    run           = p.add_run("  " + text.upper())
    run.bold      = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p


def add_body(doc, text, size=10,
             bold=False, indent=0.5):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent)
    run           = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold      = bold
    return p


def add_bullet(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    return p


def safe_list(raw, fallback=None):
    """Robust JSON list parser with fallback."""
    if fallback is None:
        fallback = []
    raw = re.sub(
        r'```json|```', '', raw
    ).strip()
    try:
        r = json.loads(raw)
        if isinstance(r, list):
            return r
        if isinstance(r, dict):
            for v in r.values():
                if isinstance(v, list):
                    return v
    except Exception:
        pass
    s = raw.find('[')
    e = raw.rfind(']') + 1
    if s != -1 and e > 0:
        try:
            r = json.loads(raw[s:e])
            if isinstance(r, list):
                return r
        except Exception:
            pass
    objs   = re.findall(r'\{[^\{\}]+\}', raw)
    parsed = []
    for o in objs:
        try:
            parsed.append(json.loads(o))
        except Exception:
            pass
    return parsed if parsed else fallback


def ask(client, prompt, max_tok=1000):
    """Single AI call — short and reliable."""
    r = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tok,
        messages=[{
            "role":    "user",
            "content": prompt
        }]
    )
    return r.content[0].text.strip()


# ════════════════════════════════════════════════════════
# BUILD FULL ERP DOCX
# ════════════════════════════════════════════════════════

def build_erp_docx(erp_data, meta,
                   emirate, scenarios_list):
    doc = Document()

    # A4 Portrait
    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    nums = EMERGENCY_NUMBERS.get(emirate, {})

    # ── COVER PAGE ────────────────────────────
    cover           = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(cover, '8B0000')
    cr = cover.add_run(
        "\n\n  🚨  EMERGENCY RESPONSE PLAN  🚨"
        "\n\n  "
        + meta.get('company', 'ENGC')
        + "  \n\n"
    )
    cr.bold           = True
    cr.font.size      = Pt(20)
    cr.font.color.rgb = RGBColor(255, 255, 255)

    sub           = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(sub, 'C00000')
    sr = sub.add_run(
        "\n  PROJECT: "
        + meta.get('project', '').upper()
        + "  \n  "
        + meta.get('location', '').upper()
        + "  \n  "
        + emirate.upper()
        + ", UAE  \n"
    )
    sr.bold           = True
    sr.font.size      = Pt(14)
    sr.font.color.rgb = RGBColor(255, 255, 255)

    # Emergency numbers on cover
    doc.add_paragraph()
    emg_cover = doc.add_paragraph()
    shd_para(emg_cover, '1A1A1A')
    er = emg_cover.add_run(
        "  🚨 EMERGENCY NUMBERS — CALL IMMEDIATELY:"
    )
    er.bold           = True
    er.font.size      = Pt(11)
    er.font.color.rgb = RGBColor(255, 200, 0)

    num_tbl       = doc.add_table(
        rows=0, cols=4
    )
    num_tbl.style = 'Table Grid'
    num_items     = list(nums.items())
    for i in range(0, len(num_items), 2):
        row   = num_tbl.add_row()
        cells = row.cells
        for j in range(2):
            if i + j < len(num_items):
                svc, num = num_items[i + j]
                lc = cells[j * 2]
                vc = cells[j * 2 + 1]
                lc.text = svc
                lc.paragraphs[0].runs[0].bold = True
                lc.paragraphs[0].runs[0]\
                    .font.size = Pt(9)
                lc.paragraphs[0].runs[0]\
                    .font.color.rgb = RGBColor(
                    255, 255, 255
                )
                shd_cell(lc, '1A1A1A')
                cell_borders(lc, 'C00000')
                vc.text = "📞 " + num
                vc.paragraphs[0].runs[0]\
                    .font.size = Pt(10)
                vc.paragraphs[0].runs[0].bold = True
                vc.paragraphs[0].runs[0]\
                    .font.color.rgb = RGBColor(
                    255, 50, 50
                )
                shd_cell(vc, '2A2A2A')
                cell_borders(vc, 'C00000')

    doc.add_page_break()

    # ── DOCUMENT CONTROL ──────────────────────
    section_hdr(
        doc, "1. DOCUMENT CONTROL",
        fill="1F3864", size=10
    )
    ctrl_tbl       = doc.add_table(rows=0, cols=4)
    ctrl_tbl.style = 'Table Grid'

    def ctrl_row(tbl, l1, v1, l2, v2):
        row   = tbl.add_row()
        cells = row.cells
        for i, (l, v) in enumerate(
            [(l1, v1), (l2, v2)]
        ):
            lc = cells[i * 2]
            vc = cells[i * 2 + 1]
            lc.text = l
            lc.paragraphs[0].runs[0].bold      = True
            lc.paragraphs[0].runs[0].font.size = Pt(9)
            shd_cell(lc, 'FFEBEE')
            cell_borders(lc, '1F3864')
            vc.text = str(v) if v else "—"
            vc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_borders(vc, '1F3864')

    reg_ref = (
        "ADOSH CoP 44.0 | CoP 53.1 | "
        "ADOSH-SF v4.0 | UAE Civil Defence"
        if emirate == "Abu Dhabi"
        else "DM Code Ch.12 | Dubai Civil Defence"
             " | UAE Fed. Law 33/2021"
    )

    ctrl_row(ctrl_tbl,
             "Document Title",
             "Emergency Response Plan (ERP)",
             "Document No.", meta.get('doc_no',''))
    ctrl_row(ctrl_tbl,
             "Project", meta.get('project',''),
             "Date Issued",
             str(meta.get('erp_date', date.today())))
    ctrl_row(ctrl_tbl,
             "Location",
             meta.get('location',''),
             "Revision",
             meta.get('revision','Rev 00'))
    ctrl_row(ctrl_tbl,
             "Company",
             meta.get('company','ENGC'),
             "Prepared By",
             meta.get('prepared_by',''))
    ctrl_row(ctrl_tbl,
             "Emergency Coordinator",
             meta.get('emergency_coord',''),
             "Deputy Coordinator",
             meta.get('deputy_coord',''))
    ctrl_row(ctrl_tbl,
             "Regulatory Reference",
             reg_ref,
             "Review Frequency",
             "Quarterly or after any incident")

    doc.add_paragraph()

    # ── PURPOSE & SCOPE ───────────────────────
    section_hdr(
        doc, "2. PURPOSE, SCOPE & OBJECTIVES",
        fill="8B0000"
    )
    add_body(doc, erp_data.get('purpose', ''))
    doc.add_paragraph()

    # ── KEY PERSONNEL ─────────────────────────
    section_hdr(
        doc, "3. EMERGENCY RESPONSE TEAM "
             "& KEY CONTACTS",
        fill="1F3864"
    )
    ert_tbl       = doc.add_table(rows=1, cols=5)
    ert_tbl.style = 'Table Grid'
    for cell, h in zip(
        ert_tbl.rows[0].cells,
        ["Role / Title", "Name",
         "Mobile / Radio",
         "Responsibility",
         "Backup / Deputy"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F3864')
        cell_borders(cell, '1F3864')
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    personnel = erp_data.get('ert_personnel', [])
    for i, p in enumerate(personnel):
        row  = ert_tbl.add_row()
        vals = [
            p.get('role',''),
            p.get('name',''),
            p.get('contact',''),
            p.get('responsibility',''),
            p.get('backup',''),
        ]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'FFEBEE')
            cell_borders(c, '1F3864')

    doc.add_paragraph()

    # ── EMERGENCY NUMBERS ─────────────────────
    section_hdr(
        doc,
        "4. EMERGENCY CONTACT NUMBERS — "
        + emirate.upper(),
        fill="8B0000"
    )

    em_tbl       = doc.add_table(rows=0, cols=4)
    em_tbl.style = 'Table Grid'
    num_items2   = list(nums.items())
    for i in range(0, len(num_items2), 2):
        row   = em_tbl.add_row()
        cells = row.cells
        for j in range(2):
            if i + j < len(num_items2):
                svc, num = num_items2[i + j]
                lc = cells[j * 2]
                vc = cells[j * 2 + 1]
                lc.text = svc
                lc.paragraphs[0].runs[0].bold      = True
                lc.paragraphs[0].runs[0].font.size = Pt(10)
                shd_cell(lc, 'FFEBEE')
                cell_borders(lc, 'C00000')
                vc.text = "📞  " + num
                vc.paragraphs[0].runs[0].font.size = Pt(11)
                vc.paragraphs[0].runs[0].bold      = True
                vc.paragraphs[0].runs[0]\
                    .font.color.rgb = RGBColor(
                    192, 0, 0
                )
                cell_borders(vc, 'C00000')

    # Internal contacts
    doc.add_paragraph()
    section_hdr(
        doc,
        "4B. INTERNAL PROJECT CONTACTS",
        fill="37474F", size=10
    )
    int_tbl       = doc.add_table(rows=1, cols=3)
    int_tbl.style = 'Table Grid'
    for cell, h in zip(
        int_tbl.rows[0].cells,
        ["Role", "Name", "Mobile Number"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '37474F')
        cell_borders(cell, '37474F')

    int_contacts = [
        ("Projects Director",
         meta.get('projects_director',''), ""),
        ("Project Manager",
         meta.get('project_manager',''), ""),
        ("HSE Manager / Officer",
         meta.get('emergency_coord',''), ""),
        ("Site Supervisor / Foreman",
         "", ""),
        ("First Aider (Primary)",
         meta.get('first_aider',''), ""),
        ("First Aider (Secondary)", "", ""),
        ("Site Security",  "", ""),
        ("Subcontractor HSE Rep", "", ""),
    ]
    for i, (role, name, mob) in enumerate(
        int_contacts
    ):
        row   = int_tbl.add_row()
        vals  = [role, name, mob]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v if v else "_______________"
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'F5F5F5')
            cell_borders(c, '37474F')

    doc.add_paragraph()

    # ── MUSTER POINTS & SITE LAYOUT ───────────
    section_hdr(
        doc,
        "5. MUSTER POINTS & EVACUATION ROUTES",
        fill="1F5864"
    )
    mus_tbl       = doc.add_table(rows=1, cols=4)
    mus_tbl.style = 'Table Grid'
    for cell, h in zip(
        mus_tbl.rows[0].cells,
        ["Muster Point",
         "Location Description",
         "Assembly Leader",
         "Capacity"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F5864')
        cell_borders(cell, '1F5864')

    muster_pts = erp_data.get('muster_points', [
        {"id": "MP-01",
         "location": "Main Site Entrance — "
                     "Near Security Hut",
         "leader": "Site Supervisor",
         "capacity": "All workers"},
        {"id": "MP-02",
         "location": "North corner of site — "
                     "Adjacent to welfare block",
         "leader": "HSE Officer",
         "capacity": "50 persons"},
    ])
    for i, mp in enumerate(muster_pts):
        row   = mus_tbl.add_row()
        vals  = [
            mp.get('id','MP-0' + str(i+1)),
            mp.get('location',''),
            mp.get('leader','Site Supervisor'),
            mp.get('capacity',''),
        ]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E0F2F1')
            cell_borders(c, '1F5864')

    doc.add_paragraph()

    # ── GENERAL EMERGENCY PROCEDURE ───────────
    section_hdr(
        doc,
        "6. GENERAL EMERGENCY RESPONSE "
        "PROCEDURE (ALL EMERGENCIES)",
        fill="8B0000"
    )

    gen_steps = [
        ("DISCOVER",
         "Discover or witness an emergency situation"),
        ("STOP WORK",
         "STOP all work immediately in the area"),
        ("RAISE ALARM",
         "Sound the alarm — shout EMERGENCY, "
         "use radio, call emergency number"),
        ("CALL",
         "Call emergency services: "
         + ("Ambulance 998 | Police 999 | "
            "Civil Defence 997"
            if emirate == "Abu Dhabi"
            else
            "Police 999 | Civil Defence 997 | "
            "Ambulance 998 | DM 800900")),
        ("NOTIFY",
         "Notify HSE Officer and Project Manager "
         "immediately"),
        ("EVACUATE",
         "Evacuate all personnel to nearest "
         "Muster Point — do NOT use elevators"),
        ("ACCOUNT",
         "Emergency Coordinator takes headcount "
         "at Muster Point — report missing persons"),
        ("FIRST AID",
         "Qualified First Aider provides immediate "
         "first aid — do NOT move seriously injured "
         "unless in danger"),
        ("SECURE",
         "Secure the area — prevent unauthorised "
         "re-entry until all-clear given"),
        ("COOPERATE",
         "Cooperate fully with emergency services "
         "on arrival — brief incident commander"),
        ("INVESTIGATE",
         "After emergency: preserve scene, "
         "document evidence, report to authorities"),
        ("REPORT",
         "Complete incident report — notify "
         + ("ADOSH within required timeframe "
            "per ADOSH-SF v4.0"
            if emirate == "Abu Dhabi"
            else
            "Dubai Municipality within 72 hours "
            "per DM Code Art. 2.4.5")),
    ]

    gen_tbl       = doc.add_table(rows=1, cols=3)
    gen_tbl.style = 'Table Grid'
    for cell, h in zip(
        gen_tbl.rows[0].cells,
        ["Step", "Action", "Detail"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '8B0000')
        cell_borders(cell, '8B0000')
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    step_colors = [
        'FF0000','C00000','B71C1C','8B0000',
        '7B1FA2','1565C0','1B5E20','E65100',
        '37474F','4A148C','1F3864','37474F',
    ]
    for i, (step, detail) in enumerate(gen_steps):
        row   = gen_tbl.add_row()
        cells = row.cells
        color = step_colors[
            i % len(step_colors)
        ]

        cells[0].width = Cm(1.5)
        cells[0].text  = str(i + 1)
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0]\
            .font.size = Pt(14)
        cells[0].paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cells[0], color)
        cell_borders(cells[0], '8B0000')
        cells[0].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        cells[1].width = Cm(3.5)
        cells[1].text  = step
        cells[1].paragraphs[0].runs[0].bold = True
        cells[1].paragraphs[0].runs[0]\
            .font.size = Pt(10)
        cells[1].paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cells[1], color)
        cell_borders(cells[1], '8B0000')
        cells[1].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        cells[2].width = Cm(12)
        cells[2].text  = detail
        cells[2].paragraphs[0].runs[0]\
            .font.size = Pt(9)
        if i % 2 == 0:
            shd_cell(cells[2], 'FFF5F5')
        cell_borders(cells[2], '8B0000')

    doc.add_paragraph()

    # ── SCENARIO-SPECIFIC PROCEDURES ──────────
    section_hdr(
        doc,
        "7. SCENARIO-SPECIFIC EMERGENCY "
        "PROCEDURES",
        fill="8B0000"
    )

    scenarios = erp_data.get('scenarios', [])
    for scn in scenarios:
        name  = scn.get('name','')
        color = scn.get('color', 'C00000')
        cop   = scn.get('cop', '')
        steps = scn.get('steps', [])

        # Scenario header
        scn_p = doc.add_paragraph()
        scn_p.paragraph_format.space_before = Pt(8)
        scn_p.paragraph_format.space_after  = Pt(2)
        shd_para(scn_p, color)
        scn_r = scn_p.add_run(
            "  🚨  " + name.upper()
        )
        scn_r.bold      = True
        scn_r.font.size = Pt(10)
        scn_r.font.color.rgb = RGBColor(
            255, 255, 255
        )

        # CoP reference
        if cop:
            ref_p = doc.add_paragraph()
            shd_para(ref_p, '37474F')
            ref_r = ref_p.add_run(
                "  📋 Reference: " + cop
            )
            ref_r.font.size  = Pt(8)
            ref_r.font.italic = True
            ref_r.font.color.rgb = RGBColor(
                255, 255, 200
            )

        # Steps table
        if steps:
            st_tbl       = doc.add_table(
                rows=1, cols=3
            )
            st_tbl.style = 'Table Grid'
            for cell, h in zip(
                st_tbl.rows[0].cells,
                ["Step", "Action",
                 "Responsible Person"]
            ):
                cell.text = h
                cell.paragraphs[0].runs[0]\
                    .bold = True
                cell.paragraphs[0].runs[0]\
                    .font.size = Pt(8)
                cell.paragraphs[0].runs[0]\
                    .font.color.rgb = RGBColor(
                    255, 255, 255
                )
                shd_cell(cell, color)
                cell_borders(cell, color)

            for i, step in enumerate(steps):
                row  = st_tbl.add_row()
                if isinstance(step, dict):
                    s_num  = str(i + 1)
                    s_act  = step.get('action','')
                    s_resp = step.get(
                        'responsible',
                        'HSE Officer'
                    )
                else:
                    s_num  = str(i + 1)
                    s_act  = str(step)
                    s_resp = 'HSE Officer'

                vals   = [s_num, s_act, s_resp]
                widths = [Cm(1.2), Cm(11.8),
                          Cm(4.0)]
                for j, (c, v, w) in enumerate(
                    zip(row.cells, vals, widths)
                ):
                    c.width = w
                    c.text  = v
                    c.paragraphs[0].runs[0]\
                        .font.size = Pt(9)
                    if j == 0:
                        c.paragraphs[0].runs[0]\
                            .bold = True
                        c.paragraphs[0].alignment\
                            = WD_ALIGN_PARAGRAPH.CENTER
                    if i % 2 == 0:
                        shd_cell(c, 'FFF5F5')
                    cell_borders(c, color)

        doc.add_paragraph()

    # ── FIRST AID ─────────────────────────────
    section_hdr(
        doc, "8. FIRST AID ARRANGEMENTS",
        fill="1B5E20"
    )
    fa_tbl       = doc.add_table(rows=1, cols=3)
    fa_tbl.style = 'Table Grid'
    for cell, h in zip(
        fa_tbl.rows[0].cells,
        ["Item", "Requirement", "Location / Detail"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1B5E20')
        cell_borders(cell, '1B5E20')

    fa_items = [
        ("First Aid Room",
         "Mandatory on site",
         "Site welfare block — clearly signed"),
        ("First Aid Kits",
         "Minimum 1 per 25 workers",
         "Site office, welfare block, "
         "each work area"),
        ("Stretcher / Spine Board",
         "1 minimum on site",
         "First Aid Room"),
        ("AED (Defibrillator)",
         "Recommended — mandatory >100 workers",
         "Site office / main welfare area"),
        ("Eye Wash Station",
         "Required where chemicals used",
         "Chemical storage area, lab, "
         "concrete areas"),
        ("Oxygen Kit",
         "Required for confined space works",
         "Near confined space entry point"),
        ("Qualified First Aiders",
         "Min 1 per 50 workers at all times",
         "Named persons with valid certificate"),
        ("First Aid Log",
         "All treatment must be recorded",
         "First Aid Room logbook"),
        ("Emergency Medical Evacuation",
         "Nearest hospital pre-identified",
         meta.get('nearest_hospital',
                  'Nearest hospital — '
                  'pre-identify and post route')),
    ]
    for i, (item, req, loc) in enumerate(fa_items):
        row  = fa_tbl.add_row()
        vals = [item, req, loc]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F5E9')
            cell_borders(c, '1B5E20')

    doc.add_paragraph()

    # ── FIRE PREVENTION ───────────────────────
    section_hdr(
        doc,
        "9. FIRE PREVENTION & PROTECTION",
        fill="FF6F00"
    )
    fire_items = erp_data.get('fire_prevention', [
        "Fire risk assessment conducted "
        "before works commence",
        "Fire extinguishers positioned at max "
        "15m intervals — inspected monthly",
        "Flammable/combustible materials stored "
        "in designated areas — minimum 10m from "
        "ignition sources",
        "Hot Work Permit required for all welding, "
        "cutting, and burning operations",
        "30-minute fire watch after all hot work",
        "No smoking except in designated areas",
        "LPG cylinders stored upright, chained, "
        "in ventilated cage — away from heat sources",
        "Electrical switchboards kept clear — "
        "2m exclusion zone maintained",
        "Waste and combustible materials removed "
        "daily from work areas",
        "Fire alarm system tested weekly",
        "All fire exits and routes kept clear "
        "at all times",
        "Fire marshal identified and trained "
        "for each work area",
    ])
    for item in fire_items:
        add_bullet(doc, str(item))

    doc.add_paragraph()

    # ── INCIDENT REPORTING ────────────────────
    section_hdr(
        doc,
        "10. INCIDENT REPORTING & NOTIFICATION "
        "REQUIREMENTS",
        fill="4A148C"
    )
    rep_tbl       = doc.add_table(rows=1, cols=4)
    rep_tbl.style = 'Table Grid'
    for cell, h in zip(
        rep_tbl.rows[0].cells,
        ["Incident Type",
         "Notify Internally",
         "Notify Authority",
         "Timeframe"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '4A148C')
        cell_borders(cell, '4A148C')

    if emirate == "Abu Dhabi":
        rep_rows = [
            ("Fatality",
             "PM, PD immediately",
             "ADOSH + Police 999",
             "IMMEDIATELY"),
            ("Serious Injury / Hospitalisation",
             "PM, PD, HSE Manager",
             "ADOSH within 24 hours",
             "Within 24 hours"),
            ("Dangerous Occurrence",
             "PM, HSE Officer",
             "ADOSH within 48 hours",
             "Within 48 hours"),
            ("Lost Time Injury (LTI)",
             "PM, HSE Officer",
             "ADOSH monthly report",
             "Monthly report"),
            ("Near Miss",
             "HSE Officer same day",
             "Internal only",
             "Same day"),
            ("Property Damage",
             "PM, Site Manager",
             "Internal — assess severity",
             "Same day"),
            ("Environmental Incident",
             "PM, HSE Officer",
             "ADOSH + EAD if significant",
             "Within 24 hours"),
            ("Fire",
             "All management immediately",
             "Civil Defence 997 immediately",
             "IMMEDIATELY"),
        ]
    else:
        rep_rows = [
            ("Fatality",
             "PM, PD immediately",
             "DM + Police 999 + DCD 997",
             "IMMEDIATELY"),
            ("Serious Injury",
             "PM, PD, HSE Manager",
             "DM within 24 hours",
             "Within 24 hours"),
            ("Dangerous Occurrence",
             "PM, HSE Officer",
             "DM within 72 hours",
             "Within 72 hours"),
            ("Lost Time Injury",
             "PM, HSE Officer",
             "DM monthly report",
             "Monthly"),
            ("Near Miss",
             "HSE Officer same day",
             "Internal only",
             "Same day"),
            ("Fire",
             "All management immediately",
             "Civil Defence 997 immediately",
             "IMMEDIATELY"),
            ("Environmental Incident",
             "PM, HSE Officer",
             "DM Environment Dept",
             "Within 24 hours"),
            ("Vehicle Accident",
             "PM, HSE Officer",
             "Dubai Police 999",
             "IMMEDIATELY"),
        ]

    urgent_colors = {
        "IMMEDIATELY": "FF0000",
        "Within 24 hours": "FFC000",
        "Within 48 hours": "FFFF00",
        "Within 72 hours": "FFFF00",
    }

    for i, (inc, internal, auth, time) in enumerate(
        rep_rows
    ):
        row   = rep_tbl.add_row()
        vals  = [inc, internal, auth, time]
        t_color = urgent_colors.get(time, '')
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if j == 3 and t_color:
                shd_cell(c, t_color)
                c.paragraphs[0].runs[0].bold = True
            elif i % 2 == 0:
                shd_cell(c, 'F3E5F5')
            cell_borders(c, '4A148C')

    doc.add_paragraph()

    # ── DRILLS ────────────────────────────────
    section_hdr(
        doc, "11. EMERGENCY DRILLS & EXERCISES",
        fill="1F3864"
    )
    drill_tbl       = doc.add_table(rows=1, cols=4)
    drill_tbl.style = 'Table Grid'
    for cell, h in zip(
        drill_tbl.rows[0].cells,
        ["Drill Type", "Frequency",
         "Organized By", "Record Kept"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F3864')
        cell_borders(cell, '1F3864')

    drills = [
        ("Full Site Evacuation Drill",
         "Quarterly (min)",
         "HSE Officer",
         "Drill Report + Register"),
        ("Fire Drill",
         "Quarterly",
         "HSE Officer + Fire Marshal",
         "Drill Report"),
        ("First Aid Scenario",
         "Bi-annual",
         "First Aider / HSE Officer",
         "Training Record"),
        ("Toolbox Talk — Emergency Procedures",
         "Monthly",
         "HSE Officer / Supervisor",
         "Attendance Register"),
        ("New Worker Emergency Briefing",
         "On induction (Day 1)",
         "HSE Officer",
         "Induction Record"),
        ("Confined Space Rescue Drill",
         "Before each CS entry",
         "CS Supervisor + Rescue Team",
         "Drill Record"),
        ("Communication Test",
         "Weekly",
         "HSE Officer",
         "Communication Log"),
        ("Emergency Equipment Check",
         "Monthly",
         "HSE Officer",
         "Inspection Checklist"),
    ]
    for i, (drill, freq, org, rec) in enumerate(
        drills
    ):
        row  = drill_tbl.add_row()
        vals = [drill, freq, org, rec]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F0FE')
            cell_borders(c, '1F3864')

    doc.add_paragraph()

    # ── EQUIPMENT ─────────────────────────────
    section_hdr(
        doc,
        "12. EMERGENCY EQUIPMENT & RESOURCES",
        fill="37474F"
    )
    eq_tbl       = doc.add_table(rows=1, cols=4)
    eq_tbl.style = 'Table Grid'
    for cell, h in zip(
        eq_tbl.rows[0].cells,
        ["Equipment", "Quantity",
         "Location on Site",
         "Inspection Frequency"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '37474F')
        cell_borders(cell, '37474F')

    eq_items = [
        ("Fire Extinguishers (ABC Dry Powder)",
         "As per fire risk assessment",
         "Throughout site — max 15m apart",
         "Monthly inspection"),
        ("Fire Extinguishers (CO2 — Electrical)",
         "Min 2 per DB / electrical area",
         "All DB panels and generator areas",
         "Monthly inspection"),
        ("Fire Blanket",
         "1 per welfare block / kitchen",
         "Kitchen / canteen area",
         "Quarterly"),
        ("First Aid Kits (Class C)",
         "1 per 25 workers",
         "Site office, work areas, vehicles",
         "Weekly check"),
        ("Stretcher / Spine Board",
         "Minimum 2 on site",
         "First Aid Room + Site Office",
         "Monthly"),
        ("AED (Defibrillator)",
         "1 minimum (>50 workers: 2)",
         "Site office / main welfare area",
         "Monthly battery check"),
        ("Eye Wash Station",
         "Per hazard area",
         "Chemical areas, concrete zones",
         "Weekly flush test"),
        ("Emergency Shower",
         "Where chemicals used",
         "Chemical storage / mixing area",
         "Weekly test"),
        ("Breathing Apparatus (SCBA)",
         "Min 2 sets for confined space",
         "Confined space entry point",
         "Before each use"),
        ("Spill Kit",
         "Min 1 per fuel/chemical area",
         "Fuel storage, chemical area",
         "Monthly check"),
        ("Emergency Lighting",
         "All escape routes and exits",
         "Throughout site",
         "Monthly test"),
        ("Assembly Point Signs",
         "Min 2 per muster point",
         "Muster points — clearly visible",
         "Inspect quarterly"),
        ("Emergency Contact Notice Boards",
         "Min 3 — site office, entry, welfare",
         "Main locations on site",
         "Update immediately on change"),
        ("Rescue Ropes / Harness Kit",
         "1 set per work-at-height team",
         "Height work areas",
         "Before each use"),
        ("Emergency Radio / Walkie-Talkie",
         "Min 4 on site",
         "HSE Officer, PM, Security, Supervisor",
         "Daily charge check"),
    ]
    for i, (eq, qty, loc, freq) in enumerate(
        eq_items
    ):
        row  = eq_tbl.add_row()
        vals = [eq, qty, loc, freq]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'F5F5F5')
            cell_borders(c, '37474F')

    doc.add_paragraph()

    # ── NEAREST HOSPITAL ──────────────────────
    section_hdr(
        doc,
        "13. NEAREST MEDICAL FACILITIES "
        "& ROUTE",
        fill="1B5E20"
    )
    hosp_tbl       = doc.add_table(rows=1, cols=4)
    hosp_tbl.style = 'Table Grid'
    for cell, h in zip(
        hosp_tbl.rows[0].cells,
        ["Facility", "Address / Location",
         "Phone", "Distance / Time"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1B5E20')
        cell_borders(cell, '1B5E20')

    hosp_rows = erp_data.get('hospitals', [
        {
            "name":     "Nearest Government Hospital",
            "address":  meta.get(
                'nearest_hospital',
                "To be confirmed by site team"
            ),
            "phone":    "Check locally",
            "distance": "TBC — verify before works",
        },
        {
            "name":     "Nearest Private Hospital",
            "address":  "To be confirmed",
            "phone":    "Check locally",
            "distance": "TBC",
        },
        {
            "name":     "Ambulance (Emergency)",
            "address":  "Call 998",
            "phone":    "998",
            "distance": "Response: ~15-20 min",
        },
    ])
    for i, hosp in enumerate(hosp_rows):
        row   = hosp_tbl.add_row()
        vals  = [
            hosp.get('name',''),
            hosp.get('address',''),
            hosp.get('phone',''),
            hosp.get('distance',''),
        ]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F5E9')
            cell_borders(c, '1B5E20')

    doc.add_paragraph()

    # ── REGULATORY REQUIREMENTS ───────────────
    section_hdr(
        doc,
        "14. REGULATORY & REPORTING "
        "REQUIREMENTS — " + emirate.upper(),
        fill="1F3864"
    )
    reg_items = erp_data.get('reg_requirements', [])
    for item in reg_items:
        add_bullet(doc, str(item))

    doc.add_paragraph()

    # ── ACKNOWLEDGEMENT ───────────────────────
    section_hdr(
        doc,
        "15. WORKER ACKNOWLEDGEMENT "
        "& TRAINING RECORD",
        fill="37474F"
    )
    add_body(
        doc,
        "All workers must be briefed on this "
        "Emergency Response Plan during site "
        "induction and sign below to confirm "
        "they have understood the emergency "
        "procedures, muster points, and their "
        "individual responsibilities."
    )
    doc.add_paragraph()

    ack_tbl       = doc.add_table(rows=1, cols=5)
    ack_tbl.style = 'Table Grid'
    for cell, h in zip(
        ack_tbl.rows[0].cells,
        ["No.", "Worker Name",
         "Company / Trade",
         "Date of Briefing",
         "Signature"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '37474F')
        cell_borders(cell, '37474F')

    for i in range(20):
        row  = ack_tbl.add_row()
        vals = [str(i+1), "", "", "", ""]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                shd_cell(c, 'F5F5F5')
            cell_borders(c, '37474F')
        row.cells[0].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    doc.add_paragraph()

    # ── APPROVAL SIGNATURES ───────────────────
    section_hdr(
        doc, "16. APPROVAL & SIGNATURES",
        fill="1F3864"
    )
    sig_tbl       = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    for i, role in enumerate([
        "Prepared By\n(HSE Officer / ERT Lead)",
        "Reviewed By\n(Project Manager)",
        "Approved By\n(Projects Director)"
    ]):
        c = sig_tbl.rows[0].cells[i]
        c.text = role
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, '1F3864')
        cell_borders(c, '1F3864')
    for i in range(3):
        c = sig_tbl.rows[1].cells[i]
        c.text = "\n\nName: _______________\n"
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c, '1F3864')
        c = sig_tbl.rows[2].cells[i]
        c.text = (
            "Designation: _______________\n\n"
            "Signature: _______________\n\n"
            "Date: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c, '1F3864')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════
# AI CONTENT GENERATORS
# ════════════════════════════════════════════════════════

def gen_erp_content(client, meta,
                    emirate, selected_scenarios):
    """
    Generate ERP content using small focused
    AI calls — avoids JSON overflow errors.
    """
    reg = (
        "ADOSH-SF v4.0, ADOSH CoP 44.0"
        if emirate == "Abu Dhabi"
        else "Dubai Municipality Code Chapter 12, "
             "Dubai Civil Defence requirements"
    )
    ctx = (
        "Project: " + meta.get('project','')
        + " | Location: " + meta.get('location','')
        + " | Company: " + meta.get('company','ENGC')
        + " | Emirate: " + emirate
    )

    data = {}

    # Purpose
    raw = ask(
        client,
        "Write 2-3 sentences for the PURPOSE "
        "section of an Emergency Response Plan "
        "for a construction project in "
        + emirate + ", UAE. Context: " + ctx
        + ". Reference: " + reg
        + ". Plain text only, no JSON."
    )
    data['purpose'] = raw[:500].strip()

    # ERT Personnel
    raw = ask(
        client,
        "Return ONLY a JSON array of 6 Emergency "
        "Response Team members for a construction "
        "site in " + emirate + ". Context: " + ctx
        + ".\n"
        'Format: [{"role":"Emergency Coordinator",'
        '"name":"' + meta.get('emergency_coord','TBD')
        + '","contact":"Radio Ch.1",'
        '"responsibility":"Lead all emergency response",'
        '"backup":"Deputy HSE Officer"}]\n'
        "Include: Emergency Coordinator, "
        "Deputy Coordinator, First Aider, "
        "Fire Marshal, Evacuation Warden, "
        "Security. "
        "Return ONLY the JSON array."
    )
    data['ert_personnel'] = safe_list(raw, [
        {"role":           "Emergency Coordinator",
         "name":           meta.get(
             'emergency_coord', 'HSE Officer'
         ),
         "contact":        "Radio Ch.1",
         "responsibility": "Lead all emergency "
                           "response",
         "backup":         "Deputy"},
        {"role":           "First Aider",
         "name":           meta.get(
             'first_aider', 'TBD'
         ),
         "contact":        "Radio Ch.2",
         "responsibility": "Provide first aid",
         "backup":         "Second First Aider"},
        {"role":           "Fire Marshal",
         "name":           "TBD",
         "contact":        "Radio Ch.3",
         "responsibility": "Control fire response "
                           "and evacuation",
         "backup":         "Assistant Fire Marshal"},
    ])

    # Scenario procedures
    data['scenarios'] = []
    for scn_name in selected_scenarios:
        scn_info = SCENARIOS.get(scn_name, {})
        color    = scn_info.get('color', 'C00000')
        cop      = scn_info.get('cop', '')

        raw = ask(
            client,
            "Return ONLY a JSON array of 6 "
            "emergency response steps for: "
            + scn_name + " on a construction site "
            "in " + emirate + ".\n"
            'Format: [{"action":"What to do",'
            '"responsible":"Who does it"}]\n'
            "Include: stop work, call emergency, "
            "specific response actions for "
            + scn_name + ". "
            "Use " + emirate + " emergency numbers "
            + str(
                list(
                    EMERGENCY_NUMBERS.get(
                        emirate, {}
                    ).items()
                )[:3]
            )
            + ". Return ONLY the JSON array."
        )
        steps = safe_list(raw, [
            {"action":      "STOP WORK — raise alarm",
             "responsible": "All workers"},
            {"action":      "Call emergency services",
             "responsible": "Emergency Coordinator"},
            {"action":      "Evacuate to Muster Point",
             "responsible": "Evacuation Warden"},
            {"action":      "Account for all workers",
             "responsible": "Emergency Coordinator"},
            {"action":      "Notify management",
             "responsible": "HSE Officer"},
            {"action":      "Complete incident report",
             "responsible": "HSE Officer"},
        ])

        data['scenarios'].append({
            "name":  scn_name,
            "color": color,
            "cop":   cop,
            "steps": steps,
        })

    # Fire prevention
    raw = ask(
        client,
        "Return ONLY a JSON array of 8 fire "
        "prevention measures (as plain strings) "
        "for a construction site in "
        + emirate + ". Reference: " + reg + ".\n"
        'Example: ["Hot work permit required for '
        'all welding operations"]\n'
        "Return ONLY the JSON array."
    )
    data['fire_prevention'] = safe_list(raw, [
        "Hot work permit required for welding",
        "Fire extinguishers checked monthly",
        "Flammable materials stored away from "
        "ignition sources",
        "No smoking except in designated areas",
        "Fire exits kept clear at all times",
        "Weekly fire alarm test conducted",
        "Fire marshal identified and trained",
        "Combustible waste removed daily",
    ])

    # Regulatory requirements
    if emirate == "Abu Dhabi":
        data['reg_requirements'] = [
            "ADOSH-SF Version 4.0 (July 2024) — "
            "Emergency preparedness requirements",
            "ADOSH CoP 44.0 — Emergency Management "
            "on Construction Sites",
            "ADOSH CoP 53.1 — OSH Construction "
            "Management Plan (ERP mandatory section)",
            "UAE Civil Defence Requirements — "
            "fire safety and evacuation",
            "MOHRE Resolution 44/2022 — "
            "Heat emergency protocols "
            "(15 Jun–15 Sep)",
            "ADOSH CoP 3.0 — First Aid requirements",
            "Fatalities must be reported to ADOSH "
            "IMMEDIATELY",
            "Serious injuries: ADOSH within 24 hours",
            "Dangerous occurrences: "
            "ADOSH within 48 hours",
            "Monthly HSE statistics report to ADOSH",
            "Emergency drill records must be "
            "maintained for minimum 3 years",
            "ERP must be reviewed quarterly and "
            "after every emergency or near miss",
        ]
    else:
        data['reg_requirements'] = [
            "Dubai Municipality Code of Construction "
            "Safety Practice — Chapter 12",
            "Dubai Civil Defence (DCD) — "
            "Fire safety and emergency requirements",
            "UAE Federal Decree-Law 33/2021 — "
            "Worker safety obligations",
            "All accidents/incidents reported to "
            "DM within 72 hours",
            "Fatalities: Dubai Police 999 + "
            "DM IMMEDIATELY",
            "Fire emergencies: Civil Defence 997",
            "Safety meetings every 15 days "
            "(DM Code Art. 2.2.8)",
            "Emergency drill records maintained "
            "for minimum 3 years",
            "ERP must be approved by Engineer "
            "before works commence",
            "DM Environment Dept for environmental "
            "incidents: 800 900",
            "DEWA emergency for utility incidents: "
            "991",
        ]

    # Hospital info (AI-generated, location-based)
    raw = ask(
        client,
        "Name 2 nearest hospitals to a construction "
        "site in " + meta.get(
            'location', emirate
        ) + ", " + emirate + ", UAE. "
        "Return ONLY a JSON array:\n"
        '[{"name":"Hospital name",'
        '"address":"General area/district",'
        '"phone":"Hospital number or 998",'
        '"distance":"Approx distance/time"}]\n'
        "Return ONLY the JSON array."
    )
    data['hospitals'] = safe_list(raw, [
        {
            "name":     "Nearest Government Hospital",
            "address":  "Confirm location with "
                        "local authorities",
            "phone":    "998 (Ambulance)",
            "distance": "TBC — verify before works",
        },
        {
            "name":     "Ambulance Emergency",
            "address":  "Call immediately",
            "phone":    "998",
            "distance": "Response ~15-20 min",
        },
    ])

    return data


# ════════════════════════════════════════════════════════
# MAIN UI — THREE TABS
# ════════════════════════════════════════════════════════

tab_gen, tab_quick, tab_upload = st.tabs([
    "🚨 Generate Full ERP",
    "⚡ Quick Emergency Cards",
    "📤 Upload Your ERP Format",
])


# ════════════════════════════════════════════════════════
# TAB 1 — FULL ERP GENERATOR
# ════════════════════════════════════════════════════════

with tab_gen:
    st.markdown(
        "### 🚨 Generate Complete Emergency "
        "Response Plan"
    )
    st.info(
        "Fill in your project details and select "
        "the emergency scenarios relevant to your "
        "site. The AI generates a complete, "
        + emirate + "-specific ERP with all "
        "response procedures, correct emergency "
        "numbers, regulatory references, "
        "and worker acknowledgement register."
    )

    # Show emirate numbers
    nums = EMERGENCY_NUMBERS.get(emirate, {})
    st.markdown(
        "#### 🚨 " + emirate + " Emergency Numbers"
    )
    num_cols = st.columns(
        min(5, len(nums))
    )
    for i, (svc, num) in enumerate(nums.items()):
        num_cols[i % 5].metric(svc, num)

    st.divider()
    st.subheader("📋 Project Details")

    c1, c2 = st.columns(2)
    with c1:
        g_project  = st.text_input(
            "Project Name *",
            placeholder="e.g. Bloom Living Almeria",
            key="erp_proj"
        )
        g_location = st.text_input(
            "Site Location / Address *",
            placeholder=(
                "e.g. Plot C13, Zayed City, "
                "Abu Dhabi"
            ),
            key="erp_loc"
        )
        g_company  = st.text_input(
            "Main Contractor",
            value="ENGC",
            key="erp_co"
        )
        g_client   = st.text_input(
            "Client",
            placeholder="e.g. Bloom District Properties",
            key="erp_client"
        )
    with c2:
        g_doc_no   = st.text_input(
            "Document Number",
            placeholder="e.g. ENGC-ERP-001",
            key="erp_docno"
        )
        g_rev      = st.selectbox(
            "Revision",
            ["Rev 00","Rev 01","Rev 02","Rev 03"],
            key="erp_rev"
        )
        g_date     = st.date_input(
            "Issue Date",
            value=date.today(),
            key="erp_date"
        )
        g_workers  = st.number_input(
            "Peak Workforce (number)",
            min_value=5,
            max_value=5000,
            value=150,
            key="erp_workers"
        )

    st.divider()
    st.subheader("👥 Key Personnel")
    p1, p2 = st.columns(2)
    with p1:
        g_prep   = st.text_input(
            "Prepared By (HSE Officer)",
            placeholder="e.g. Aftab Qamar",
            key="erp_prep"
        )
        g_coord  = st.text_input(
            "Emergency Coordinator (Primary)",
            placeholder="e.g. Aftab Qamar — "
                        "Senior HSSE Engineer",
            key="erp_coord"
        )
        g_deputy = st.text_input(
            "Deputy Emergency Coordinator",
            placeholder="e.g. Usman Qaisar — "
                        "Sr. Safety Officer",
            key="erp_deputy"
        )
    with p2:
        g_pm       = st.text_input(
            "Project Manager",
            placeholder="e.g. Ahmad Abdelrahman",
            key="erp_pm"
        )
        g_pd       = st.text_input(
            "Projects Director",
            placeholder="e.g. Rami Kamal Yassin",
            key="erp_pd"
        )
        g_fa       = st.text_input(
            "First Aider (Qualified)",
            placeholder="e.g. Gurpal Singh — "
                        "HSE Engineer",
            key="erp_fa"
        )

    st.divider()
    st.subheader("🏥 Site Information")
    s1, s2 = st.columns(2)
    with s1:
        g_hospital = st.text_input(
            "Nearest Hospital Name & Location",
            placeholder=(
                "e.g. Sheikh Khalifa Medical City, "
                "Karama St, Abu Dhabi"
            ),
            key="erp_hosp"
        )
        g_muster1  = st.text_input(
            "Muster Point 1 Location",
            placeholder="e.g. Main site entrance — "
                        "near security hut",
            key="erp_mp1"
        )
    with s2:
        g_approved = st.text_input(
            "Approved By (Projects Director)",
            placeholder="e.g. Rami Kamal Yassin",
            key="erp_approve"
        )
        g_muster2  = st.text_input(
            "Muster Point 2 Location (optional)",
            placeholder="e.g. North corner of site "
                        "— near welfare block",
            key="erp_mp2"
        )

    st.divider()
    st.subheader(
        "⚠️ Select Emergency Scenarios for This Site"
    )
    st.caption(
        "Select ALL scenarios applicable to your "
        "project. The AI will generate specific "
        "response procedures for each one."
    )

    # Show in columns
    scn_names   = list(SCENARIOS.keys())
    selected    = []
    cols_scn    = st.columns(2)
    for i, scn in enumerate(scn_names):
        col = cols_scn[i % 2]
        if col.checkbox(
            scn,
            value=(i < 6),
            key="scn_" + str(i)
        ):
            selected.append(scn)

    if st.button(
        "🚨 Generate Complete ERP (.docx)",
        type="primary",
        use_container_width=True,
        key="erp_gen_btn"
    ):
        if not g_project or not g_location:
            st.error(
                "Please enter the project name "
                "and site location."
            )
        elif not selected:
            st.error(
                "Please select at least one "
                "emergency scenario."
            )
        else:
            prog = st.progress(
                0,
                text="Starting ERP generation..."
            )
            try:
                client = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                meta = {
                    "project":          g_project,
                    "location":         g_location,
                    "company":          g_company,
                    "client":           g_client,
                    "doc_no":           g_doc_no,
                    "revision":         g_rev,
                    "erp_date":         g_date,
                    "prepared_by":      g_prep,
                    "approved_by":      g_approved,
                    "emergency_coord":  g_coord,
                    "deputy_coord":     g_deputy,
                    "project_manager":  g_pm,
                    "projects_director": g_pd,
                    "first_aider":      g_fa,
                    "nearest_hospital": g_hospital,
                }

                muster_pts = []
                if g_muster1:
                    muster_pts.append({
                        "id":       "MP-01",
                        "location": g_muster1,
                        "leader":   g_coord or
                                    "HSE Officer",
                        "capacity": "All workers",
                    })
                if g_muster2:
                    muster_pts.append({
                        "id":       "MP-02",
                        "location": g_muster2,
                        "leader":   g_deputy or
                                    "Deputy HSE",
                        "capacity": str(
                            g_workers // 2
                        ) + " workers",
                    })

                total = len(selected) + 5
                step  = [0]

                def upd(msg):
                    step[0] += 1
                    prog.progress(
                        min(
                            90,
                            int(step[0] / total * 88)
                        ),
                        text=msg
                    )

                upd("Generating purpose & scope...")
                upd("Building ERT personnel...")

                erp_data = gen_erp_content(
                    client, meta, emirate, selected
                )

                if muster_pts:
                    erp_data['muster_points'] = (
                        muster_pts
                    )

                for scn in selected:
                    upd(
                        "Generating: "
                        + scn[:40] + "..."
                    )

                upd("Building Word document...")

                docx_buf = build_erp_docx(
                    erp_data, meta,
                    emirate, selected
                )

                prog.progress(
                    100,
                    text="✅ ERP complete!"
                )

                st.success(
                    "✅ Emergency Response Plan "
                    "generated — "
                    + str(len(selected))
                    + " scenarios | 16 sections | "
                    "Worker acknowledgement register!"
                )

                proj_fn = g_project.replace(' ','_')
                em_fn   = emirate.replace(' ','_')
                fname   = (
                    "ERP_"
                    + proj_fn + "_"
                    + em_fn + "_"
                    + str(g_date) + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download Emergency "
                        "Response Plan (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="erp_dl"
                )

                st.warning(
                    "📋 **MANDATORY NEXT STEPS:**\n\n"
                    "1. ✅ Brief ALL workers on "
                    "this ERP during site induction\n\n"
                    "2. ✅ Post Emergency Numbers "
                    "at site office, entrance, "
                    "and welfare block\n\n"
                    "3. ✅ Conduct evacuation drill "
                    "within first week of site "
                    "mobilization\n\n"
                    "4. ✅ Submit to Consultant "
                    "(" + (
                        "Dar Al-Handasah"
                        if "Bloom" in g_project
                        else "your consultant"
                    ) + ") for approval\n\n"
                    "5. ✅ Review quarterly or "
                    "after any incident\n\n"
                    "6. ✅ Keep signed "
                    "acknowledgements on file\n\n"
                    "⚠️ *This ERP must be verified "
                    "by a competent HSE professional "
                    "before use on site.*"
                )

            except Exception as e:
                prog.progress(0)
                st.error("Error: " + str(e))


# ════════════════════════════════════════════════════════
# TAB 2 — QUICK EMERGENCY CARDS
# ════════════════════════════════════════════════════════

with tab_quick:
    st.markdown(
        "### ⚡ Quick Emergency Reference Cards"
    )
    st.info(
        "Generate single-page emergency reference "
        "cards for posting around the site — "
        "notice boards, site office, welfare block, "
        "and work areas. "
        "Simple, clear, and bilingual-friendly."
    )

    nums_q = EMERGENCY_NUMBERS.get(emirate, {})

    q1, q2 = st.columns(2)
    with q1:
        q_project  = st.text_input(
            "Project Name",
            placeholder="e.g. Bloom Living Almeria",
            key="q_proj"
        )
        q_coord    = st.text_input(
            "Emergency Coordinator Name",
            placeholder="e.g. Aftab Qamar",
            key="q_coord"
        )
        q_coord_no = st.text_input(
            "Coordinator Mobile",
            placeholder="e.g. +971 52 267 1122",
            key="q_coordno"
        )
    with q2:
        q_company  = st.text_input(
            "Company",
            value="ENGC",
            key="q_co"
        )
        q_muster   = st.text_input(
            "Muster Point Location",
            placeholder="e.g. Main entrance — "
                        "near security hut",
            key="q_muster"
        )
        q_hospital = st.text_input(
            "Nearest Hospital",
            placeholder="e.g. Sheikh Khalifa "
                        "Medical City",
            key="q_hosp"
        )

    if st.button(
        "⚡ Generate Emergency Card (.docx)",
        type="primary",
        use_container_width=True,
        key="q_gen_btn"
    ):
        if not q_project:
            st.error(
                "Please enter a project name."
            )
        else:
            with st.spinner(
                "Building emergency card..."
            ):
                doc = Document()
                sec               = doc.sections[0]
                sec.page_width    = Cm(21.0)
                sec.page_height   = Cm(29.7)
                sec.left_margin   = Cm(1.5)
                sec.right_margin  = Cm(1.5)
                sec.top_margin    = Cm(1.5)
                sec.bottom_margin = Cm(1.5)

                # Big red header
                hdr           = doc.add_paragraph()
                hdr.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
                shd_para(hdr, '8B0000')
                hr = hdr.add_run(
                    "\n  🚨  IN AN EMERGENCY  🚨"
                    "\n  FOLLOW THESE STEPS  \n"
                )
                hr.bold           = True
                hr.font.size      = Pt(24)
                hr.font.color.rgb = RGBColor(
                    255, 255, 255
                )

                # STOP-CALL-HELP steps
                steps_q = [
                    ("1. STOP", "C00000",
                     "STOP ALL WORK IMMEDIATELY\n"
                     "Move away from danger zone"),
                    ("2. RAISE ALARM", "B71C1C",
                     "SHOUT 'EMERGENCY'\n"
                     "Use radio — Ch. 1\n"
                     "Press alarm button if available"),
                    ("3. CALL", "8B0000",
                     "CALL EMERGENCY SERVICES NOW"),
                    ("4. EVACUATE", "1565C0",
                     "WALK — DO NOT RUN\n"
                     "Go to Muster Point:\n"
                     + q_muster),
                    ("5. REPORT", "1B5E20",
                     "REPORT TO:\n"
                     + q_coord + "\n"
                     + q_coord_no),
                ]

                for step_title, col, detail in (
                    steps_q
                ):
                    p   = doc.add_paragraph()
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                    )
                    shd_para(p, col)
                    r = p.add_run(
                        "\n  " + step_title
                        + "  \n  " + detail
                        + "  \n"
                    )
                    r.bold      = True
                    r.font.size = Pt(16)
                    r.font.color.rgb = RGBColor(
                        255, 255, 255
                    )

                # Emergency numbers
                doc.add_paragraph()
                num_hdr           = doc.add_paragraph()
                num_hdr.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
                shd_para(num_hdr, '1F3864')
                nr = num_hdr.add_run(
                    "\n  📞  EMERGENCY NUMBERS  \n"
                )
                nr.bold           = True
                nr.font.size      = Pt(18)
                nr.font.color.rgb = RGBColor(
                    255, 255, 0
                )

                num_tbl       = doc.add_table(
                    rows=0, cols=2
                )
                num_tbl.style = 'Table Grid'
                for svc, num in nums_q.items():
                    row   = num_tbl.add_row()
                    cells = row.cells
                    cells[0].text = svc
                    cells[0].paragraphs[0]\
                        .runs[0].bold      = True
                    cells[0].paragraphs[0]\
                        .runs[0].font.size = Pt(13)
                    shd_cell(cells[0], '1A1A1A')
                    cells[0].paragraphs[0]\
                        .runs[0].font.color.rgb = (
                        RGBColor(255, 255, 255)
                    )
                    cell_borders(cells[0], 'C00000')
                    cells[1].text = "📞  " + num
                    cells[1].paragraphs[0]\
                        .runs[0].bold      = True
                    cells[1].paragraphs[0]\
                        .runs[0].font.size = Pt(16)
                    cells[1].paragraphs[0]\
                        .runs[0].font.color.rgb = (
                        RGBColor(255, 50, 50)
                    )
                    cell_borders(cells[1], 'C00000')

                # Internal contact
                doc.add_paragraph()
                ic           = doc.add_paragraph()
                ic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shd_para(ic, '1B5E20')
                ir = ic.add_run(
                    "\n  👷 EMERGENCY COORDINATOR\n"
                    "  " + q_coord + "  \n"
                    "  📱 " + q_coord_no + "  \n\n"
                    "  🏥 NEAREST HOSPITAL: "
                    + q_hospital + "  \n"
                    "  📞 Call Ambulance: 998  \n"
                )
                ir.bold           = True
                ir.font.size      = Pt(14)
                ir.font.color.rgb = RGBColor(
                    255, 255, 255
                )

                # Footer
                ft           = doc.add_paragraph()
                ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shd_para(ft, '1F3864')
                fr = ft.add_run(
                    "\n  " + q_company
                    + " | " + q_project
                    + " | " + emirate
                    + ", UAE  \n"
                    "  🚨 MUSTER POINT: "
                    + q_muster + "  \n"
                )
                fr.bold           = True
                fr.font.size      = Pt(11)
                fr.font.color.rgb = RGBColor(
                    255, 255, 255
                )

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)

                proj_fn = q_project.replace(' ','_')
                fname   = (
                    "Emergency_Card_"
                    + proj_fn + "_"
                    + str(date.today()) + ".docx"
                )

                st.success(
                    "✅ Emergency Card generated! "
                    "Print A3 size and post at: "
                    "Site entrance, site office, "
                    "welfare block, and each "
                    "work area."
                )

                st.download_button(
                    label=(
                        "⬇️ Download Emergency "
                        "Reference Card (.docx)"
                    ),
                    data=buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="q_dl"
                )

                st.info(
                    "🖨️ **Print Instructions:**\n\n"
                    "Print at **A3 size** for "
                    "maximum visibility.\n\n"
                    "Laminate and post at:\n"
                    "• Site main entrance\n"
                    "• Site office door\n"
                    "• Welfare block / rest area\n"
                    "• Each major work area\n"
                    "• Inside all site vehicles"
                )


# ════════════════════════════════════════════════════════
# TAB 3 — UPLOAD YOUR ERP FORMAT
# ════════════════════════════════════════════════════════

with tab_upload:
    st.markdown(
        "### 📤 Upload Your Own ERP Format"
    )
    st.info(
        "Upload your company's existing ERP "
        "template (Word or PDF). The AI reads "
        "your EXACT format — all sections, "
        "tables, headings — and fills it with "
        "project-specific content and correct "
        + emirate + " emergency numbers and "
        "regulatory references."
    )

    up_file = st.file_uploader(
        "📎 Upload ERP Template or Sample",
        type=["docx", "pdf"],
        help=(
            "Upload a blank ERP template or "
            "a previously completed ERP. "
            "Word (.docx) gives best results."
        ),
        key="erp_up_file"
    )

    if up_file:
        st.success(
            "✅ Uploaded: **"
            + up_file.name + "** ("
            + "{:,}".format(up_file.size)
            + " bytes)"
        )

        st.divider()
        st.subheader("📋 Project Details")

        u1, u2 = st.columns(2)
        with u1:
            u_project  = st.text_input(
                "Project Name",
                placeholder="e.g. Bloom Living Almeria",
                key="u_erp_proj"
            )
            u_location = st.text_input(
                "Site Location",
                placeholder="e.g. Zayed City, "
                            "Abu Dhabi",
                key="u_erp_loc"
            )
            u_company  = st.text_input(
                "Company",
                value="ENGC",
                key="u_erp_co"
            )
            u_coord    = st.text_input(
                "Emergency Coordinator",
                placeholder="e.g. Aftab Qamar",
                key="u_erp_coord"
            )
        with u2:
            u_workers  = st.number_input(
                "Peak Workforce",
                min_value=5,
                max_value=5000,
                value=150,
                key="u_erp_workers"
            )
            u_pm       = st.text_input(
                "Project Manager",
                placeholder="e.g. Ahmad Abdelrahman",
                key="u_erp_pm"
            )
            u_hospital = st.text_input(
                "Nearest Hospital",
                placeholder="e.g. Sheikh Khalifa "
                            "Medical City",
                key="u_erp_hosp"
            )
            u_date     = st.date_input(
                "Issue Date",
                value=date.today(),
                key="u_erp_date"
            )

        if st.button(
            "⚡ Fill My ERP Format",
            type="primary",
            use_container_width=True,
            key="u_erp_btn"
        ):
            if not u_project:
                st.error(
                    "Please enter a project name."
                )
            else:
                with st.spinner(
                    "Reading format and filling "
                    "ERP content..."
                ):
                    try:
                        client = anthropic.Anthropic(
                            api_key=os.getenv(
                                "ANTHROPIC_API_KEY"
                            )
                        )

                        nums_str = "\n".join([
                            svc + ": " + num
                            for svc, num
                            in EMERGENCY_NUMBERS.get(
                                emirate, {}
                            ).items()
                        ])

                        reg = (
                            "ADOSH-SF v4.0, "
                            "ADOSH CoP 44.0"
                            if emirate == "Abu Dhabi"
                            else
                            "Dubai Municipality Code "
                            "Chapter 12, "
                            "Dubai Civil Defence"
                        )

                        instr = (
                            "This is an Emergency "
                            "Response Plan (ERP) "
                            "template for "
                            + emirate + ", UAE.\n\n"
                            "Fill it for:\n"
                            "Project: " + u_project + "\n"
                            "Location: " + u_location + "\n"
                            "Company: " + u_company + "\n"
                            "Emergency Coordinator: "
                            + u_coord + "\n"
                            "Project Manager: " + u_pm + "\n"
                            "Peak Workers: "
                            + str(u_workers) + "\n"
                            "Nearest Hospital: "
                            + u_hospital + "\n"
                            "Date: " + str(u_date) + "\n"
                            "Regulation: " + reg + "\n\n"
                            "EMERGENCY NUMBERS FOR "
                            + emirate + ":\n"
                            + nums_str + "\n\n"
                            "INSTRUCTIONS:\n"
                            "1. Keep EXACT same "
                            "headings and structure\n"
                            "2. Fill ALL emergency "
                            "numbers with the correct "
                            + emirate + " numbers above\n"
                            "3. Fill all procedures "
                            "with specific, practical "
                            "steps\n"
                            "4. Reference " + reg + "\n"
                            "5. Make it ready for "
                            "immediate site use\n"
                            "6. Include scenario-specific "
                            "procedures for fire, "
                            "medical, collapse, "
                            "and electrical emergencies"
                        )

                        file_bytes = up_file.read()

                        if up_file.name\
                                .lower()\
                                .endswith('.pdf'):
                            b64 = (
                                base64.standard_b64encode(
                                    file_bytes
                                ).decode()
                            )
                            response = (
                                client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "document",
                                            "source": {
                                                "type": "base64",
                                                "media_type": "application/pdf",
                                                "data": b64,
                                            }
                                        },
                                        {
                                            "type": "text",
                                            "text": instr
                                        }
                                    ]
                                }]
                            ))
                        else:
                            doc_tmp = Document(
                                BytesIO(file_bytes)
                            )
                            extracted = "\n".join([
                                p.text
                                for p in doc_tmp.paragraphs
                                if p.text.strip()
                            ])
                            tbl_txt = ""
                            for tbl in doc_tmp.tables:
                                for row in tbl.rows:
                                    rt = " | ".join([
                                        c.text.strip()
                                        for c in row.cells
                                        if c.text.strip()
                                    ])
                                    if rt:
                                        tbl_txt += (
                                            rt + "\n"
                                        )
                            full = (
                                extracted[:2500]
                                + (
                                    "\n\nTABLES:\n"
                                    + tbl_txt[:500]
                                    if tbl_txt else ""
                                )
                            )
                            response = (
                                client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": (
                                        "ERP TEMPLATE:\n\n"
                                        "---\n" + full
                                        + "\n---\n\n"
                                        + instr
                                    )
                                }]
                            ))

                        filled = (
                            response.content[0].text
                        )

                        st.success(
                            "✅ ERP filled for: "
                            + u_project + "!"
                        )

                        with st.expander(
                            "📄 Preview Generated ERP",
                            expanded=True
                        ):
                            st.markdown(filled)

                        proj_fn = u_project.replace(
                            ' ', '_'
                        )
                        fname   = (
                            "ERP_"
                            + proj_fn + "_"
                            + str(u_date)
                            + "_filled.txt"
                        )

                        st.download_button(
                            label=(
                                "⬇️ Download Filled "
                                "ERP (.txt)"
                            ),
                            data=filled,
                            file_name=fname,
                            mime="text/plain",
                            key="u_erp_dl"
                        )

                        st.info(
                            "💡 Copy content above "
                            "into your Word template "
                            "for the final "
                            "formatted ERP."
                        )

                    except Exception as e:
                        st.error(
                            "Error: " + str(e)
                        )
    else:
        ca, cb = st.columns(2)
        with ca:
            st.markdown("""
**✅ Accepted Formats:**
- Word Document (.docx) — Best
- PDF (.pdf)

**📋 What to Upload:**
- Your company ERP template
- Consultant-approved ERP format
- Previous project ERP to update
- ADOSH/DM standard ERP format
""")
        with cb:
            st.markdown("""
**🤖 What the AI Does:**
- Fills correct emergency numbers
  for """ + emirate + """
- Adds scenario-specific procedures
- References correct regulations
- Fills all contact details
- Keeps your EXACT format
""")

        st.divider()
        st.info(
            "🚨 **Don't have an ERP template?**\n\n"
            "Use the **Generate Full ERP** tab "
            "to create a complete, professional "
            "Emergency Response Plan from scratch — "
            "with 16 sections, all "
            + emirate + " emergency numbers, "
            "scenario procedures, and worker "
            "acknowledgement register."
        )