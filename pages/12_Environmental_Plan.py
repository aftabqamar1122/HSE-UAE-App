import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import json
import re

load_dotenv()

st.set_page_config(
    page_title="Environmental Management Plan",
    page_icon="🌿"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning(
        "Please go to Home and select "
        "your Emirate first."
    )
    st.stop()

flag = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title(
    "🌿 Environmental Management Plan — "
    + flag + " " + emirate
)

if emirate == "Abu Dhabi":
    st.markdown(
        "**Reference: ADOSH CoP 54.0 — Waste "
        "Management | EAD Environmental Regulations "
        "| Abu Dhabi Environmental Quality Standards "
        "| ISO 14001:2015 | UAE Federal Law "
        "No. 24 of 1999 (Environment Protection)**"
    )
else:
    st.markdown(
        "**Reference: Dubai Municipality "
        "Environmental Protection Regulations | "
        "DM Green Building Regulations | "
        "Trakhees/EPDA Requirements | "
        "UAE Federal Law No. 24 of 1999 | "
        "ISO 14001:2015**"
    )

st.info(
    "🌿 **Legal Requirement:** All construction "
    "projects in the UAE must comply with "
    "environmental protection legislation under "
    "UAE Federal Law No. 24 of 1999 and emirate-"
    "specific regulations. This Environmental "
    "Management Plan (EMP) identifies "
    "environmental aspects, impacts, and controls "
    "in accordance with ISO 14001:2015 and local "
    "regulatory requirements.\n\n"
    "⚠️ *This EMP must be reviewed by a "
    "qualified Environmental professional and "
    "approved by the relevant authority before "
    "works commence.*"
)

st.divider()


# ════════════════════════════════════════════════════════
# REGULATORY FRAMEWORK
# ════════════════════════════════════════════════════════

REGULATIONS = {
    "Abu Dhabi": {
        "Primary Law":
            "UAE Federal Law No. 24 of 1999 "
            "— Environmental Protection",
        "Local Authority":
            "Environment Agency Abu Dhabi (EAD)",
        "Waste Regulation":
            "ADOSH CoP 54.0 — Waste Management",
        "Air Quality":
            "EAD Air Quality Standards "
            "(Abu Dhabi Quality System)",
        "Water/Dewatering":
            "EAD Water Quality Standards | "
            "AD Sewerage Co. Discharge Limits",
        "Noise":
            "EAD Noise Standards | "
            "Municipality Noise Bylaws",
        "Hazardous Waste":
            "EAD Hazardous Waste Regulations | "
            "Tadweer (Centre of Waste Management)",
        "Environmental Permit":
            "EAD Environmental Permit required "
            "for significant works",
        "Building":
            "Abu Dhabi Department of Urban "
            "Planning & Municipalities (DPM)",
        "ISO Standard":
            "ISO 14001:2015 EMS",
    },
    "Dubai": {
        "Primary Law":
            "UAE Federal Law No. 24 of 1999 "
            "— Environmental Protection",
        "Local Authority":
            "Dubai Municipality — "
            "Environmental Mgmt Department",
        "Waste Regulation":
            "DM Environmental Protection & Safety "
            "Section — Waste Management Regs",
        "Air Quality":
            "Dubai Municipality Air Quality "
            "Standards | UAE Federal Standard",
        "Water/Dewatering":
            "Dubai Municipality Discharge to "
            "Drainage Regulations",
        "Noise":
            "Dubai Municipality Noise "
            "Limit Regulations",
        "Hazardous Waste":
            "DM Hazardous Waste Regulations | "
            "Bee'ah (Waste Management)",
        "Environmental Permit":
            "DM NOC required for construction "
            "environmental aspects",
        "Building":
            "Dubai Building Permit — "
            "Dubai Municipality",
        "ISO Standard":
            "ISO 14001:2015 EMS | "
            "DM Green Building Regulations",
    },
}

# ════════════════════════════════════════════════════════
# ENVIRONMENTAL ASPECTS LIBRARY
# ════════════════════════════════════════════════════════

ENV_ASPECTS = {
    "💨 Air Quality & Dust": {
        "color":   "5D4037",
        "aspects": [
            "Dust generation from excavation",
            "Dust from concrete cutting & grinding",
            "Vehicle exhaust emissions",
            "Generator exhaust & fumes",
            "Spray painting & solvent vapours",
            "Burning & smoke emissions",
            "Asbestos fibre release (if applicable)",
            "Diesel fume from plant & equipment",
        ],
        "controls": [
            "Water spraying on dusty surfaces — "
            "minimum 3 times daily in hot/dry weather",
            "Cover all vehicles transporting "
            "materials with tarpaulin",
            "Suppress dust on site roads — "
            "wheel wash at site exit",
            "Wet cutting methods for concrete "
            "and masonry",
            "Limit vehicle speed to 15 km/h on "
            "unpaved roads",
            "Shut down generators when not in use",
            "Use low-sulphur diesel (50 ppm max) "
            "for all plant",
            "Site boundary hoarding to contain dust",
            "Wind monitoring — stop dusty works "
            "in wind speeds >20 km/h",
            "Vegetation clearance to minimum "
            "area required",
        ],
    },
    "💧 Water Management": {
        "color":   "0D47A1",
        "aspects": [
            "Dewatering discharge to drainage",
            "Concrete washout water",
            "Surface water runoff contamination",
            "Sewage from welfare facilities",
            "Chemical/oil contamination of water",
            "Groundwater contamination",
            "Stormwater management",
            "Water wastage / conservation",
        ],
        "controls": [
            "Dewatering discharge quality tested "
            "before discharge — pH 6-9, "
            "TSS <200 mg/L",
            "Concrete washout areas — lined, "
            "contained, separated from drainage",
            "Drip trays under all plant and "
            "equipment — minimum 110% capacity",
            "Chemical storage on impermeable bunds",
            "No direct discharge to wadis, sea, "
            "or natural watercourses",
            "Sewage via sealed portable toilets "
            "— licensed vacuum tanker collection",
            "Stormwater diversion channels around "
            "excavations",
            "Water reuse for dust suppression "
            "where quality permits",
            "Track water consumption monthly",
        ],
    },
    "🗑️ Waste Management": {
        "color":   "1B5E20",
        "aspects": [
            "Construction waste — inert debris",
            "Hazardous waste — chemicals, oils",
            "Mixed solid waste — general rubbish",
            "Excavated material — soil/rock",
            "Concrete and masonry waste",
            "Packaging and plastic waste",
            "Scrap metal and steel",
            "Electrical / WEEE waste",
            "Medical waste from first aid",
            "Sewage and sanitary waste",
        ],
        "controls": [
            "Waste hierarchy: Reduce → Reuse → "
            "Recycle → Recover → Dispose",
            "Segregated waste skips: General, "
            "Recyclable, Hazardous — clearly labelled",
            "Licensed waste contractor for "
            "all disposal — waste transfer notes kept",
            "Hazardous waste manifest system "
            "— cradle-to-grave tracking",
            "No burning of waste on site — "
            "strictly prohibited",
            "Target minimum 70% waste recycling "
            "/ diversion from landfill",
            "Excess excavated material to "
            "approved licensed facility",
            "Waste generation tracking — monthly "
            "report with quantities by type",
            "Scrap metal to licensed scrap dealer",
            "Packaging minimisation — "
            "bulk delivery preferred",
        ],
    },
    "🛢️ Spill Prevention & Soil": {
        "color":   "BF360C",
        "aspects": [
            "Fuel storage and handling spills",
            "Oil and lubricant leaks from plant",
            "Chemical storage and use spills",
            "Concrete spills",
            "Paint and solvents",
            "Sewage spills",
            "Soil contamination",
        ],
        "controls": [
            "All fuel tanks on bunded areas — "
            "110% secondary containment",
            "Spill kits at each fuel storage area, "
            "chemical area, and plant depot",
            "Emergency spill response plan "
            "— 24-hour response",
            "Drip trays under all refuelling "
            "activities",
            "No refuelling within 30m of "
            "drainage, watercourse, or wells",
            "Chemical inventory maintained — "
            "SDS/MSDS available for all chemicals",
            "Concrete pump washout in designated "
            "lined area only",
            "Contaminated soil segregated, tested, "
            "and disposed via licensed contractor",
            "All spills reported to EMP coordinator "
            "within 1 hour",
            "Spill response training for all "
            "relevant workers",
        ],
    },
    "🔊 Noise & Vibration": {
        "color":   "4A148C",
        "aspects": [
            "Construction equipment noise",
            "Piling and percussive works",
            "Concrete breaking / demolition",
            "Generator noise",
            "Vehicle and traffic noise",
            "Ground vibration from piling/blasting",
            "Vibration impact on adjacent structures",
        ],
        "controls": [
            "Daytime only for noisy works — "
            "Sunday to Thursday 7:00–20:00 "
            "(check local bylaws)",
            "No loud works on Fridays and "
            "public holidays without permit",
            "Plant and equipment — use "
            "lowest-noise option available",
            "Generators fitted with acoustic "
            "enclosures where practicable",
            "Noise monitoring at site boundary — "
            "monthly or when complaints received",
            "Noise barrier/hoarding between "
            "source and receptors",
            "Vibration monitoring during piling — "
            "limit 5 mm/s PPV adjacent structures",
            "Community complaint procedure "
            "posted at site entrance",
            "Night work — separate permit and "
            "noise assessment required",
        ],
    },
    "🌡️ Heat & Energy": {
        "color":   "E65100",
        "aspects": [
            "Generator fuel consumption",
            "Vehicle fleet fuel use",
            "Electricity consumption — offices",
            "Carbon emissions from plant",
            "Heat island effect on local environment",
            "Summer heat impact on workers",
        ],
        "controls": [
            "Switch off all plant and generators "
            "when not in use",
            "Regular maintenance — reduces fuel "
            "consumption by up to 15%",
            "LED lighting throughout site",
            "Solar-powered welfare and office units "
            "where feasible",
            "Track fuel consumption monthly — "
            "carbon footprint reporting",
            "Vehicle idle time limit — max 3 "
            "minutes when stationary",
            "Journey planning to reduce vehicle "
            "movements",
            "Shade structures for all work areas "
            "during summer",
        ],
    },
    "🌱 Ecology & Biodiversity": {
        "color":   "2E7D32",
        "aspects": [
            "Vegetation clearance",
            "Wildlife habitat disturbance",
            "Protected species impact",
            "Light pollution at night",
            "Invasive species introduction",
            "Tree removal / damage",
        ],
        "controls": [
            "Pre-clearance ecological survey — "
            "check for protected species",
            "Tree protection orders — "
            "no removal without approval",
            "Minimum footprint clearing — "
            "only clear what is necessary",
            "Wildlife displacement survey before "
            "starting work",
            "Night lighting directed downward — "
            "no light spill beyond site boundary",
            "Invasive species removed from "
            "excavated material if found",
            "Post-construction landscaping "
            "to restore cleared areas",
            "No introduction of non-native plants "
            "in landscaping",
        ],
    },
    "🚛 Traffic & Transport": {
        "color":   "37474F",
        "aspects": [
            "Traffic congestion on local roads",
            "Road surface damage from heavy vehicles",
            "Mud/debris tracked onto public roads",
            "Air pollution from vehicle movements",
            "Pedestrian conflict with site vehicles",
            "Delivery vehicle idling",
        ],
        "controls": [
            "Wheel wash at site exit — mandatory "
            "for all vehicles leaving site",
            "Construction traffic management plan "
            "agreed with local authorities",
            "No HGV movements during peak hours "
            "(07:00–09:00 and 17:00–19:00)",
            "Vehicle maintenance records — "
            "clean, maintained engines only",
            "Mud on public road — immediate "
            "cleaning by site team",
            "Designated haul routes — "
            "signed and enforced",
            "Delivery scheduling — "
            "minimize simultaneous arrivals",
        ],
    },
}

# ════════════════════════════════════════════════════════
# SIGNIFICANCE RATING MATRIX
# ════════════════════════════════════════════════════════

SIGNIFICANCE_MATRIX = {
    # (likelihood, severity): rating
    (1, 1): "LOW",    (1, 2): "LOW",    (1, 3): "LOW",
    (1, 4): "MEDIUM", (1, 5): "MEDIUM",
    (2, 1): "LOW",    (2, 2): "LOW",    (2, 3): "MEDIUM",
    (2, 4): "MEDIUM", (2, 5): "HIGH",
    (3, 1): "LOW",    (3, 2): "MEDIUM", (3, 3): "MEDIUM",
    (3, 4): "HIGH",   (3, 5): "HIGH",
    (4, 1): "MEDIUM", (4, 2): "MEDIUM", (4, 3): "HIGH",
    (4, 4): "HIGH",   (4, 5): "SIGNIFICANT",
    (5, 1): "MEDIUM", (5, 2): "HIGH",   (5, 3): "HIGH",
    (5, 4): "SIGNIFICANT", (5, 5): "SIGNIFICANT",
}

SIG_COLORS = {
    "LOW":         "00AF50",
    "MEDIUM":      "FFFF00",
    "HIGH":        "FFC000",
    "SIGNIFICANT": "FF0000",
}


# ════════════════════════════════════════════════════════
# DOCX HELPERS
# ════════════════════════════════════════════════════════

def shd_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)


def shd_cell(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def cell_borders(cell, color="1B5E20"):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def sec_hdr(doc, text, fill="1B5E20", size=11):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    shd_para(p, fill)
    run           = p.add_run("  " + text.upper())
    run.bold      = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p


def add_body(doc, text, size=10,
             bold=False, indent=0.5):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent)
    run           = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold      = bold
    return p


def add_bullet(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    return p


def safe_list(raw, fallback=None):
    """Robust JSON list parser with fallback."""
    if fallback is None:
        fallback = []
    raw = re.sub(r'```json|```', '', raw).strip()
    try:
        r = json.loads(raw)
        if isinstance(r, list):
            return r
        if isinstance(r, dict):
            for v in r.values():
                if isinstance(v, list):
                    return v
    except Exception:
        pass
    s = raw.find('[')
    e = raw.rfind(']') + 1
    if s != -1 and e > 0:
        try:
            r = json.loads(raw[s:e])
            if isinstance(r, list):
                return r
        except Exception:
            pass
    objs   = re.findall(r'\{[^\{\}]+\}', raw)
    parsed = []
    for o in objs:
        try:
            parsed.append(json.loads(o))
        except Exception:
            pass
    return parsed if parsed else fallback


def ask(client, prompt, max_tok=1000):
    """Single AI call — short and reliable."""
    r = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tok,
        messages=[{
            "role":    "user",
            "content": prompt
        }]
    )
    return r.content[0].text.strip()


# ════════════════════════════════════════════════════════
# BUILD FULL EMP DOCX
# ════════════════════════════════════════════════════════

def build_emp_docx(emp_data, meta,
                   emirate, selected_aspects):
    doc = Document()

    # A4 Portrait
    sec               = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    regs = REGULATIONS.get(emirate, {})

    # ── COVER ─────────────────────────────────
    cover           = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(cover, '1B5E20')
    cr = cover.add_run(
        "\n\n  🌿  ENVIRONMENTAL MANAGEMENT PLAN"
        "\n\n  "
        + meta.get('company', 'ENGC')
        + "  \n\n"
    )
    cr.bold           = True
    cr.font.size      = Pt(20)
    cr.font.color.rgb = RGBColor(255, 255, 255)

    sub           = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(sub, '2E7D32')
    sr = sub.add_run(
        "\n  PROJECT:  "
        + meta.get('project', '').upper()
        + "  \n  "
        + meta.get('location', '').upper()
        + "  |  "
        + emirate.upper()
        + ", UAE  \n  "
        + regs.get(
            'Local Authority',
            'Environmental Authority'
        )
        + "  \n"
    )
    sr.bold           = True
    sr.font.size      = Pt(13)
    sr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_page_break()

    # ── DOCUMENT CONTROL ──────────────────────
    sec_hdr(
        doc, "1. DOCUMENT CONTROL",
        fill="1F3864", size=10
    )
    ctrl_tbl       = doc.add_table(rows=0, cols=4)
    ctrl_tbl.style = 'Table Grid'

    def ctrl_row(tbl, l1, v1, l2, v2):
        row   = tbl.add_row()
        cells = row.cells
        for i, (l, v) in enumerate(
            [(l1, v1), (l2, v2)]
        ):
            lc = cells[i * 2]
            vc = cells[i * 2 + 1]
            lc.text = l
            lc.paragraphs[0].runs[0].bold      = True
            lc.paragraphs[0].runs[0].font.size = Pt(9)
            shd_cell(lc, 'E8F5E9')
            cell_borders(lc, '1F3864')
            vc.text = str(v) if v else "—"
            vc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_borders(vc, '1F3864')

    ctrl_row(ctrl_tbl,
             "Document Title",
             "Environmental Management Plan (EMP)",
             "Document No.", meta.get('doc_no', ''))
    ctrl_row(ctrl_tbl,
             "Project", meta.get('project', ''),
             "Date Issued",
             str(meta.get('emp_date', date.today())))
    ctrl_row(ctrl_tbl,
             "Location", meta.get('location', ''),
             "Revision", meta.get('revision', 'Rev 00'))
    ctrl_row(ctrl_tbl,
             "Company", meta.get('company', 'ENGC'),
             "Prepared By", meta.get('prepared_by', ''))
    ctrl_row(ctrl_tbl,
             "Client", meta.get('client', ''),
             "Approved By", meta.get('approved_by', ''))
    ctrl_row(ctrl_tbl,
             "Environmental Authority",
             regs.get('Local Authority', ''),
             "ISO Standard",
             regs.get('ISO Standard', 'ISO 14001:2015'))
    ctrl_row(ctrl_tbl,
             "Primary Legislation",
             regs.get('Primary Law', ''),
             "Review Frequency",
             "Quarterly or after significant change")

    doc.add_paragraph()

    # ── ENVIRONMENTAL POLICY ──────────────────
    sec_hdr(
        doc,
        "2. ENVIRONMENTAL POLICY STATEMENT",
        fill="1B5E20"
    )
    policy = emp_data.get(
        'policy',
        meta.get('company', 'ENGC')
        + " is committed to protecting the "
        "environment and minimising the "
        "environmental impact of all construction "
        "activities. We comply with all applicable "
        "UAE environmental legislation and "
        "continuously improve our environmental "
        "performance in accordance with "
        "ISO 14001:2015."
    )
    add_body(doc, policy, size=10)
    doc.add_paragraph()

    # ── PROJECT DESCRIPTION ───────────────────
    sec_hdr(
        doc, "3. PROJECT DESCRIPTION & SCOPE",
        fill="1F3864"
    )
    proj_tbl       = doc.add_table(rows=0, cols=4)
    proj_tbl.style = 'Table Grid'
    proj_info = [
        ("Project Name",   meta.get('project', ''),
         "Location",       meta.get('location', '')),
        ("Client",         meta.get('client', ''),
         "Contract No.",   meta.get('contract_no', '')),
        ("Contractor",     meta.get('company', ''),
         "PMC",            meta.get('pmc', '')),
        ("Project Type",   meta.get('project_type', ''),
         "Duration",       meta.get('duration', '')),
        ("Site Area",      meta.get('site_area', ''),
         "Peak Workforce", meta.get('workers', '')),
        ("Consultant",     meta.get('consultant', ''),
         "Env. Coordinator",
         meta.get('env_coord', '')),
    ]
    for l1, v1, l2, v2 in proj_info:
        row   = proj_tbl.add_row()
        cells = row.cells
        for i, (l, v) in enumerate(
            [(l1, v1), (l2, v2)]
        ):
            lc = cells[i * 2]
            vc = cells[i * 2 + 1]
            lc.text = l
            lc.paragraphs[0].runs[0].bold      = True
            lc.paragraphs[0].runs[0].font.size = Pt(9)
            shd_cell(lc, 'E8F5E9')
            cell_borders(lc, '1F3864')
            vc.text = str(v) if v else "—"
            vc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_borders(vc, '1F3864')

    doc.add_paragraph()
    add_body(doc,
             emp_data.get(
                 'project_description',
                 'Refer to contract documents.'
             ))
    doc.add_paragraph()

    # ── LEGAL REGISTER ────────────────────────
    sec_hdr(
        doc,
        "4. LEGAL REGISTER & REGULATORY "
        "REQUIREMENTS — " + emirate.upper(),
        fill="1F3864"
    )
    leg_tbl       = doc.add_table(rows=1, cols=4)
    leg_tbl.style = 'Table Grid'
    for cell, h in zip(
        leg_tbl.rows[0].cells,
        ["Regulation / Standard",
         "Requirement Summary",
         "Applies To",
         "Compliance Method"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F3864')
        cell_borders(cell, '1F3864')
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    legal_reqs = emp_data.get('legal_reqs', [])
    for i, req in enumerate(legal_reqs):
        row = leg_tbl.add_row()
        if isinstance(req, dict):
            r_reg = req.get('regulation', '')
            r_req = req.get('requirement', '')
            r_app = req.get('applies_to', '')
            r_com = req.get('compliance', '')
        else:
            r_reg = str(req)
            r_req = ''
            r_app = 'All activities'
            r_com = 'Implementation & monitoring'
        vals = [r_reg, r_req, r_app, r_com]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F5E9')
            cell_borders(c, '1F3864')

    doc.add_paragraph()

    # ── ENVIRONMENTAL ASPECTS & IMPACTS ───────
    sec_hdr(
        doc,
        "5. ENVIRONMENTAL ASPECTS, IMPACTS "
        "& SIGNIFICANCE ASSESSMENT (ISO 14001:2015)",
        fill="1B5E20"
    )

    add_body(
        doc,
        "Significance Rating = Likelihood (1-5) × "
        "Severity (1-5). Rating > 12 = SIGNIFICANT "
        "(requires enhanced controls). "
        "4-12 = HIGH. 2-4 = MEDIUM. 1-2 = LOW.",
        size=9,
        bold=False
    )
    doc.add_paragraph()

    # Legend
    leg_p = doc.add_paragraph()
    shd_para(leg_p, '37474F')
    lr = leg_p.add_run(
        "  SIGNIFICANCE LEGEND:  "
        "🟢 LOW (1-2)  "
        "🟡 MEDIUM (3-6)  "
        "🟠 HIGH (8-12)  "
        "🔴 SIGNIFICANT (>12)  "
        "★ = Legal requirement"
    )
    lr.bold           = True
    lr.font.size      = Pt(8)
    lr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    aspects_tbl       = doc.add_table(rows=1, cols=8)
    aspects_tbl.style = 'Table Grid'
    asp_hdrs = [
        ("Category",     Cm(2.2)),
        ("Aspect",       Cm(3.0)),
        ("Impact",       Cm(2.5)),
        ("L\n(1-5)",     Cm(0.8)),
        ("S\n(1-5)",     Cm(0.8)),
        ("Score",        Cm(0.8)),
        ("Significance", Cm(1.5)),
        ("Key Control Measure", Cm(5.4)),
    ]
    for (h, w), cell in zip(
        asp_hdrs, aspects_tbl.rows[0].cells
    ):
        cell.width = w
        cell.text  = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1B5E20')
        cell_borders(cell, '1B5E20')
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    aspects_data = emp_data.get('aspects', [])
    for i, asp in enumerate(aspects_data):
        if not isinstance(asp, dict):
            continue
        row = aspects_tbl.add_row()
        cat  = asp.get('category', '')
        aspt = asp.get('aspect', '')
        imp  = asp.get('impact', '')
        lik  = int(asp.get('likelihood', 3))
        sev  = int(asp.get('severity', 3))
        scr  = lik * sev
        sig  = SIGNIFICANCE_MATRIX.get(
            (min(lik,5), min(sev,5)), "MEDIUM"
        )
        if scr > 12:
            sig = "SIGNIFICANT"
        elif scr > 6:
            sig = "HIGH"
        elif scr > 3:
            sig = "MEDIUM"
        else:
            sig = "LOW"
        ctrl = asp.get('control', '')

        vals   = [cat, aspt, imp,
                  str(lik), str(sev),
                  str(scr), sig, ctrl]
        widths = [Cm(2.2), Cm(3.0), Cm(2.5),
                  Cm(0.8), Cm(0.8), Cm(0.8),
                  Cm(1.5), Cm(5.4)]
        for j, (c, v, w) in enumerate(
            zip(row.cells, vals, widths)
        ):
            c.width = w
            c.text  = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 6:  # Significance
                fill = SIG_COLORS.get(sig, 'FFFF00')
                shd_cell(c, fill)
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
            elif j in [3, 4, 5]:
                c.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
                if i % 2 == 0:
                    shd_cell(c, 'E8F5E9')
            elif i % 2 == 0:
                shd_cell(c, 'E8F5E9')
            cell_borders(c, '1B5E20')

    doc.add_paragraph()

    # ── ASPECT-SPECIFIC CONTROL SECTIONS ──────
    sec_hdr(
        doc,
        "6. ENVIRONMENTAL CONTROLS BY CATEGORY",
        fill="1B5E20"
    )

    for asp_name in selected_aspects:
        asp_data = ENV_ASPECTS.get(asp_name, {})
        color    = asp_data.get('color', '1B5E20')
        controls = asp_data.get('controls', [])

        # Category header
        cat_p = doc.add_paragraph()
        cat_p.paragraph_format.space_before = Pt(8)
        shd_para(cat_p, color)
        cat_r = cat_p.add_run(
            "  🌿  " + asp_name.upper()
        )
        cat_r.bold      = True
        cat_r.font.size = Pt(10)
        cat_r.font.color.rgb = RGBColor(255, 255, 255)

        # Controls table
        ctrl_tbl_a       = doc.add_table(
            rows=1, cols=3
        )
        ctrl_tbl_a.style = 'Table Grid'
        for cell, h in zip(
            ctrl_tbl_a.rows[0].cells,
            ["Control Measure",
             "Responsible Person",
             "Monitoring / Evidence Required"]
        ):
            cell.text = h
            cell.paragraphs[0].runs[0].bold      = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].runs[0]\
                .font.color.rgb = RGBColor(255,255,255)
            shd_cell(cell, color)
            cell_borders(cell, color)

        for i, ctrl in enumerate(controls):
            row  = ctrl_tbl_a.add_row()
            vals = [
                str(ctrl),
                "Environmental Coordinator / "
                "Site Supervisor",
                "Daily / Weekly inspection record"
            ]
            for j, (c, v) in enumerate(
                zip(row.cells, vals)
            ):
                c.text = v
                c.paragraphs[0].runs[0].font.size = Pt(8)
                if j == 0:
                    c.paragraphs[0].runs[0].bold = True
                if i % 2 == 0:
                    shd_cell(c, 'E8F5E9')
                cell_borders(c, color)

        doc.add_paragraph()

    # ── WASTE MANAGEMENT PLAN ─────────────────
    sec_hdr(
        doc,
        "7. WASTE MANAGEMENT PLAN — "
        "STREAMS & DISPOSAL",
        fill="2E7D32"
    )
    waste_tbl       = doc.add_table(rows=1, cols=5)
    waste_tbl.style = 'Table Grid'
    for cell, h in zip(
        waste_tbl.rows[0].cells,
        ["Waste Stream",
         "Classification",
         "Segregation",
         "Licensed Disposal Route",
         "Target (Divert from Landfill)"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '2E7D32')
        cell_borders(cell, '2E7D32')
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    waste_streams = [
        ("Concrete/Masonry Rubble",
         "Non-Hazardous / Inert",
         "Grey skip — separate",
         "Licensed inert waste facility / "
         "Crusher / Reuse as fill",
         "90% recycling target"),
        ("Steel / Rebar Scrap",
         "Non-Hazardous",
         "Metal skip — separate",
         "Licensed scrap metal dealer",
         "100% recycling"),
        ("Timber / Formwork",
         "Non-Hazardous",
         "Timber skip — separate",
         "Reuse → Wood chip / Energy recovery",
         "80% diversion"),
        ("Packaging / Plastic",
         "Non-Hazardous",
         "Recycling skip — blue",
         "Licensed recycler",
         "70% recycling"),
        ("General Site Waste",
         "Non-Hazardous — Mixed",
         "Black skip",
         "Licensed waste facility — MRF sorting",
         "50% diversion"),
        ("Excavated Soil (Clean)",
         "Non-Hazardous / Inert",
         "Stockpile — separate from contaminated",
         "Reuse on site / Licensed soil facility",
         "Max reuse on site"),
        ("Excavated Soil (Contaminated)",
         "Potentially Hazardous",
         "Separate covered stockpile",
         "Licensed contaminated soil contractor — "
         "test first",
         "Treatment > Disposal"),
        ("Oils / Lubricants / Filters",
         "Hazardous",
         "Sealed containers — hazardous area",
         "Licensed hazardous waste contractor",
         "Reuse / Re-refine"),
        ("Fuel / Chemical Containers",
         "Hazardous",
         "Sealed — hazardous waste area",
         "Licensed chemical waste contractor",
         "0% to general waste"),
        ("Paint / Solvents",
         "Hazardous",
         "Sealed — hazardous waste area",
         "Licensed hazardous waste contractor",
         "0% to general waste"),
        ("Sewage / Sanitary Waste",
         "Hazardous / Biological",
         "Sealed portable toilets",
         "Licensed vacuum tanker — "
         "approved treatment plant",
         "100% proper treatment"),
        ("Medical / First Aid Waste",
         "Hazardous / Clinical",
         "Yellow clinical waste bag",
         "Licensed clinical waste contractor",
         "0% to general waste"),
        ("WEEE / Electrical Waste",
         "Hazardous",
         "Separate WEEE collection point",
         "Authorised WEEE handler",
         "100% recycling"),
        ("Concrete Washout Water",
         "Non-Hazardous — pH >9",
         "Dedicated lined washout pit",
         "Allow to settle, neutralise to pH 6-9, "
         "then discharge",
         "Zero direct discharge to drains"),
    ]
    for i, (ws, cls, seg, disp, tgt) in enumerate(
        waste_streams
    ):
        row  = waste_tbl.add_row()
        vals = [ws, cls, seg, disp, tgt]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if j == 1:
                if "Hazardous" in cls:
                    shd_cell(c, 'FFF9C4')
                else:
                    shd_cell(c, 'E8F5E9')
            elif i % 2 == 0:
                shd_cell(c, 'F1F8E9')
            cell_borders(c, '2E7D32')

    doc.add_paragraph()

    # ── ENVIRONMENTAL MONITORING ──────────────
    sec_hdr(
        doc,
        "8. ENVIRONMENTAL MONITORING PROGRAMME",
        fill="0D47A1"
    )
    mon_tbl       = doc.add_table(rows=1, cols=5)
    mon_tbl.style = 'Table Grid'
    for cell, h in zip(
        mon_tbl.rows[0].cells,
        ["Parameter",
         "Method / Equipment",
         "Standard / Limit",
         "Frequency",
         "Record / Action if Exceeded"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '0D47A1')
        cell_borders(cell, '0D47A1')

    if emirate == "Abu Dhabi":
        dust_std  = "EAD: PM10 <150 μg/m³ (24hr)"
        noise_std = "EAD: 55 dB(A) day / 45 night"
        water_std = "EAD: pH 6-9, TSS <200 mg/L"
        vib_std   = "5 mm/s PPV adjacent structures"
    else:
        dust_std  = "DM: PM10 <150 μg/m³ (24hr)"
        noise_std = "DM: 55 dB(A) residential day"
        water_std = "DM: pH 6-9, TSS <200 mg/L"
        vib_std   = "5 mm/s PPV (BS 5228 Guidance)"

    mon_items = [
        ("Dust / PM10",
         "Visual observation + dust monitor "
         "at boundary",
         dust_std,
         "Daily visual / Monthly monitor",
         "Stop dusty works — water spray — "
         "investigate if limit exceeded"),
        ("Noise at Boundary",
         "Calibrated sound level meter "
         "(dB(A) — slow response)",
         noise_std,
         "Monthly + when complaints received",
         "Reduce source — add barriers — "
         "reschedule if limit exceeded"),
        ("Vibration",
         "Seismograph at nearest structure",
         vib_std,
         "Continuous during piling/demolition",
         "Stop work — structural survey — "
         "reduce energy if limit exceeded"),
        ("Dewatering Discharge Quality",
         "pH meter + turbidity meter on site",
         water_std,
         "Every discharge event — before release",
         "Do NOT discharge if out of limits — "
         "treat further then retest"),
        ("Waste Generation",
         "Waste transfer notes + skip weights",
         "Track monthly by stream",
         "Monthly — per waste stream",
         "Monthly environmental report"),
        ("Fuel / Chemical Storage",
         "Visual inspection — bund integrity",
         "0 spills / 110% containment",
         "Weekly inspection",
         "Repair bund — clean spill — "
         "report to Env. Coordinator"),
        ("Vehicle Emissions",
         "Visual check — black smoke / idle time",
         "No visible black smoke / "
         "idle max 3 minutes",
         "Monthly",
         "Remove vehicle from site for repair"),
        ("Water Consumption",
         "Meter readings on tankers / supplies",
         "Track monthly — "
         "set reduction target",
         "Monthly",
         "Investigate high usage — "
         "repair leaks"),
        ("Environmental Complaints",
         "Community complaint log",
         "Zero valid complaints target",
         "Ongoing",
         "Respond within 24 hours — "
         "investigate — report"),
    ]
    for i, (param, method, std, freq, action) in (
        enumerate(mon_items)
    ):
        row  = mon_tbl.add_row()
        vals = [param, method, std, freq, action]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8EAF6')
            cell_borders(c, '0D47A1')

    doc.add_paragraph()

    # ── ENVIRONMENTAL OBJECTIVES & TARGETS ────
    sec_hdr(
        doc,
        "9. ENVIRONMENTAL OBJECTIVES, "
        "TARGETS & PROGRAMMES (ISO 14001:2015)",
        fill="2E7D32"
    )
    obj_tbl       = doc.add_table(rows=1, cols=5)
    obj_tbl.style = 'Table Grid'
    for cell, h in zip(
        obj_tbl.rows[0].cells,
        ["Objective",
         "Target",
         "KPI",
         "Responsible",
         "Completion"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '2E7D32')
        cell_borders(cell, '2E7D32')

    objectives = [
        ("Minimise waste to landfill",
         "70% waste diversion from landfill",
         "% waste recycled/reused monthly",
         "Env. Coordinator",
         "Monthly — project end"),
        ("Control dust emissions",
         "Zero valid dust complaints",
         "No. complaints per month",
         "HSE Officer",
         "Ongoing"),
        ("Prevent water pollution",
         "Zero unauthorised discharges",
         "No. spill incidents per month",
         "Site Supervisor",
         "Ongoing"),
        ("Reduce energy / fuel use",
         "10% reduction vs baseline",
         "Litres fuel / m² completed monthly",
         "Plant Manager",
         "Quarterly review"),
        ("Minimise noise nuisance",
         "Zero valid noise complaints",
         "No. complaints per month",
         "HSE Officer",
         "Ongoing"),
        ("Full legal compliance",
         "Zero regulatory enforcement actions",
         "No. enforcement notices",
         "Env. Coordinator",
         "Ongoing"),
        ("Protect soil from contamination",
         "Zero soil contamination incidents",
         "No. spill/contamination events",
         "Site Supervisor",
         "Ongoing"),
        ("Carbon footprint tracking",
         "Monthly carbon emissions reported",
         "Tonnes CO2e per month",
         "Env. Coordinator",
         "Monthly report"),
    ]
    for i, (obj, tgt, kpi, resp, comp) in enumerate(
        objectives
    ):
        row  = obj_tbl.add_row()
        vals = [obj, tgt, kpi, resp, comp]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F5E9')
            cell_borders(c, '2E7D32')

    doc.add_paragraph()

    # ── ROLES & RESPONSIBILITIES ──────────────
    sec_hdr(
        doc,
        "10. ENVIRONMENTAL ROLES & RESPONSIBILITIES",
        fill="1F3864"
    )
    resp_tbl       = doc.add_table(rows=1, cols=3)
    resp_tbl.style = 'Table Grid'
    for cell, h in zip(
        resp_tbl.rows[0].cells,
        ["Role / Title",
         "Environmental Responsibilities",
         "Authority & Accountability"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F3864')
        cell_borders(cell, '1F3864')

    responsibilities = [
        ("Projects Director",
         "Approve and endorse the EMP; "
         "ensure adequate environmental resources; "
         "review environmental performance quarterly",
         "Full environmental authority; "
         "Stop Work Authority"),
        ("Project Manager",
         "Implement the EMP; ensure all site "
         "activities comply; integrate environmental "
         "controls into daily operations",
         "Site authority for environmental "
         "decisions; escalate to PD if required"),
        ("HSE / Environmental Coordinator ("
         + meta.get('env_coord', 'Named Person')
         + ")",
         "Day-to-day implementation of EMP; "
         "conduct environmental inspections; "
         "maintain monitoring records; "
         "investigate incidents; submit reports",
         "Stop Work Authority for "
         "environmental violations; report to PM"),
        ("Site Supervisor / Foreman",
         "Enforce controls at work front level; "
         "brief workers on environmental "
         "requirements; report incidents immediately",
         "Stop Work Authority at work front level"),
        ("All Workers",
         "Follow environmental procedures; "
         "segregate waste correctly; "
         "report spills immediately; "
         "attend environmental induction",
         "Stop Work Authority for "
         "clear environmental violations"),
        ("Subcontractors",
         "Comply with client EMP requirements; "
         "submit subcontractor environmental plan; "
         "attend environmental inductions",
         "Accountable for their own operations"),
    ]
    for i, (role, duty, auth) in enumerate(
        responsibilities
    ):
        row  = resp_tbl.add_row()
        vals = [role, duty, auth]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F0FE')
            cell_borders(c, '1F3864')

    doc.add_paragraph()

    # ── TRAINING ──────────────────────────────
    sec_hdr(
        doc,
        "11. ENVIRONMENTAL TRAINING & AWARENESS",
        fill="4A148C"
    )
    train_tbl       = doc.add_table(rows=1, cols=4)
    train_tbl.style = 'Table Grid'
    for cell, h in zip(
        train_tbl.rows[0].cells,
        ["Training Topic",
         "Who",
         "Frequency",
         "Delivered By"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '4A148C')
        cell_borders(cell, '4A148C')

    training = [
        ("Environmental Induction (EMP overview, "
         "waste segregation, spill response)",
         "ALL workers on Day 1",
         "Once on joining",
         "Env. Coordinator / HSE Officer"),
        ("Waste Segregation & Disposal Procedure",
         "All site workers",
         "Toolbox talk — monthly",
         "Site Supervisor / HSE"),
        ("Spill Prevention & Emergency Response",
         "All workers handling chemicals/fuel",
         "Before first exposure + annual",
         "Env. Coordinator"),
        ("Dust & Noise Control",
         "Operators and supervisors",
         "Annual + when method changes",
         "HSE Officer"),
        ("Environmental Legal Requirements — "
         + emirate,
         "Supervisors and Env. Coordinator",
         "Annual",
         "External Environmental Consultant"),
        ("Toolbox Talk — Environmental Awareness",
         "All workers",
         "Monthly",
         "HSE Officer / Supervisor"),
        ("ISO 14001:2015 Awareness",
         "Management and coordinators",
         "Annual",
         "Env. Coordinator / External"),
        ("Incident Reporting & Investigation",
         "Supervisors + Env. Coordinator",
         "Annual",
         "HSE Manager"),
    ]
    for i, (topic, who, freq, by) in enumerate(
        training
    ):
        row  = train_tbl.add_row()
        vals = [topic, who, freq, by]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'F3E5F5')
            cell_borders(c, '4A148C')

    doc.add_paragraph()

    # ── INCIDENT REPORTING ────────────────────
    sec_hdr(
        doc,
        "12. ENVIRONMENTAL INCIDENT REPORTING "
        "& NOTIFICATION",
        fill="C00000"
    )
    inc_tbl       = doc.add_table(rows=1, cols=4)
    inc_tbl.style = 'Table Grid'
    for cell, h in zip(
        inc_tbl.rows[0].cells,
        ["Incident Type",
         "Internal Notification",
         "Regulatory Notification",
         "Timeframe"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, 'C00000')
        cell_borders(cell, 'C00000')

    auth = regs.get('Local Authority',
                    'Environmental Authority')
    inc_rows = [
        ("Major Spill (>200L fuel/chemical)",
         "PM + PD immediately",
         auth + " + Civil Defence 997",
         "IMMEDIATELY"),
        ("Minor Spill (contained, <200L)",
         "Env. Coordinator same day",
         "Internal — manage & report in "
         "monthly env. report",
         "Same day"),
        ("Unauthorised Discharge to Drainage",
         "PM + Env. Coordinator immediately",
         auth + " immediately",
         "IMMEDIATELY"),
        ("Significant Dust Nuisance / Complaint",
         "Env. Coordinator within 1 hour",
         auth + " if serious",
         "Within 24 hours"),
        ("Noise Complaint from Neighbour",
         "Env. Coordinator same day",
         "Log + respond to complainant",
         "Same day — respond within 24hrs"),
        ("Tree / Vegetation Damage",
         "PM + Env. Coordinator",
         auth + " — tree protection order check",
         "Within 24 hours"),
        ("Discovery of Contaminated Ground",
         "PM + Env. Coordinator immediately",
         auth + " — stop disturbing site",
         "IMMEDIATELY — stop works"),
        ("Protected Species Found",
         "PM + Env. Coordinator immediately",
         auth + " — ecological survey required",
         "IMMEDIATELY — stop works"),
    ]
    for i, (inc, internal, reg_not, tf) in enumerate(
        inc_rows
    ):
        row  = inc_tbl.add_row()
        vals = [inc, internal, reg_not, tf]
        tf_color = (
            'FF0000'
            if 'IMMEDIATELY' in tf
            else 'FFC000'
            if '24' in tf
            else 'FFFF00'
        )
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if j == 3:
                shd_cell(c, tf_color)
                c.paragraphs[0].runs[0].bold = True
            elif i % 2 == 0:
                shd_cell(c, 'FFEBEE')
            cell_borders(c, 'C00000')

    doc.add_paragraph()

    # ── INSPECTION SCHEDULE ───────────────────
    sec_hdr(
        doc,
        "13. ENVIRONMENTAL INSPECTION SCHEDULE",
        fill="1F3864"
    )
    insp_tbl       = doc.add_table(rows=1, cols=4)
    insp_tbl.style = 'Table Grid'
    for cell, h in zip(
        insp_tbl.rows[0].cells,
        ["Inspection Activity",
         "Frequency",
         "Conducted By",
         "Record / Form"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '1F3864')
        cell_borders(cell, '1F3864')

    inspections = [
        ("Daily site environmental walk",
         "Daily",
         "HSE Officer / Env. Coordinator",
         "Daily inspection checklist"),
        ("Waste skips — correct segregation",
         "Daily",
         "Site Supervisor",
         "Daily checklist"),
        ("Bund integrity & spill kits",
         "Weekly",
         "Env. Coordinator",
         "Bund inspection form"),
        ("Dust suppression effectiveness",
         "Daily in dry/windy weather",
         "HSE Officer",
         "Daily checklist"),
        ("Wheel wash operation",
         "Daily during operations",
         "Security / Site Supervisor",
         "Daily checklist"),
        ("Dewatering discharge quality",
         "Every discharge event",
         "Env. Coordinator",
         "Discharge monitoring form"),
        ("Monthly environmental audit",
         "Monthly",
         "Env. Coordinator",
         "Monthly EMP Audit Report"),
        ("Quarterly management review",
         "Quarterly",
         "Project Manager",
         "Quarterly review minutes"),
        ("Environmental KPI report",
         "Monthly",
         "Env. Coordinator",
         "Monthly KPI dashboard"),
        ("Subcontractor env. compliance check",
         "Monthly",
         "Env. Coordinator",
         "Subcontractor env. audit form"),
        ("Regulatory authority inspection",
         "As required by authority",
         "PM + Env. Coordinator",
         "Authority inspection record"),
    ]
    for i, (insp, freq, by, rec) in enumerate(
        inspections
    ):
        row  = insp_tbl.add_row()
        vals = [insp, freq, by, rec]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if i % 2 == 0:
                shd_cell(c, 'E8F0FE')
            cell_borders(c, '1F3864')

    doc.add_paragraph()

    # ── ENVIRONMENTAL KPI DASHBOARD ───────────
    sec_hdr(
        doc,
        "14. ENVIRONMENTAL KPI TRACKER "
        "(Complete Monthly)",
        fill="37474F"
    )
    kpi_tbl       = doc.add_table(rows=1, cols=6)
    kpi_tbl.style = 'Table Grid'
    for cell, h in zip(
        kpi_tbl.rows[0].cells,
        ["KPI",
         "Unit",
         "Target",
         "Month 1",
         "Month 2",
         "Month 3"]
    ):
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, '37474F')
        cell_borders(cell, '37474F')
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    kpis = [
        ("Total Waste Generated",
         "Tonnes", "Track monthly", "", "", ""),
        ("Waste Recycled/Diverted",
         "%", ">70%", "", "", ""),
        ("Hazardous Waste Disposed",
         "Kg", "Track monthly", "", "", ""),
        ("Fuel Consumed (all plant)",
         "Litres", "Track monthly", "", "", ""),
        ("Water Consumed",
         "m³ or L", "Track monthly", "", "", ""),
        ("Dust Complaints Received",
         "No.", "0 valid", "", "", ""),
        ("Noise Complaints Received",
         "No.", "0 valid", "", "", ""),
        ("Environmental Incidents (Spills etc.)",
         "No.", "0 significant", "", "", ""),
        ("Near Miss Reports",
         "No.", ">0 reported", "", "", ""),
        ("Environmental Inspections Completed",
         "No.", "Per schedule", "", "", ""),
        ("Enforcement Notices / Fines",
         "No.", "0", "", "", ""),
        ("Environmental Training Delivered",
         "Hours", "Per plan", "", "", ""),
    ]
    for i, (kpi, unit, tgt, m1, m2, m3) in (
        enumerate(kpis)
    ):
        row  = kpi_tbl.add_row()
        vals = [kpi, unit, tgt, m1, m2, m3]
        for j, (c, v) in enumerate(
            zip(row.cells, vals)
        ):
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
            if j == 2:
                shd_cell(c, 'E8F5E9')
            elif j > 2:
                shd_cell(c, 'FFFDE7')
            elif i % 2 == 0:
                shd_cell(c, 'F5F5F5')
            cell_borders(c, '37474F')

    doc.add_paragraph()

    # ── SIGNATURES ────────────────────────────
    sec_hdr(
        doc,
        "15. APPROVAL & SIGNATURES",
        fill="1F3864"
    )
    sig_tbl       = doc.add_table(rows=3, cols=3)
    sig_tbl.style = 'Table Grid'
    for i, role in enumerate([
        "Prepared By\n(Environmental / HSE Coordinator)",
        "Reviewed By\n(Project Manager)",
        "Approved By\n(Projects Director)"
    ]):
        c = sig_tbl.rows[0].cells[i]
        c.text = role
        c.paragraphs[0].runs[0].bold      = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0]\
            .font.color.rgb = RGBColor(255,255,255)
        shd_cell(c, '1B5E20')
        cell_borders(c, '1F3864')
    for i in range(3):
        c = sig_tbl.rows[1].cells[i]
        c.text = "\n\nName: _______________\n"
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c, '1F3864')
        c = sig_tbl.rows[2].cells[i]
        c.text = (
            "Designation: _______________\n\n"
            "Signature: _______________\n\n"
            "Date: _______________"
        )
        c.paragraphs[0].runs[0].font.size = Pt(9)
        cell_borders(c, '1F3864')

    doc.add_paragraph()

    # ── DISCLAIMER ────────────────────────────
    disc = doc.add_paragraph()
    dr   = disc.add_run(
        "⚠️ IMPORTANT NOTICE: This Environmental "
        "Management Plan has been generated with "
        "AI assistance and must be reviewed and "
        "verified by a qualified Environmental "
        "Professional before implementation. "
        "All monitoring standards must be "
        "confirmed against current "
        + regs.get('Local Authority','authority')
        + " and EAD/DM requirements. "
        "This document does not constitute "
        "regulatory approval."
    )
    dr.font.size   = Pt(8)
    dr.font.italic = True
    dr.font.color.rgb = RGBColor(128, 128, 128)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════
# AI CONTENT GENERATORS
# ════════════════════════════════════════════════════════

def gen_emp_content(client, meta,
                    emirate, selected_aspects):
    """Generate EMP content via focused AI calls."""
    regs = REGULATIONS.get(emirate, {})
    reg  = (
        regs.get('Primary Law', '')
        + " | "
        + regs.get('Local Authority', '')
        + " | ISO 14001:2015"
    )
    ctx  = (
        "Project: " + meta.get('project', '')
        + " | Type: " + meta.get('project_type', '')
        + " | Emirate: " + emirate
        + " | Company: " + meta.get('company', 'ENGC')
    )

    data = {}

    # Policy
    raw = ask(
        client,
        "Write a 3-sentence Environmental Policy "
        "Statement for a construction company in "
        + emirate + ", UAE. Context: " + ctx
        + ". Reference: " + reg
        + ". Professional, committed, specific. "
        "Plain text only."
    )
    data['policy'] = raw[:600].strip()

    # Project description
    raw = ask(
        client,
        "Write 2 sentences describing the "
        "environmental scope for this construction "
        "project in " + emirate + ": " + ctx
        + ". What environmental aspects are "
        "significant for this project type? "
        "Plain text only."
    )
    data['project_description'] = raw[:400].strip()

    # Legal requirements
    raw = ask(
        client,
        "Return ONLY a JSON array of 8 "
        "environmental legal requirements for "
        "a construction project in "
        + emirate + ", UAE.\n"
        'Format: [{"regulation":"Name/number",'
        '"requirement":"What it requires",'
        '"applies_to":"What activity",'
        '"compliance":"How to comply"}]\n'
        "Include: UAE Federal Law 24/1999, "
        + regs.get('Local Authority','') + " reqs, "
        + regs.get('Waste Regulation','') + ", "
        "noise, air quality, water discharge. "
        "Return ONLY the JSON array."
    )
    data['legal_reqs'] = safe_list(raw, [
        {"regulation": "UAE Federal Law 24/1999",
         "requirement": "Environmental protection",
         "applies_to": "All activities",
         "compliance": "EMP implementation"},
        {"regulation":
             regs.get('Waste Regulation',''),
         "requirement": "Waste management",
         "applies_to": "Waste generation",
         "compliance": "Segregation + licensed disposal"},
    ])

    # Aspects & impacts
    all_aspects_rows = []
    for asp_name in selected_aspects:
        asp_data = ENV_ASPECTS.get(asp_name, {})
        raw = ask(
            client,
            "Return ONLY a JSON array of 4 "
            "environmental aspects for: "
            + asp_name + " on a construction "
            "site in " + emirate + ".\n"
            'Format: [{"category":"' + asp_name[:20]
            + '","aspect":"Activity causing impact",'
            '"impact":"Environmental impact",'
            '"likelihood":3,'
            '"severity":3,'
            '"control":"Key control measure"}]\n'
            "Use likelihood 1-5 and severity 1-5. "
            "Be specific. "
            "Return ONLY the JSON array."
        )
        items = safe_list(raw, [
            {
                "category":   asp_name[:20],
                "aspect":     "Site activities",
                "impact":     "Environmental impact",
                "likelihood": 3,
                "severity":   3,
                "control":    "Standard controls apply",
            }
        ])
        all_aspects_rows.extend(items)

    data['aspects'] = all_aspects_rows

    return data


# ════════════════════════════════════════════════════════
# MAIN UI — THREE TABS
# ════════════════════════════════════════════════════════

tab_gen, tab_register, tab_upload = st.tabs([
    "🌿 Generate Full EMP",
    "📊 Aspects Register Only",
    "📤 Upload Your EMP Format",
])


# ════════════════════════════════════════════════════════
# TAB 1 — FULL EMP GENERATOR
# ════════════════════════════════════════════════════════

with tab_gen:
    st.markdown(
        "### 🌿 Generate Complete "
        "Environmental Management Plan"
    )
    st.info(
        "Fill in your project details and select "
        "the environmental aspects relevant to "
        "your site. The AI generates a complete "
        "EMP with aspects register, legal register, "
        "waste management plan, monitoring programme, "
        "KPI tracker, and " + emirate
        + " regulatory references "
        "(ISO 14001:2015 aligned)."
    )

    # Show regulatory framework
    regs_display = REGULATIONS.get(emirate, {})
    with st.expander(
        "📋 " + emirate
        + " Environmental Regulatory Framework",
        expanded=False
    ):
        reg_cols = st.columns(2)
        for i, (k, v) in enumerate(
            regs_display.items()
        ):
            reg_cols[i % 2].markdown(
                "**" + k + ":** " + v
            )

    st.divider()
    st.subheader("📋 Project Information")

    g1, g2 = st.columns(2)
    with g1:
        g_project  = st.text_input(
            "Project Name *",
            placeholder="e.g. Bloom Living Almeria",
            key="emp_proj"
        )
        g_location = st.text_input(
            "Site Location / Address *",
            placeholder=(
                "e.g. Plot C13, Zayed City, "
                "Abu Dhabi"
            ),
            key="emp_loc"
        )
        g_client   = st.text_input(
            "Client",
            placeholder="e.g. Bloom District "
                        "Properties",
            key="emp_client"
        )
        g_company  = st.text_input(
            "Main Contractor",
            value="ENGC",
            key="emp_co"
        )
    with g2:
        g_pmc      = st.text_input(
            "PMC / Project Manager",
            placeholder="e.g. Times Development",
            key="emp_pmc"
        )
        g_consultant = st.text_input(
            "Consultant",
            placeholder="e.g. Dar Al-Handasah",
            key="emp_cons"
        )
        g_contract = st.text_input(
            "Contract Number",
            placeholder="e.g. BL-2024-001",
            key="emp_contract"
        )
        g_doc_no   = st.text_input(
            "Document Number",
            placeholder="e.g. ENGC-EMP-001",
            key="emp_docno"
        )

    st.divider()
    st.subheader("🏗️ Project Details")
    g3, g4 = st.columns(2)
    with g3:
        g_type     = st.multiselect(
            "Project Type",
            [
                "Residential Villas",
                "Community Center",
                "Service Station",
                "Commercial Building",
                "Infrastructure",
                "MEP Works",
                "Fit-Out Works",
                "Demolition",
            ],
            default=["Residential Villas"],
            key="emp_type"
        )
        g_workers  = st.number_input(
            "Peak Workforce",
            min_value=5,
            max_value=5000,
            value=150,
            key="emp_workers"
        )
        g_site_area = st.text_input(
            "Site Area",
            placeholder="e.g. 25,000 m²",
            key="emp_area"
        )
    with g4:
        g_duration = st.text_input(
            "Project Duration",
            placeholder="e.g. 24 months "
                        "(Jan 2025 – Dec 2026)",
            key="emp_duration"
        )
        g_sensitive = st.multiselect(
            "Sensitive Receptors Near Site",
            [
                "Residential properties",
                "Schools / Hospitals",
                "Mosques",
                "Protected land / Wildlife area",
                "Wadi / Water course",
                "Heritage / Archaeological site",
                "Commercial / Industrial area",
            ],
            default=["Residential properties"],
            key="emp_sensitive"
        )

    st.divider()
    st.subheader("👤 Key Personnel")
    p1, p2 = st.columns(2)
    with p1:
        g_prep     = st.text_input(
            "Prepared By (HSE/Env. Coordinator)",
            placeholder="e.g. Aftab Qamar",
            key="emp_prep"
        )
        g_env_coord = st.text_input(
            "Environmental Coordinator Name",
            placeholder="e.g. Aftab Qamar — "
                        "Senior HSSE Engineer",
            key="emp_coord"
        )
    with p2:
        g_approve  = st.text_input(
            "Approved By (Projects Director)",
            placeholder="e.g. Rami Kamal Yassin",
            key="emp_approve"
        )
        g_date     = st.date_input(
            "Issue Date",
            value=date.today(),
            key="emp_date"
        )

    g_rev = st.selectbox(
        "Revision",
        ["Rev 00","Rev 01","Rev 02","Rev 03"],
        key="emp_rev"
    )

    st.divider()
    st.subheader(
        "🌿 Select Environmental Aspects for "
        "This Site"
    )
    st.caption(
        "Select ALL aspects relevant to your "
        "project. The EMP will include specific "
        "controls and monitoring for each."
    )

    asp_names    = list(ENV_ASPECTS.keys())
    selected_asp = []
    asp_cols     = st.columns(2)
    for i, asp in enumerate(asp_names):
        col = asp_cols[i % 2]
        if col.checkbox(
            asp,
            value=True,
            key="asp_" + str(i)
        ):
            selected_asp.append(asp)

    if st.button(
        "🌿 Generate Complete EMP (.docx)",
        type="primary",
        use_container_width=True,
        key="emp_gen_btn"
    ):
        if not g_project or not g_location:
            st.error(
                "Please enter project name "
                "and location."
            )
        elif not selected_asp:
            st.error(
                "Please select at least one "
                "environmental aspect."
            )
        else:
            prog = st.progress(
                0,
                text="Starting EMP generation..."
            )
            try:
                client = anthropic.Anthropic(
                    api_key=os.getenv(
                        "ANTHROPIC_API_KEY"
                    )
                )

                meta = {
                    "project":      g_project,
                    "location":     g_location,
                    "client":       g_client,
                    "company":      g_company,
                    "pmc":          g_pmc,
                    "consultant":   g_consultant,
                    "contract_no":  g_contract,
                    "doc_no":       g_doc_no,
                    "project_type": ", ".join(g_type),
                    "workers":      str(g_workers),
                    "site_area":    g_site_area,
                    "duration":     g_duration,
                    "revision":     g_rev,
                    "prepared_by":  g_prep,
                    "approved_by":  g_approve,
                    "emp_date":     g_date,
                    "env_coord":    g_env_coord,
                }

                total = len(selected_asp) + 4
                step  = [0]

                def upd(msg):
                    step[0] += 1
                    prog.progress(
                        min(90,
                            int(step[0]/total*88)),
                        text=msg
                    )

                upd("Writing Environmental Policy...")
                upd("Building Legal Register...")

                emp_data = gen_emp_content(
                    client, meta,
                    emirate, selected_asp
                )

                for asp in selected_asp:
                    upd("Assessing: "
                        + asp[:40] + "...")

                upd("Building EMP document...")

                docx_buf = build_emp_docx(
                    emp_data, meta,
                    emirate, selected_asp
                )

                prog.progress(
                    100,
                    text="✅ EMP complete!"
                )

                st.success(
                    "✅ Environmental Management Plan "
                    "generated — "
                    + str(len(selected_asp))
                    + " aspects | 15 sections | "
                    "ISO 14001:2015 aligned!"
                )

                proj_fn = g_project.replace(' ','_')
                em_fn   = emirate.replace(' ','_')
                fname   = (
                    "EMP_"
                    + proj_fn + "_"
                    + em_fn + "_"
                    + str(g_date) + ".docx"
                )

                st.download_button(
                    label=(
                        "⬇️ Download Environmental "
                        "Management Plan (.docx)"
                    ),
                    data=docx_buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="emp_dl"
                )

                st.warning(
                    "📋 **MANDATORY NEXT STEPS "
                    "(ISO 14001:2015):**\n\n"
                    "1. ✅ Review with qualified "
                    "Environmental Professional\n\n"
                    "2. ✅ Submit to Consultant "
                    "(" + g_consultant + ") "
                    "for approval before works\n\n"
                    "3. ✅ Register with "
                    + REGULATIONS.get(
                        emirate, {}
                    ).get('Local Authority','authority')
                    + " if environmental permit "
                    "required\n\n"
                    "4. ✅ Brief ALL workers during "
                    "site induction\n\n"
                    "5. ✅ Start KPI monitoring "
                    "from Day 1 of site mobilisation\n\n"
                    "6. ✅ Review quarterly and "
                    "after any environmental incident\n\n"
                    "7. ✅ Conduct monthly "
                    "environmental inspections\n\n"
                    "⚠️ *Verify all monitoring limits "
                    "against current "
                    + REGULATIONS.get(
                        emirate,{}
                    ).get('Local Authority','')
                    + " standards before use.*"
                )

            except Exception as e:
                prog.progress(0)
                st.error("Error: " + str(e))


# ════════════════════════════════════════════════════════
# TAB 2 — ASPECTS REGISTER ONLY
# ════════════════════════════════════════════════════════

with tab_register:
    st.markdown(
        "### 📊 Environmental Aspects & "
        "Impacts Register"
    )
    st.info(
        "Generate a standalone Aspects & Impacts "
        "Register (ISO 14001:2015 Clause 6.1.2). "
        "Select your activities and the app "
        "builds the full register with significance "
        "scoring, controls, and monitoring "
        "requirements."
    )

    r_proj = st.text_input(
        "Project Name",
        placeholder="e.g. Bloom Living Almeria",
        key="reg_proj"
    )
    r_co   = st.text_input(
        "Company", value="ENGC",
        key="reg_co"
    )
    r_date = st.date_input(
        "Date", value=date.today(),
        key="reg_date"
    )

    st.divider()
    st.markdown(
        "#### Select Aspects to Include:"
    )
    r_selected = []
    r_cols     = st.columns(2)
    for i, asp in enumerate(
        list(ENV_ASPECTS.keys())
    ):
        col = r_cols[i % 2]
        if col.checkbox(
            asp, value=True,
            key="reg_asp_" + str(i)
        ):
            r_selected.append(asp)

    if st.button(
        "📊 Generate Aspects Register (.docx)",
        type="primary",
        use_container_width=True,
        key="reg_gen_btn"
    ):
        if not r_proj:
            st.error("Please enter a project name.")
        elif not r_selected:
            st.error(
                "Please select at least one aspect."
            )
        else:
            with st.spinner(
                "Building aspects register..."
            ):
                doc = Document()
                sec               = doc.sections[0]
                sec.page_width    = Cm(29.7)
                sec.page_height   = Cm(21.0)
                sec.left_margin   = Cm(1.5)
                sec.right_margin  = Cm(1.5)
                sec.top_margin    = Cm(1.5)
                sec.bottom_margin = Cm(1.5)

                # Header
                hdr           = doc.add_paragraph()
                hdr.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )
                shd_para(hdr, '1B5E20')
                hr = hdr.add_run(
                    "  🌿  ENVIRONMENTAL ASPECTS "
                    "& IMPACTS REGISTER  |  "
                    + r_proj.upper()
                    + "  |  " + emirate.upper()
                    + "  |  ISO 14001:2015  "
                )
                hr.bold           = True
                hr.font.size      = Pt(13)
                hr.font.color.rgb = RGBColor(
                    255, 255, 255
                )

                # Info row
                info_p = doc.add_paragraph()
                shd_para(info_p, '37474F')
                ir = info_p.add_run(
                    "  Company: " + r_co
                    + "  |  Date: " + str(r_date)
                    + "  |  Significance: "
                    "L×S Score — "
                    "LOW(1-2) | MEDIUM(3-6) | "
                    "HIGH(8-12) | SIGNIFICANT(>12)  "
                    "  |  ★ = Legal Requirement  "
                )
                ir.font.size      = Pt(9)
                ir.font.color.rgb = RGBColor(
                    255, 255, 200
                )

                doc.add_paragraph()

                # Main table — A3 Landscape
                asp_tbl       = doc.add_table(
                    rows=1, cols=9
                )
                asp_tbl.style = 'Table Grid'
                hdrs = [
                    ("No.",            Cm(0.7)),
                    ("Category",       Cm(2.5)),
                    ("Aspect",         Cm(3.5)),
                    ("Impact",         Cm(3.0)),
                    ("L\n(1-5)",       Cm(0.9)),
                    ("S\n(1-5)",       Cm(0.9)),
                    ("Score",          Cm(0.9)),
                    ("Significance",   Cm(1.8)),
                    ("Key Control",    Cm(7.5)),
                ]
                for (h, w), cell in zip(
                    hdrs, asp_tbl.rows[0].cells
                ):
                    cell.width = w
                    cell.text  = h
                    cell.paragraphs[0].runs[0]\
                        .bold = True
                    cell.paragraphs[0].runs[0]\
                        .font.size = Pt(8)
                    cell.paragraphs[0].runs[0]\
                        .font.color.rgb = RGBColor(
                        255, 255, 255
                    )
                    shd_cell(cell, '1B5E20')
                    cell_borders(cell, '1B5E20')
                    cell.paragraphs[0].alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                    )

                row_num = 1
                for asp_name in r_selected:
                    asp_data = ENV_ASPECTS.get(
                        asp_name, {}
                    )
                    color    = asp_data.get(
                        'color', '1B5E20'
                    )
                    aspects  = asp_data.get(
                        'aspects', []
                    )
                    controls = asp_data.get(
                        'controls', []
                    )

                    for i, (aspt, ctrl) in enumerate(
                        zip(
                            aspects,
                            controls + [''] * 20
                        )
                    ):
                        lik = 3
                        sev = 3
                        if "major" in aspt.lower() \
                           or "significant" in (
                               aspt.lower()
                           ):
                            lik = 4
                        if "contamination" in (
                            aspt.lower()
                        ) or "hazardous" in (
                            aspt.lower()
                        ):
                            sev = 4

                        scr = lik * sev
                        if scr > 12:
                            sig   = "SIGNIFICANT"
                            s_col = 'FF0000'
                        elif scr > 6:
                            sig   = "HIGH"
                            s_col = 'FFC000'
                        elif scr > 3:
                            sig   = "MEDIUM"
                            s_col = 'FFFF00'
                        else:
                            sig   = "LOW"
                            s_col = '00AF50'

                        row   = asp_tbl.add_row()
                        cells = row.cells
                        vals  = [
                            str(row_num),
                            asp_name.split(' ', 1)[
                                -1
                            ][:20],
                            aspt,
                            "Environmental impact",
                            str(lik), str(sev),
                            str(scr),
                            sig,
                            ctrl,
                        ]
                        widths = [
                            Cm(0.7), Cm(2.5),
                            Cm(3.5), Cm(3.0),
                            Cm(0.9), Cm(0.9),
                            Cm(0.9), Cm(1.8),
                            Cm(7.5),
                        ]
                        for j, (c, v, w) in enumerate(
                            zip(cells, vals, widths)
                        ):
                            c.width = w
                            c.text  = v
                            c.paragraphs[0]\
                                .runs[0]\
                                .font.size = Pt(7)
                            if j == 7:
                                shd_cell(c, s_col)
                                c.paragraphs[0]\
                                    .runs[0]\
                                    .bold = True
                                c.paragraphs[0]\
                                    .alignment = (
                                    WD_ALIGN_PARAGRAPH
                                    .CENTER
                                )
                            elif j in [4, 5, 6]:
                                c.paragraphs[0]\
                                    .alignment = (
                                    WD_ALIGN_PARAGRAPH
                                    .CENTER
                                )
                                if i % 2 == 0:
                                    shd_cell(
                                        c, 'E8F5E9'
                                    )
                            elif i % 2 == 0:
                                shd_cell(c, 'E8F5E9')
                            cell_borders(c, '1B5E20')

                        row_num += 1

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)

                proj_fn = r_proj.replace(' ','_')
                fname   = (
                    "ENV_Aspects_Register_"
                    + proj_fn + "_"
                    + str(r_date) + ".docx"
                )

                st.success(
                    "✅ Aspects Register generated — "
                    + str(row_num - 1)
                    + " aspects across "
                    + str(len(r_selected))
                    + " categories!"
                )

                st.download_button(
                    label=(
                        "⬇️ Download Aspects "
                        "Register (.docx)"
                    ),
                    data=buf,
                    file_name=fname,
                    mime=(
                        "application/vnd.openxmlformats"
                        "-officedocument"
                        ".wordprocessingml.document"
                    ),
                    key="reg_dl"
                )


# ════════════════════════════════════════════════════════
# TAB 3 — UPLOAD YOUR EMP FORMAT
# ════════════════════════════════════════════════════════

with tab_upload:
    st.markdown(
        "### 📤 Upload Your Own EMP Format"
    )
    st.info(
        "Upload your company's existing EMP "
        "template (Word or PDF). The AI reads "
        "your EXACT format and fills it with "
        "project-specific content including "
        + emirate + " regulatory requirements, "
        "environmental monitoring standards, "
        "and ISO 14001:2015 compliant content."
    )

    up_file = st.file_uploader(
        "📎 Upload EMP Template or Sample",
        type=["docx", "pdf"],
        help=(
            "Upload a blank EMP template or a "
            "previously completed EMP. "
            "Word (.docx) gives best results."
        ),
        key="emp_up_file"
    )

    if up_file:
        st.success(
            "✅ Uploaded: **"
            + up_file.name + "** ("
            + "{:,}".format(up_file.size)
            + " bytes)"
        )

        st.divider()
        st.subheader("📋 Project Details")

        u1, u2 = st.columns(2)
        with u1:
            u_project = st.text_input(
                "Project Name",
                placeholder="e.g. Bloom Living Almeria",
                key="u_emp_proj"
            )
            u_location = st.text_input(
                "Location",
                placeholder="e.g. Zayed City, AD",
                key="u_emp_loc"
            )
            u_company = st.text_input(
                "Company", value="ENGC",
                key="u_emp_co"
            )
            u_coord = st.text_input(
                "Environmental Coordinator",
                placeholder="e.g. Aftab Qamar",
                key="u_emp_coord"
            )
        with u2:
            u_type = st.text_input(
                "Project Type",
                placeholder=(
                    "e.g. Residential Villas"
                ),
                key="u_emp_type"
            )
            u_workers = st.number_input(
                "Peak Workers",
                min_value=5,
                max_value=5000,
                value=150,
                key="u_emp_workers"
            )
            u_date = st.date_input(
                "Issue Date",
                value=date.today(),
                key="u_emp_date"
            )
            u_sensit = st.text_input(
                "Sensitive receptors near site",
                placeholder="e.g. Residential, "
                            "Mosque, School",
                key="u_emp_sensit"
            )

        if st.button(
            "⚡ Fill My EMP Format",
            type="primary",
            use_container_width=True,
            key="u_emp_btn"
        ):
            if not u_project:
                st.error(
                    "Please enter a project name."
                )
            else:
                with st.spinner(
                    "Reading format and generating "
                    "EMP content..."
                ):
                    try:
                        client = anthropic.Anthropic(
                            api_key=os.getenv(
                                "ANTHROPIC_API_KEY"
                            )
                        )

                        regs_u = REGULATIONS.get(
                            emirate, {}
                        )
                        reg_str = "\n".join([
                            k + ": " + v
                            for k, v in regs_u.items()
                        ])

                        instr = (
                            "This is an Environmental "
                            "Management Plan template "
                            "for " + emirate + ", UAE.\n\n"
                            "Fill it for:\n"
                            "Project: " + u_project + "\n"
                            "Location: " + u_location + "\n"
                            "Company: " + u_company + "\n"
                            "Project Type: " + u_type + "\n"
                            "Env. Coordinator: "
                            + u_coord + "\n"
                            "Workers: " + str(u_workers) + "\n"
                            "Sensitive Receptors: "
                            + u_sensit + "\n"
                            "Date: " + str(u_date) + "\n\n"
                            + emirate.upper()
                            + " REGULATIONS:\n"
                            + reg_str + "\n\n"
                            "INSTRUCTIONS:\n"
                            "1. Keep EXACT same "
                            "structure and headings\n"
                            "2. Fill ALL sections with "
                            "specific content\n"
                            "3. Reference correct "
                            + emirate + " regulations\n"
                            "4. Include environmental "
                            "aspects, impacts, and "
                            "control measures\n"
                            "5. Use ISO 14001:2015 "
                            "terminology throughout\n"
                            "6. Include monitoring "
                            "standards from "
                            + regs_u.get(
                                'Local Authority', ''
                            )
                        )

                        file_bytes = up_file.read()

                        if up_file.name\
                                .lower()\
                                .endswith('.pdf'):
                            b64 = (
                                base64.standard_b64encode(
                                    file_bytes
                                ).decode()
                            )
                            response = (
                                client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "document",
                                            "source": {
                                                "type": "base64",
                                                "media_type": "application/pdf",
                                                "data": b64,
                                            }
                                        },
                                        {
                                            "type": "text",
                                            "text": instr
                                        }
                                    ]
                                }]
                            ))
                        else:
                            doc_tmp = Document(
                                BytesIO(file_bytes)
                            )
                            extracted = "\n".join([
                                p.text
                                for p in doc_tmp.paragraphs
                                if p.text.strip()
                            ])
                            tbl_txt = ""
                            for tbl in doc_tmp.tables:
                                for row in tbl.rows:
                                    rt = " | ".join([
                                        c.text.strip()
                                        for c in row.cells
                                        if c.text.strip()
                                    ])
                                    if rt:
                                        tbl_txt += (
                                            rt + "\n"
                                        )
                            full = (
                                extracted[:2500]
                                + (
                                    "\n\nTABLES:\n"
                                    + tbl_txt[:500]
                                    if tbl_txt else ""
                                )
                            )
                            response = (
                                client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=3000,
                                messages=[{
                                    "role": "user",
                                    "content": (
                                        "EMP TEMPLATE:\n\n"
                                        "---\n" + full
                                        + "\n---\n\n"
                                        + instr
                                    )
                                }]
                            ))

                        filled = (
                            response.content[0].text
                        )

                        st.success(
                            "✅ EMP filled for: "
                            + u_project + "!"
                        )

                        with st.expander(
                            "📄 Preview Generated EMP",
                            expanded=True
                        ):
                            st.markdown(filled)

                        proj_fn = u_project.replace(
                            ' ','_'
                        )
                        fname = (
                            "EMP_"
                            + proj_fn + "_"
                            + emirate.replace(' ','_')
                            + "_" + str(u_date)
                            + "_filled.txt"
                        )

                        st.download_button(
                            label=(
                                "⬇️ Download Filled "
                                "EMP (.txt)"
                            ),
                            data=filled,
                            file_name=fname,
                            mime="text/plain",
                            key="u_emp_dl"
                        )

                        st.info(
                            "💡 Copy content above "
                            "into your Word template "
                            "for the final formatted "
                            "EMP.\n\n"
                            "⚠️ Verify all monitoring "
                            "limits against current "
                            + REGULATIONS.get(
                                emirate, {}
                            ).get(
                                'Local Authority', ''
                            )
                            + " standards."
                        )

                    except Exception as e:
                        st.error(
                            "Error: " + str(e)
                        )
    else:
        ca, cb = st.columns(2)
        with ca:
            st.markdown("""
**✅ Accepted Formats:**
- Word Document (.docx) — Best
- PDF (.pdf)

**📋 What to Upload:**
- Your company EMP template
- Consultant-approved EMP format
- Previous project EMP to update
- Client-specific EMP format
""")
        with cb:
            st.markdown("""
**🌿 What the AI Does:**
- Fills correct """ + emirate + """ regulations
- Adds environmental aspects & impacts
- References ISO 14001:2015 clauses
- Includes monitoring standards & limits
- Keeps your EXACT format
""")

        st.divider()
        st.info(
            "🌿 **Don't have an EMP template?**\n\n"
            "Use the **Generate Full EMP** tab "
            "to create a complete, professional "
            "Environmental Management Plan from "
            "scratch — ISO 14001:2015 aligned, "
            "with legal register, aspects register, "
            "waste plan, monitoring programme, "
            "and KPI tracker."
        )