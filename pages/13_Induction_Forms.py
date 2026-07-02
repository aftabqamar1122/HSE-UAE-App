import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import json
import re

load_dotenv()

st.set_page_config(
    page_title="Worker Induction Forms",
    page_icon="👷"
)

emirate = st.session_state.get("emirate", None)
if not emirate:
    st.warning("Please go to Home and select your Emirate first.")
    st.stop()

flag = "🔵" if emirate == "Abu Dhabi" else "🔴"
st.title("👷 Worker Induction Forms & Register — " + flag + " " + emirate)

if emirate == "Abu Dhabi":
    st.markdown("**Reference: ADOSH CoP 53.1 | ADOSH-SF v4.0 | ADOSH CoP 1.0 | ISO 45001:2018 Clause 7.2 & 7.3 | UAE Federal Decree-Law No. 33/2021**")
else:
    st.markdown("**Reference: Dubai Municipality Code Art. 2.2.5 | ISO 45001:2018 Clause 7.2 & 7.3 | UAE Federal Decree-Law No. 33/2021**")

st.info(
    "👷 **Legal Requirement:** Every worker must receive a **site-specific safety induction** before starting work. "
    "This is mandatory under "
    + ("ADOSH CoP 53.1" if emirate == "Abu Dhabi" else "Dubai Municipality Code Art. 2.2.5")
    + " and ISO 45001:2018 Clause 7.3. Workers must sign the induction register to confirm understanding. "
    "**No signed induction = No entry to site.**"
)

st.divider()

# ── EMERGENCY NUMBERS ──────────────────────────────────────────────────────
EMERGENCY_NUMBERS = {
    "Abu Dhabi": "Ambulance: 998 | Police: 999 | Civil Defence: 997 | ADOSH: 800 60",
    "Dubai":     "Ambulance: 998 | Police: 999 | Civil Defence: 997 | DM: 800 900",
}

# ── DOCX HELPERS ───────────────────────────────────────────────────────────

def shd_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def shd_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def cell_borders(cell, color="1F3864"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement("w:" + side)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tcPr.append(borders)

def sec_hdr(doc, text, fill="1F3864", size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    shd_para(p, fill)
    run = p.add_run("  " + text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p

def add_body(doc, text, size=10, bold=False, indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    return p

# ── INDUCTION TOPICS LIBRARY ───────────────────────────────────────────────

def get_topics(emirate):
    midday_ban = (
        "NO outdoor work 12:30-15:00, 15 June to 15 September. MOHRE Resolution 44/2022."
        if emirate == "Abu Dhabi"
        else "NO outdoor work 12:30-15:00, 15 June to 15 September. MOHRE Resolution 44/2022."
    )
    excavation_rule = (
        "Never enter unshored excavation >0.5m. ADOSH CoP 29.0."
        if emirate == "Abu Dhabi"
        else "Never enter unshored excavation >1.25m. DM Code Chapter 7."
    )
    return {
        "🚨 Emergency Procedures": {
            "color": "C00000", "critical": True,
            "topics": [
                {"topic": "Emergency Numbers",
                 "detail": EMERGENCY_NUMBERS.get(emirate, ""),
                 "quiz": "What number do you call for an ambulance?"},
                {"topic": "Muster Points",
                 "detail": "Location of muster points, evacuation routes, assembly procedure",
                 "quiz": "Where is the nearest muster point?"},
                {"topic": "Raise the Alarm",
                 "detail": "Shout EMERGENCY, use radio, press alarm button",
                 "quiz": "What do you do first when you discover a fire?"},
                {"topic": "Stop Work Authority",
                 "detail": "Every worker has the RIGHT and DUTY to stop work if unsafe. No punishment allowed.",
                 "quiz": "Can you stop work if you think it is unsafe?"},
                {"topic": "First Aid Location",
                 "detail": "First aid room location, first aider names, first aid kit locations",
                 "quiz": "Who is the first aider for your work area?"},
            ],
        },
        "📋 Site Rules & Access": {
            "color": "1F3864", "critical": True,
            "topics": [
                {"topic": "Site Boundaries", "detail": "Authorised access points, restricted areas, no trespassing",
                 "quiz": "Are you allowed in restricted areas?"},
                {"topic": "Site ID & Pass", "detail": "Must carry site ID/pass at all times",
                 "quiz": "What must you carry on site at all times?"},
                {"topic": "Prohibited Items", "detail": "No alcohol, drugs, or mobile phones while operating plant",
                 "quiz": "Is alcohol allowed on site?"},
                {"topic": "Speed Limits", "detail": "15 km/h maximum on site — pedestrian routes always",
                 "quiz": "What is the site speed limit?"},
                {"topic": "Drug & Alcohol Policy", "detail": "Zero tolerance — random testing — immediate removal",
                 "quiz": "What happens if you are found with alcohol on site?"},
                {"topic": "Working Hours", "detail": "Authorised hours, rest breaks, overtime requirements",
                 "quiz": "Can you work outside authorised hours without approval?"},
            ],
        },
        "🦺 PPE Requirements": {
            "color": "1B5E20", "critical": True,
            "topics": [
                {"topic": "Mandatory PPE", "detail": "Safety helmet (EN 397), safety boots (EN ISO 20345), hi-vis vest — ALWAYS",
                 "quiz": "What PPE must be worn at all times on site?"},
                {"topic": "PPE Inspection", "detail": "Inspect PPE before use, report damaged PPE, replacement procedure",
                 "quiz": "What do you do if your PPE is damaged?"},
                {"topic": "Activity PPE", "detail": "Heights: full harness | Cutting: goggles + gloves | Chemicals: face shield",
                 "quiz": "What extra PPE is needed at heights?"},
                {"topic": "Respiratory Protection", "detail": "Dust masks for cutting, SCBA for confined spaces",
                 "quiz": "When must you wear a dust mask?"},
            ],
        },
        "☀️ Heat Stress & Welfare": {
            "color": "E65100", "critical": True,
            "topics": [
                {"topic": "Midday Work Ban", "detail": midday_ban,
                 "quiz": "Can you work outdoors at 1:00 PM in August?"},
                {"topic": "Heat Stroke Signs", "detail": "Dizziness, confusion, no sweating, hot red skin = EMERGENCY call 998",
                 "quiz": "What are the warning signs of heat stroke?"},
                {"topic": "Hydration", "detail": "Drink 250ml water every 20 minutes in heat — cool water provided",
                 "quiz": "How often should you drink water in the heat?"},
                {"topic": "Acclimatisation", "detail": "New workers: 50% work day 1, 80% day 3, full from day 5",
                 "quiz": "Can a new worker do a full day of outdoor work on Day 1?"},
                {"topic": "Rest Areas", "detail": "Shaded rest areas, cool water locations, welfare block",
                 "quiz": "Where is the nearest shaded rest area?"},
            ],
        },
        "⚠️ Hazard Awareness": {
            "color": "BF8F00", "critical": True,
            "topics": [
                {"topic": "Hazard Reporting", "detail": "Report ALL hazards immediately to supervisor or HSE Officer",
                 "quiz": "What should you do if you see a hazard?"},
                {"topic": "Excavation Safety", "detail": excavation_rule,
                 "quiz": "Can you enter an unshored excavation?"},
                {"topic": "Working at Heights", "detail": "Fall protection >1.8m, inspect harness, use anchor points",
                 "quiz": "At what height is fall protection required?"},
                {"topic": "Lifting Operations", "detail": "NEVER walk under suspended loads — exclusion zone required",
                 "quiz": "Is it safe to walk under a suspended load?"},
                {"topic": "Manual Handling", "detail": "Bend knees not back, max 25kg alone, use mechanical aids",
                 "quiz": "What is the safe maximum weight for one person to lift?"},
                {"topic": "Chemical Safety", "detail": "Read SDS/MSDS before use, correct PPE, no eating near chemicals",
                 "quiz": "What must you read before using chemicals?"},
            ],
        },
        "📝 Permits & Procedures": {
            "color": "4A148C", "critical": False,
            "topics": [
                {"topic": "Permit to Work", "detail": "Hot work, confined space, excavation, heights — ALL require PTW",
                 "quiz": "Do you need a permit to do hot work?"},
                {"topic": "Take 5 / Last-Minute Risk Assessment", "detail": "STOP-THINK-ACT before every task",
                 "quiz": "What is the Take 5 procedure?"},
                {"topic": "Incident Reporting", "detail": "Report ALL incidents IMMEDIATELY — injuries, near misses, damage",
                 "quiz": "Must you report a near miss?"},
                {"topic": "Toolbox Talks", "detail": "Daily briefing — attendance mandatory — sign the register",
                 "quiz": "Are toolbox talks mandatory?"},
            ],
        },
        "🌿 Environmental Rules": {
            "color": "2E7D32", "critical": False,
            "topics": [
                {"topic": "Waste Segregation", "detail": "General (black), Recyclable (blue), Hazardous (yellow) — NO mixing",
                 "quiz": "Which bin does plastic packaging go in?"},
                {"topic": "Spill Reporting", "detail": "Report ALL spills immediately to HSE Officer",
                 "quiz": "What do you do if you spill diesel?"},
                {"topic": "No Burning", "detail": "Burning waste is PROHIBITED — immediate fine",
                 "quiz": "Can you burn waste material on site?"},
                {"topic": "Drain Protection", "detail": "No concrete, chemicals, or oil to drains — immediate fine from authority",
                 "quiz": "Can you wash concrete into a drain?"},
            ],
        },
        "🏗️ Site-Specific Hazards": {
            "color": "37474F", "critical": False,
            "topics": [
                {"topic": "Overhead Power Lines", "detail": "3m minimum clearance — goal posts required",
                 "quiz": "What is the safe clearance from overhead power lines?"},
                {"topic": "Underground Services", "detail": "CAT scan before ALL digging — hand dig within 500mm",
                 "quiz": "What must you do before any digging on site?"},
                {"topic": "Scaffolding", "detail": "Green tag only = safe to use. Do not modify scaffold.",
                 "quiz": "What tag means scaffold is safe to use?"},
                {"topic": "Confined Spaces", "detail": "NEVER enter without PTW and atmospheric testing",
                 "quiz": "Can you enter a manhole without a permit?"},
                {"topic": "Housekeeping", "detail": "Keep work areas clean daily, clear walkways, no trip hazards",
                 "quiz": "Who is responsible for housekeeping?"},
            ],
        },
        "👤 Worker Rights": {
            "color": "5C4033", "critical": False,
            "topics": [
                {"topic": "Right to Safe Workplace", "detail": "UAE Labour Law 33/2021 — employer must provide safe workplace",
                 "quiz": "Who is responsible for providing a safe workplace?"},
                {"topic": "Right to Refuse Unsafe Work", "detail": "You cannot be punished for refusing genuinely unsafe work",
                 "quiz": "Can you be punished for refusing unsafe work?"},
                {"topic": "Grievance Procedure", "detail": "Supervisor → HSE Officer → Project Manager. No retaliation.",
                 "quiz": "Who do you report safety concerns to first?"},
                {"topic": "PPE is FREE", "detail": "PPE must be provided free. You must not be charged for PPE.",
                 "quiz": "Do you have to pay for your PPE?"},
            ],
        },
    }

# ── BUILD FUNCTIONS ────────────────────────────────────────────────────────

def build_induction_booklet(meta, selected_topics, emirate):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    nums = EMERGENCY_NUMBERS.get(emirate, "")

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(cover, "1F3864")
    cr = cover.add_run(
        "\n\n  👷  SITE SAFETY INDUCTION\n\n"
        "  WELCOME TO "
        + meta.get("project", "").upper()
        + "  \n\n"
    )
    cr.bold = True
    cr.font.size = Pt(22)
    cr.font.color.rgb = RGBColor(255, 255, 255)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(sub, "C00000")
    sr = sub.add_run(
        "\n  🚨 EMERGENCY: " + nums + "  \n\n"
        "  Muster Point: " + meta.get("muster", "Main Site Entrance") + "  \n"
    )
    sr.bold = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Info table
    info_tbl = doc.add_table(rows=0, cols=4)
    info_tbl.style = "Table Grid"
    def add_info(tbl, l1, v1, l2, v2):
        row = tbl.add_row()
        cells = row.cells
        for i, (l, v) in enumerate([(l1, v1), (l2, v2)]):
            lc = cells[i * 2]
            vc = cells[i * 2 + 1]
            lc.text = l
            lc.paragraphs[0].runs[0].bold = True
            lc.paragraphs[0].runs[0].font.size = Pt(10)
            shd_cell(lc, "E8F0FE")
            cell_borders(lc, "1F3864")
            vc.text = str(v) if v else "—"
            vc.paragraphs[0].runs[0].font.size = Pt(10)
            cell_borders(vc, "1F3864")

    add_info(info_tbl, "Project", meta.get("project",""), "Location", meta.get("location",""))
    add_info(info_tbl, "Main Contractor", meta.get("company","ENGC"), "Client", meta.get("client",""))
    add_info(info_tbl, "HSE Officer", meta.get("hse_officer",""), "Project Manager", meta.get("pm",""))
    add_info(info_tbl, "Version", meta.get("version","Rev 00"), "Date", str(meta.get("ind_date", date.today())))
    doc.add_paragraph()

    # Worker pledge
    pledge_p = doc.add_paragraph()
    shd_para(pledge_p, "1B5E20")
    pr = pledge_p.add_run("  WORKER SAFETY PLEDGE")
    pr.bold = True
    pr.font.size = Pt(12)
    pr.font.color.rgb = RGBColor(255, 255, 255)

    pledge_box = doc.add_paragraph()
    shd_para(pledge_box, "E8F5E9")
    pb = pledge_box.add_run(
        "\n  I confirm that I:\n\n"
        "  ✅  Have received this safety induction and understood the content\n"
        "  ✅  Will comply with all site safety rules at all times\n"
        "  ✅  Will wear required PPE at all times\n"
        "  ✅  Will report all hazards, incidents, and near misses immediately\n"
        "  ✅  Will not work under influence of alcohol or drugs\n"
        "  ✅  Understand and will exercise my Stop Work Authority if I see unsafe work\n"
        "  ✅  Will attend daily toolbox talks\n\n"
        "  Worker Name: ________________________  "
        "Signature: ________________________  "
        "Date: ____________\n\n"
        "  Trade: ________________________  "
        "Company: ________________________  "
        "Employee ID: ____________\n\n"
        "  Nationality: ____________  "
        "Mobile: ________________________  "
        "Language: ____________\n\n"
        "  Inducted By: ________________________  "
        "Designation: ____________  "
        "Signature: ________________________\n"
    )
    pb.font.size = Pt(10)
    doc.add_page_break()

    # All selected topic sections
    all_topics_lib = get_topics(emirate)
    all_questions = []

    for section_name in selected_topics:
        section = all_topics_lib.get(section_name, {})
        color = section.get("color", "1F3864")
        critical = section.get("critical", False)
        topics = section.get("topics", [])

        hdr_text = section_name.upper()
        if critical:
            hdr_text = "⭐ " + hdr_text + " ★ CRITICAL"
        sec_hdr(doc, hdr_text, fill=color, size=10)

        top_tbl = doc.add_table(rows=1, cols=3)
        top_tbl.style = "Table Grid"
        for cell, h in zip(top_tbl.rows[0].cells, ["Topic", "Key Points", "Quiz Question"]):
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            shd_cell(cell, color)
            cell_borders(cell, color)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, topic in enumerate(topics):
            row = top_tbl.add_row()
            vals = [topic.get("topic",""), topic.get("detail",""), "Q: " + topic.get("quiz","")]
            for j, (c, v) in enumerate(zip(row.cells, vals)):
                c.text = v
                c.paragraphs[0].runs[0].font.size = Pt(8)
                if j == 0:
                    c.paragraphs[0].runs[0].bold = True
                if j == 2:
                    shd_cell(c, "FFF9C4")
                elif i % 2 == 0:
                    shd_cell(c, "F5F5F5")
                cell_borders(c, color)
            if topic.get("quiz") and len(all_questions) < 20:
                all_questions.append({"num": len(all_questions)+1, "q": topic["quiz"], "hint": topic.get("topic","")})

        doc.add_paragraph()

    # Site-specific rules
    if meta.get("site_rules"):
        sec_hdr(doc, "SITE-SPECIFIC RULES — " + meta.get("project","").upper(), fill="37474F")
        for rule in meta["site_rules"]:
            add_bullet(doc, str(rule), size=9)
        doc.add_paragraph()

    # Competency test
    doc.add_page_break()
    sec_hdr(doc, "COMPETENCY TEST (PASS MARK: 80% — 16/20)", fill="C00000", size=11)
    add_body(doc, "Answer ALL questions. Score 80% (16/20) to pass. If you fail, you will be re-inducted before starting work.", size=9)
    doc.add_paragraph()

    # Fill to 20 questions
    extra_qs = [
        "What is the emergency number for ambulance?",
        "What is the site vehicle speed limit?",
        "How often should you drink water in heat?",
        "Must you report a near miss incident?",
        "Can you enter a confined space without a permit?",
        "What does a green scaffold tag mean?",
        "What is the maximum safe lift for one person?",
        "Can you be punished for refusing unsafe work?",
        "Where is the first aid room?",
        "What colour bin is for recyclable waste?",
    ]
    ei = 0
    while len(all_questions) < 20 and ei < len(extra_qs):
        all_questions.append({"num": len(all_questions)+1, "q": extra_qs[ei], "hint": "General Safety"})
        ei += 1

    test_tbl = doc.add_table(rows=1, cols=4)
    test_tbl.style = "Table Grid"
    for cell, h in zip(test_tbl.rows[0].cells, ["Q No.", "Question", "Your Answer", "✅/❌"]):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, "C00000")
        cell_borders(cell, "C00000")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, q in enumerate(all_questions[:20]):
        row = test_tbl.add_row()
        vals = [str(q["num"]), q["q"] + "\n(" + q["hint"] + ")", "", ""]
        widths = [Cm(1.0), Cm(10.5), Cm(4.5), Cm(1.0)]
        for j, (c, v, w) in enumerate(zip(row.cells, vals, widths)):
            c.width = w
            c.text = v
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                shd_cell(c, "FFF5F5")
            cell_borders(c, "C00000")

    doc.add_paragraph()
    res_p = doc.add_paragraph()
    shd_para(res_p, "1F3864")
    rr = res_p.add_run(
        "  TEST RESULT:  Score: _____ / 20  |  Pass (≥16) ☐   Fail (<16) ☐  |  "
        "Re-induction required ☐  |  Inducted By: _______________  Date: _______________"
    )
    rr.bold = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(255, 255, 255)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_induction_register(meta, emirate, rows=30):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(hdr, "1F3864")
    hr = hdr.add_run(
        "  👷  SITE INDUCTION ATTENDANCE REGISTER  |  "
        + meta.get("project","").upper()
        + "  |  " + meta.get("company","ENGC").upper()
        + "  |  " + emirate.upper()
    )
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = RGBColor(255, 255, 255)

    info_p = doc.add_paragraph()
    shd_para(info_p, "37474F")
    ir = info_p.add_run(
        "  HSE Officer: " + meta.get("hse_officer","___________")
        + "  |  Date: " + str(meta.get("ind_date", date.today()))
        + "  |  Location: " + meta.get("location","___________")
        + "  |  Session No.: ___________  |  Language: ___________"
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = RGBColor(255, 255, 200)
    doc.add_paragraph()

    reg_tbl = doc.add_table(rows=1, cols=9)
    reg_tbl.style = "Table Grid"
    hdrs = [
        ("No.", Cm(0.7)), ("Worker Name", Cm(3.8)), ("Trade / Role", Cm(2.8)),
        ("Company", Cm(2.8)), ("Nationality", Cm(2.0)), ("Employee ID", Cm(2.0)),
        ("Mobile No.", Cm(2.5)), ("Test Score\n(/20)", Cm(1.5)), ("Signature", Cm(8.6)),
    ]
    for (h, w), cell in zip(hdrs, reg_tbl.rows[0].cells):
        cell.width = w
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, "1F3864")
        cell_borders(cell, "1F3864")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    widths_list = [Cm(0.7), Cm(3.8), Cm(2.8), Cm(2.8), Cm(2.0), Cm(2.0), Cm(2.5), Cm(1.5), Cm(8.6)]
    for i in range(rows):
        row = reg_tbl.add_row()
        for j, (c, w) in enumerate(zip(row.cells, widths_list)):
            c.width = w
            c.text = str(i+1) if j == 0 else ""
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                shd_cell(c, "E8F0FE")
            cell_borders(c, "1F3864")

    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    shd_para(foot_p, "1B5E20")
    fr = foot_p.add_run(
        "  INDUCTING OFFICER DECLARATION: I confirm the above workers received and understood this site safety induction.  "
        "Name: _______________  Signature: _______________  Designation: _______________  Date: _______________"
    )
    fr.bold = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(255, 255, 255)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_ppe_register(meta, emirate, rows=25):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(hdr, "1B5E20")
    hr = hdr.add_run(
        "  🦺  PPE ISSUE & INSPECTION REGISTER  |  "
        + meta.get("project","").upper()
        + "  |  " + emirate.upper()
    )
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = RGBColor(255, 255, 255)

    info_p = doc.add_paragraph()
    shd_para(info_p, "2E7D32")
    ir = info_p.add_run(
        "  ADOSH CoP 2.0 — PPE Requirements  |  Company: " + meta.get("company","ENGC")
        + "  |  HSE Officer: " + meta.get("hse_officer","___________")
        + "  |  Date: " + str(meta.get("ind_date", date.today()))
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = RGBColor(255, 255, 200)
    doc.add_paragraph()

    ppe_tbl = doc.add_table(rows=1, cols=10)
    ppe_tbl.style = "Table Grid"
    ppe_hdrs = [
        ("No.", Cm(0.6)), ("Worker Name", Cm(3.2)), ("Employee ID", Cm(1.8)),
        ("Safety\nHelmet", Cm(1.6)), ("Safety\nBoots", Cm(1.6)), ("Hi-Vis\nVest", Cm(1.6)),
        ("Safety\nGloves", Cm(1.6)), ("Safety\nGoggles", Cm(1.6)),
        ("Other PPE\n(specify)", Cm(4.0)), ("Signature", Cm(8.9)),
    ]
    for (h, w), cell in zip(ppe_hdrs, ppe_tbl.rows[0].cells):
        cell.width = w
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(7)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, "1B5E20")
        cell_borders(cell, "1B5E20")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ppe_widths = [Cm(0.6), Cm(3.2), Cm(1.8), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.6), Cm(4.0), Cm(8.9)]
    for i in range(rows):
        row = ppe_tbl.add_row()
        for j, (c, w) in enumerate(zip(row.cells, ppe_widths)):
            c.width = w
            c.text = str(i+1) if j == 0 else ""
            c.paragraphs[0].runs[0].font.size = Pt(7)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                shd_cell(c, "E8F5E9")
            cell_borders(c, "1B5E20")

    doc.add_paragraph()
    note_p = doc.add_paragraph()
    shd_para(note_p, "1F3864")
    nr = note_p.add_run(
        "  ⚠️ PPE is provided FREE of charge. Workers must not be charged for PPE. "
        "Worn or damaged PPE must be replaced immediately. Ref: ADOSH CoP 2.0 | UAE Labour Law 33/2021 Art. 15"
    )
    nr.bold = True
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(255, 255, 200)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_worker_reg_form(meta, emirate):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(hdr, "1F3864")
    hr = hdr.add_run(
        "\n  👷  WORKER REGISTRATION & COMPETENCY RECORD  \n  "
        + meta.get("company","ENGC")
        + "  |  " + meta.get("project","")
        + "  |  " + emirate + "  \n"
    )
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()

    # Personal details
    sec_hdr(doc, "PERSONAL DETAILS", fill="1F3864")
    per_tbl = doc.add_table(rows=0, cols=4)
    per_tbl.style = "Table Grid"
    personal = [
        ("Full Name", "", "Date of Birth", ""),
        ("Nationality", "", "Passport / Emirates ID", ""),
        ("Visa Status", "", "Visa Expiry", ""),
        ("Trade / Occupation", "", "Company", ""),
        ("Employee ID / Badge No.", "", "Mobile Number", ""),
        ("Emergency Contact Name", "", "Emergency Contact No.", ""),
        ("Languages Spoken", "", "Blood Type", ""),
    ]
    for l1, v1, l2, v2 in personal:
        row = per_tbl.add_row()
        cells = row.cells
        for i, (l, v) in enumerate([(l1,v1),(l2,v2)]):
            lc = cells[i*2]
            vc = cells[i*2+1]
            lc.text = l
            lc.paragraphs[0].runs[0].bold = True
            lc.paragraphs[0].runs[0].font.size = Pt(9)
            shd_cell(lc, "E8F0FE")
            cell_borders(lc, "1F3864")
            vc.text = str(v) if v else ""
            vc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_borders(vc, "1F3864")
    doc.add_paragraph()

    # Certifications
    sec_hdr(doc, "QUALIFICATIONS & CERTIFICATIONS", fill="1F5864")
    qual_tbl = doc.add_table(rows=1, cols=5)
    qual_tbl.style = "Table Grid"
    for cell, h in zip(qual_tbl.rows[0].cells, ["Certificate / Licence", "Issuing Authority", "Certificate No.", "Issue Date", "Expiry Date"]):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, "1F5864")
        cell_borders(cell, "1F5864")

    certs = ["Driving Licence", "Crane / Plant Operator Licence", "Scaffolding Certificate",
             "Working at Heights Certificate", "Confined Space Certificate", "First Aid Certificate",
             "Hot Work / Welding Certificate", "Electrical Licence", "HSE Certificate (NEBOSH/IOSH)", "Other"]
    for i, cert in enumerate(certs):
        row = qual_tbl.add_row()
        row.cells[0].text = cert
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for j in range(1,5):
            row.cells[j].text = ""
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(8)
        if i % 2 == 0:
            for c in row.cells:
                shd_cell(c, "EEF6FF")
        for c in row.cells:
            cell_borders(c, "1F5864")
    doc.add_paragraph()

    # Induction record
    sec_hdr(doc, "INDUCTION & TRAINING RECORD", fill="1B5E20")
    ind_tbl = doc.add_table(rows=1, cols=5)
    ind_tbl.style = "Table Grid"
    for cell, h in zip(ind_tbl.rows[0].cells, ["Training / Induction", "Date", "Inducted By", "Test Score", "Pass/Fail"]):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, "1B5E20")
        cell_borders(cell, "1B5E20")

    ind_items = ["Site Safety Induction (General)", "Environmental Induction",
                 "Emergency Response Briefing", "Activity-Specific Induction",
                 "Annual Refresher Training"]
    for i, item in enumerate(ind_items):
        row = ind_tbl.add_row()
        row.cells[0].text = item
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for j in range(1,5):
            row.cells[j].text = ""
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
        if i % 2 == 0:
            for c in row.cells:
                shd_cell(c, "E8F5E9")
        for c in row.cells:
            cell_borders(c, "1B5E20")
    doc.add_paragraph()

    # Declaration
    sec_hdr(doc, "WORKER DECLARATION & SIGNATURE", fill="1F3864")
    decl_p = doc.add_paragraph()
    shd_para(decl_p, "E8F0FE")
    dp = decl_p.add_run(
        "\n  I confirm that:\n"
        "  ✅ The information I have provided is true and accurate\n"
        "  ✅ I have received and understood the site safety induction\n"
        "  ✅ I will comply with all site safety rules\n"
        "  ✅ I will exercise my Stop Work Authority if I see unsafe conditions\n"
        "  ✅ I understand I may be removed from site for safety violations\n\n"
        "  Worker Signature: __________________________  Date: _______________\n\n"
        "  Registered By: __________________________  Signature: _______________  Date: _______________\n"
    )
    dp.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_tbt_register(meta, emirate, rows=25):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd_para(hdr, "BF8F00")
    hr = hdr.add_run(
        "  🗣️  TOOLBOX TALK ATTENDANCE REGISTER  |  "
        + meta.get("project","").upper()
        + "  |  " + emirate.upper()
    )
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = RGBColor(255, 255, 255)

    info_p = doc.add_paragraph()
    shd_para(info_p, "37474F")
    ir = info_p.add_run(
        "  TBT Topic: ___________________________"
        "  |  Date: " + str(date.today())
        + "  |  Time: ___________"
        + "  |  Location: ___________"
        + "  |  Conducted By: ___________"
        + "  |  Duration: ___ mins"
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = RGBColor(255, 255, 200)
    doc.add_paragraph()

    top_p = doc.add_paragraph()
    shd_para(top_p, "F5F5F5")
    tp = top_p.add_run(
        "  Key Points Covered:\n"
        "  1. ___________________________________"
        "  2. ___________________________________"
        "  3. ___________________________________"
    )
    tp.font.size = Pt(9)
    doc.add_paragraph()

    tbt_tbl = doc.add_table(rows=1, cols=7)
    tbt_tbl.style = "Table Grid"
    tbt_hdrs = [
        ("No.", Cm(0.7)), ("Worker Name", Cm(4.5)), ("Trade", Cm(2.8)),
        ("Company", Cm(2.8)), ("Employee ID", Cm(2.0)),
        ("Language\nUnderstood", Cm(2.0)), ("Signature", Cm(12.2)),
    ]
    for (h, w), cell in zip(tbt_hdrs, tbt_tbl.rows[0].cells):
        cell.width = w
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        shd_cell(cell, "BF8F00")
        cell_borders(cell, "BF8F00")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbt_widths = [Cm(0.7), Cm(4.5), Cm(2.8), Cm(2.8), Cm(2.0), Cm(2.0), Cm(12.2)]
    for i in range(rows):
        row = tbt_tbl.add_row()
        for j, (c, w) in enumerate(zip(row.cells, tbt_widths)):
            c.width = w
            c.text = str(i+1) if j == 0 else ""
            c.paragraphs[0].runs[0].font.size = Pt(8)
            if j == 0:
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                shd_cell(c, "FFFDE7")
            cell_borders(c, "BF8F00")

    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    shd_para(foot_p, "BF8F00")
    fr = foot_p.add_run(
        "  Conducted By: _______________  Designation: _______________  Signature: _______________  "
        "Date: _______________  | Workers Present: ___  | Language(s): _______________"
    )
    fr.bold = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(255, 255, 255)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── SHARED FORM HELPER ─────────────────────────────────────────────────────

def project_form(prefix):
    c1, c2 = st.columns(2)
    with c1:
        proj = st.text_input("Project Name *", placeholder="e.g. Bloom Living Almeria", key=prefix+"_proj")
        loc = st.text_input("Site Location", placeholder="e.g. Zayed City, Abu Dhabi", key=prefix+"_loc")
        co = st.text_input("Main Contractor", value="ENGC", key=prefix+"_co")
        client = st.text_input("Client", placeholder="e.g. Bloom District", key=prefix+"_client")
    with c2:
        hse = st.text_input("HSE Officer Name", placeholder="e.g. Aftab Qamar", key=prefix+"_hse")
        pm = st.text_input("Project Manager", placeholder="e.g. Ahmad Abdelrahman", key=prefix+"_pm")
        ind_date = st.date_input("Date", value=date.today(), key=prefix+"_date")
        version = st.selectbox("Version", ["Rev 00","Rev 01","Rev 02"], key=prefix+"_ver")
    return {"project": proj, "location": loc, "company": co, "client": client,
            "hse_officer": hse, "pm": pm, "ind_date": ind_date, "version": version}


# ── MAIN UI ────────────────────────────────────────────────────────────────

tab_ind, tab_reg, tab_ppe, tab_worker, tab_tbt = st.tabs([
    "👷 Site Induction Booklet",
    "📋 Induction Register",
    "🦺 PPE Issue Register",
    "📁 Worker Registration Card",
    "🗣️ Toolbox Talk Register",
])

# ── TAB 1 — INDUCTION BOOKLET ─────────────────────────────────────────────
with tab_ind:
    st.markdown("### 👷 Generate Site Safety Induction Booklet")
    st.info(
        "Complete site-specific induction booklet with topic tables, quiz questions, "
        "20-question competency test, and worker pledge. Compliant with "
        + ("ADOSH CoP 53.1" if emirate == "Abu Dhabi" else "DM Code Art. 2.2.5")
        + " and ISO 45001:2018 Clause 7.3."
    )
    meta_ind = project_form("ind")
    st.divider()
    st.subheader("📋 Additional Site Information")
    ia, ib = st.columns(2)
    with ia:
        muster = st.text_input("Primary Muster Point Location", placeholder="e.g. Main site entrance — near security hut", key="ind_muster")
        first_aider = st.text_input("First Aider Name", placeholder="e.g. Gurpal Singh", key="ind_fa")
    with ib:
        hospital = st.text_input("Nearest Hospital", placeholder="e.g. Sheikh Khalifa Medical City", key="ind_hosp")
        site_speed = st.text_input("Site Speed Limit", value="15 km/h", key="ind_speed")

    st.subheader("📌 Site-Specific Rules")
    rules_raw = st.text_area("Site Rules (one per line)",
        value="No mobile phones while operating plant\nAll workers must sign in/out daily\nNo jewellery on site — entanglement risk\nVisitors must be escorted at all times\nClean up work area before leaving",
        height=100, key="ind_rules")

    st.divider()
    st.subheader("📚 Select Induction Topics")
    all_topics_lib = get_topics(emirate)
    ind_selected = []
    ind_cols = st.columns(2)
    for i, (topic_name, topic_data) in enumerate(all_topics_lib.items()):
        col = ind_cols[i % 2]
        critical = topic_data.get("critical", False)
        label = ("⭐ " + topic_name if critical else topic_name)
        if col.checkbox(label, value=True, key="ind_topic_" + str(i)):
            ind_selected.append(topic_name)

    if st.button("👷 Generate Induction Booklet (.docx)", type="primary", use_container_width=True, key="ind_gen_btn"):
        if not meta_ind.get("project"):
            st.error("Please enter a project name.")
        elif not ind_selected:
            st.error("Please select at least one topic.")
        else:
            with st.spinner("Building induction booklet..."):
                meta_ind["muster"] = muster
                meta_ind["first_aider"] = first_aider
                meta_ind["hospital"] = hospital
                meta_ind["site_speed"] = site_speed
                meta_ind["site_rules"] = [r.strip() for r in rules_raw.split("\n") if r.strip()]
                docx_buf = build_induction_booklet(meta_ind, ind_selected, emirate)
                proj_fn = meta_ind["project"].replace(" ","_")
                fname = "Induction_Booklet_" + proj_fn + "_" + str(meta_ind["ind_date"]) + ".docx"
                st.success("✅ Induction Booklet generated — " + str(len(ind_selected)) + " topic sections + 20-question competency test + worker pledge!")
                st.download_button(
                    label="⬇️ Download Site Induction Booklet (.docx)",
                    data=docx_buf, file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="ind_dl"
                )
                st.warning(
                    "📋 **BEFORE USING THIS BOOKLET:**\n\n"
                    "1. ✅ Review ALL content with your HSE team\n"
                    "2. ✅ Verify emergency numbers and muster points\n"
                    "3. ✅ Get sign-off from Project Manager\n"
                    "4. ✅ Translate into workers' languages\n"
                    "5. ✅ Conduct induction verbally — not just hand out booklet\n"
                    "6. ✅ 80% pass mark required on competency test\n\n"
                    "⚠️ *AI-generated — must be verified by a competent HSE professional.*"
                )

# ── TAB 2 — INDUCTION REGISTER ────────────────────────────────────────────
with tab_reg:
    st.markdown("### 📋 Site Induction Attendance Register")
    st.info(
        "A4 Landscape — records worker name, trade, company, nationality, employee ID, mobile, test score, and signature. "
        "Required under " + ("ADOSH CoP 53.1" if emirate == "Abu Dhabi" else "DM Code Art. 2.2.5") + " — keep for minimum 3 years."
    )
    meta_reg = project_form("reg")
    reg_rows = st.slider("Number of rows (workers)", 10, 50, 30, key="reg_rows")
    if st.button("📋 Generate Induction Register (.docx)", type="primary", use_container_width=True, key="reg_gen_btn"):
        if not meta_reg.get("project"):
            st.error("Please enter a project name.")
        else:
            with st.spinner("Building register..."):
                docx_buf = build_induction_register(meta_reg, emirate, reg_rows)
                proj_fn = meta_reg["project"].replace(" ","_")
                fname = "Induction_Register_" + proj_fn + "_" + str(meta_reg["ind_date"]) + ".docx"
                st.success("✅ Induction Register — " + str(reg_rows) + " rows ready!")
                st.download_button(label="⬇️ Download Induction Register (.docx)", data=docx_buf, file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="reg_dl")

# ── TAB 3 — PPE REGISTER ──────────────────────────────────────────────────
with tab_ppe:
    st.markdown("### 🦺 PPE Issue & Inspection Register")
    st.info(
        "Records PPE issued to each worker. A4 Landscape. PPE must be provided FREE under UAE Labour Law 33/2021. "
        "Reference: " + ("ADOSH CoP 2.0 — PPE Requirements" if emirate == "Abu Dhabi" else "DM Code Chapter 5 — PPE")
    )
    meta_ppe = project_form("ppe")
    ppe_rows = st.slider("Number of rows (workers)", 10, 50, 25, key="ppe_rows")
    if st.button("🦺 Generate PPE Register (.docx)", type="primary", use_container_width=True, key="ppe_gen_btn"):
        if not meta_ppe.get("project"):
            st.error("Please enter a project name.")
        else:
            with st.spinner("Building PPE register..."):
                docx_buf = build_ppe_register(meta_ppe, emirate, ppe_rows)
                proj_fn = meta_ppe["project"].replace(" ","_")
                fname = "PPE_Register_" + proj_fn + "_" + str(meta_ppe["ind_date"]) + ".docx"
                st.success("✅ PPE Register — " + str(ppe_rows) + " rows ready!")
                st.download_button(label="⬇️ Download PPE Register (.docx)", data=docx_buf, file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="ppe_dl")

# ── TAB 4 — WORKER REGISTRATION CARD ─────────────────────────────────────
with tab_worker:
    st.markdown("### 📁 Worker Registration & Competency Record")
    st.info(
        "Individual worker profile card — personal details, qualifications, certifications, medical fitness, "
        "induction history, and safety violation record. A4 Portrait. One form per worker. Required for ADOSH/DM audits."
    )
    meta_wrk = project_form("wrk")
    if st.button("📁 Generate Worker Registration Card (.docx)", type="primary", use_container_width=True, key="wrk_gen_btn"):
        if not meta_wrk.get("project"):
            st.error("Please enter a project name.")
        else:
            with st.spinner("Building worker registration card..."):
                docx_buf = build_worker_reg_form(meta_wrk, emirate)
                proj_fn = meta_wrk["project"].replace(" ","_")
                fname = "Worker_Registration_Card_" + proj_fn + "_" + str(meta_wrk["ind_date"]) + ".docx"
                st.success("✅ Worker Registration Card generated!")
                st.download_button(label="⬇️ Download Worker Registration Card (.docx)", data=docx_buf, file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="wrk_dl")
                st.info("💡 Print one card per worker and keep in site HSE files. Bring out during authority inspections.")

# ── TAB 5 — TOOLBOX TALK REGISTER ────────────────────────────────────────
with tab_tbt:
    st.markdown("### 🗣️ Toolbox Talk Attendance Register")
    st.info(
        "Daily toolbox talk attendance register — A4 Landscape. "
        + ("Mandatory daily under ADOSH CoP 53.1." if emirate == "Abu Dhabi"
           else "Every 15 days minimum per DM Code Art. 2.2.8 (daily recommended).")
        + " Records must be kept for 3 years."
    )
    meta_tbt = project_form("tbt")
    tbt_rows = st.slider("Number of rows (workers)", 10, 50, 25, key="tbt_rows")
    if st.button("🗣️ Generate Toolbox Talk Register (.docx)", type="primary", use_container_width=True, key="tbt_gen_btn"):
        if not meta_tbt.get("project"):
            st.error("Please enter a project name.")
        else:
            with st.spinner("Building toolbox talk register..."):
                docx_buf = build_tbt_register(meta_tbt, emirate, tbt_rows)
                proj_fn = meta_tbt["project"].replace(" ","_")
                fname = "TBT_Register_" + proj_fn + "_" + str(date.today()) + ".docx"
                st.success("✅ Toolbox Talk Register — " + str(tbt_rows) + " rows ready!")
                st.download_button(label="⬇️ Download Toolbox Talk Register (.docx)", data=docx_buf, file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="tbt_dl")

    st.divider()
    st.markdown("#### 🗓️ Suggested Weekly TBT Schedule")
    schedule = {
        "Sunday": "Heat Stress & Hydration (MOHRE Res. 44/2022)",
        "Monday": "Working at Heights & Scaffold Safety",
        "Tuesday": "Excavation & Underground Services",
        "Wednesday": "Electrical Safety & LOTO",
        "Thursday": "Lifting Operations & Rigging Safety",
        "Saturday": "Fire Safety & Emergency Procedures",
    }
    sch_cols = st.columns(3)
    for i, (day, topic) in enumerate(schedule.items()):
        sch_cols[i % 3].markdown("**" + day + ":** " + topic)